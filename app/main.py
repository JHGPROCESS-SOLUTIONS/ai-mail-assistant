import os
import json
import asyncio
import base64
import re
from datetime import datetime, timezone, timedelta
from email.mime.text import MIMEText
from urllib.parse import quote
from typing import Any, Optional

import httpx
import jwt
import stripe
from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException, Request, Body, Header, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import RedirectResponse, JSONResponse
from pydantic import BaseModel

from app.billing import router as billing_router

load_dotenv()

app = FastAPI(
    title="AI Mail Assistant API",
    docs_url="/docs",
    redoc_url="/redoc",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "https://officeflowcompany.com",
        "https://www.officeflowcompany.com",
        # Local dev only — remove for production-only deploys
        "http://localhost:3000",
        "http://127.0.0.1:5500",
    ],
    allow_credentials=True,
    allow_methods=["GET", "POST", "PATCH", "DELETE", "OPTIONS"],
    allow_headers=[
        "Authorization",
        "Content-Type",
        "X-Internal-Secret",
        "X-Requested-With",
    ],
    max_age=3600,
)

app.include_router(billing_router)

GOOGLE_CLIENT_ID = os.getenv("GOOGLE_CLIENT_ID")
GOOGLE_CLIENT_SECRET = os.getenv("GOOGLE_CLIENT_SECRET")
GOOGLE_REDIRECT_URI = os.getenv("GOOGLE_REDIRECT_URI")

SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_SERVICE_ROLE_KEY = os.getenv("SUPABASE_SERVICE_ROLE_KEY")
SUPABASE_JWT_SECRET = os.getenv("SUPABASE_JWT_SECRET")

# Internal-only API secret. /internal/* endpoints require the header
# `X-Internal-Secret: <value>` matching this env var. Used by cron jobs,
# Railway-scheduled tasks, and admin scripts. NEVER expose in frontend.
INTERNAL_API_SECRET = os.getenv("INTERNAL_API_SECRET")

# Redirect after first-time password setup via invite email
SUPABASE_SET_PASSWORD_URL = os.getenv(
    "SUPABASE_SET_PASSWORD_URL",
    "https://officeflowcompany.com/set-password.html",
)

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

FRONTEND_SUCCESS_URL = os.getenv(
    "FRONTEND_SUCCESS_URL",
    "https://officeflowcompany.com/onboarding/success",
)
FRONTEND_PRICING_URL = os.getenv(
    "FRONTEND_PRICING_URL",
    "https://officeflowcompany.com/pricing",
)

STRIPE_SECRET_KEY = os.getenv("STRIPE_SECRET_KEY")
STRIPE_WEBHOOK_SECRET = os.getenv("STRIPE_WEBHOOK_SECRET")
STRIPE_BILLING_PORTAL_RETURN_URL = os.getenv(
    "STRIPE_BILLING_PORTAL_RETURN_URL",
    "https://officeflowcompany.com/settings",
)

AUTO_PROCESS_ENABLED = os.getenv("AUTO_PROCESS_ENABLED", "true").lower() == "true"
AUTO_PROCESS_INTERVAL_SECONDS = int(os.getenv("AUTO_PROCESS_INTERVAL_SECONDS", "60"))
AUTO_PROCESS_MAX_RESULTS = int(os.getenv("AUTO_PROCESS_MAX_RESULTS", "10"))
AUTO_PROCESS_MAX_CONCURRENCY = int(os.getenv("AUTO_PROCESS_MAX_CONCURRENCY", "10"))

GMAIL_SCOPE = "openid email profile https://www.googleapis.com/auth/gmail.modify"
GMAIL_API_BASE = "https://gmail.googleapis.com/gmail/v1/users/me"
ALLOWED_SUBSCRIPTION_STATUSES = {"active", "trialing", "canceling"}

LABELS = [
    "Priority",
    "To Respond",
    "Waiting On Reply",
    "Follow Up",
    "Done",
    "FYI",
    "Notification",
    "Marketing",
    "Ignore",
]

LEGACY_LABEL_NAME_MAP = {
    "OfficeFlow/Priority": "Priority",
    "OfficeFlow/To Respond": "To Respond",
    "OfficeFlow/Waiting On Reply": "Waiting On Reply",
    "OfficeFlow/Done": "Done",
    "OfficeFlow/FYI": "FYI",
    "OfficeFlow/Notification": "Notification",
    "OfficeFlow/Marketing": "Marketing",
    "OfficeFlow/Spam": "Ignore",
}

LABEL_RULES = {
    "Priority": {"generate_draft": True},
    "To Respond": {"generate_draft": True},
    "Waiting On Reply": {"generate_draft": False},
    "Follow Up": {"generate_draft": False},
    "Done": {"generate_draft": False},
    "FYI": {"generate_draft": False},
    "Notification": {"generate_draft": False},
    "Marketing": {"generate_draft": False},
    "Ignore": {"generate_draft": False},
}

CLASSIFIER_LABELS = {
    "Priority",
    "To Respond",
    "FYI",
    "Notification",
    "Marketing",
    "Ignore",
}

FOLLOW_UP_CLASSIFIER_LABELS = {
    "Priority",
    "To Respond",
    "Waiting On Reply",
    "Done",
    "FYI",
    "Notification",
    "Ignore",
}

SENT_REPLY_STATUS_LABELS = {
    "Waiting On Reply",
    "Follow Up",
    "Done",
}

# Safe set — labels that are almost never important inbox material.
# When a mail is classified into one of these AND the mailbox has
# auto_archive_low_value=true, OfficeFlow removes the INBOX label so the
# thread disappears from Inbox but stays findable via its status label.
LOW_VALUE_LABELS = {
    "Marketing",
    "Ignore",
}

# Snooze label — applied when a user temporarily removes a mail from Inbox.
# Lives under the OfficeFlow/ namespace because it's a workflow label (not
# a classification label) and we want it visually grouped with our other
# meta-labels in Gmail.
SNOOZED_LABEL = "OfficeFlow/Snoozed"

# Notification is opt-in (notification_auto_archive on mailbox) because
# notifications are often transactional (banking, delivery, 2FA, tax) and
# users may want to keep seeing them in their primary inbox.
NOTIFICATION_AUTO_ARCHIVE_LABEL = "Notification"

# Labels that mark a sender as trusted — any prior mail from a sender with
# one of these labels means we never auto-archive further mail from them.
TRUSTED_INDICATOR_LABELS = (
    "Priority",
    "To Respond",
    "Follow Up",
)

# -----------------------------------------------------------------------------
# Classifier confidence (internal signal only — NO Gmail labels).
# The classifier returns "high" / "medium" / "low" and we surface it in the
# API response so the dashboard can render it later if we want. We do NOT
# apply it as a Gmail label because that clutters the inbox, and the
# classifier is accurate enough that the label adds more noise than value.
#
# CONFIDENCE_LEGACY_LABEL_NAMES lists Gmail labels previously created by
# older builds — setup_gmail_labels_for_mailbox deletes them on the next
# inbox run so old mails lose the tag automatically.
# -----------------------------------------------------------------------------
CONFIDENCE_VALUES = ("high", "medium", "low")

CONFIDENCE_LEGACY_LABEL_NAMES = ("AI · Sure", "AI · Check", "AI · Review")

LABEL_COLORS = {
    "Priority": {
        "textColor": "#ffffff",
        "backgroundColor": "#cc3a21",
    },
    "To Respond": {
        "textColor": "#ffffff",
        "backgroundColor": "#3c78d8",
    },
    "Waiting On Reply": {
        "textColor": "#ffffff",
        "backgroundColor": "#8e63ce",
    },
    "Follow Up": {
        "textColor": "#000000",
        "backgroundColor": "#f6bf26",
    },
    "Done": {
        "textColor": "#ffffff",
        "backgroundColor": "#16a766",
    },
    "FYI": {
        "textColor": "#000000",
        "backgroundColor": "#f3f3f3",
    },
    "Notification": {
        "textColor": "#ffffff",
        "backgroundColor": "#8e63ce",
    },
    "Marketing": {
        "textColor": "#000000",
        "backgroundColor": "#fad165",
    },
    "Ignore": {
        "textColor": "#ffffff",
        "backgroundColor": "#822111",
    },
}

LANGUAGE_NAME_MAP = {
    "nl": "Dutch",
    "nederlands": "Dutch",
    "dutch": "Dutch",
    "en": "English",
    "engels": "English",
    "english": "English",
    "de": "German",
    "duits": "German",
    "german": "German",
    "fr": "French",
    "frans": "French",
    "french": "French",
    "es": "Spanish",
    "spaans": "Spanish",
    "spanish": "Spanish",
    "it": "Italian",
    "italiaans": "Italian",
    "italian": "Italian",
    "pt": "Portuguese",
    "portugees": "Portuguese",
    "portuguese": "Portuguese",
}

if STRIPE_SECRET_KEY:
    stripe.api_key = STRIPE_SECRET_KEY


class PromptSettingsPayload(BaseModel):
    email: str
    preferred_language: str | None = None
    tone_preference: str | None = None
    formality: str | None = None
    length_preference: str | None = None
    emoji_preference: bool | None = None
    cta_preference: str | None = None
    signature_mode: str | None = None
    signature_text: str | None = None
    forbidden_phrases: list[str] | str | None = None
    preferred_phrases: list[str] | str | None = None
    custom_instructions: str | None = None
    style_learning_enabled: bool = False
    style_learning_source_limit: int = 20


class PromptSettingsBody(BaseModel):
    """JWT-auth variant: same fields as PromptSettingsPayload but no email
    (it comes from the bearer token via get_current_user)."""
    preferred_language: str | None = None
    tone_preference: str | None = None
    formality: str | None = None
    length_preference: str | None = None
    emoji_preference: bool | None = None
    cta_preference: str | None = None
    signature_mode: str | None = None
    signature_text: str | None = None
    forbidden_phrases: list[str] | str | None = None
    preferred_phrases: list[str] | str | None = None
    custom_instructions: str | None = None
    style_learning_enabled: bool = False
    style_learning_source_limit: int = 20


class OnboardingCompletePayload(BaseModel):
    email: str


class CancelSubscriptionPayload(BaseModel):
    reason: Optional[str] = None
    feedback: Optional[str] = None


def require_env(value: str | None, name: str) -> str:
    if not value:
        raise HTTPException(status_code=500, detail=f"Missing environment variable: {name}")
    return value


def utc_now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def decode_base64(data: str | None) -> str | None:
    if not data:
        return None

    padding = len(data) % 4
    if padding:
        data += "=" * (4 - padding)

    try:
        decoded_bytes = base64.urlsafe_b64decode(data)
        return decoded_bytes.decode("utf-8", errors="ignore")
    except Exception:
        return None


def get_header_value(headers: list[dict[str, Any]], name: str) -> str | None:
    for header in headers:
        if header.get("name", "").lower() == name.lower():
            return header.get("value")
    return None


def extract_email_address(from_header: str | None) -> str | None:
    if not from_header:
        return None

    if "<" in from_header and ">" in from_header:
        return from_header.split("<")[1].split(">")[0].strip()

    return from_header.strip()


# Sender addresses that explicitly say "do not reply" should never receive
# an AI-drafted response — you literally cannot reply to them. Detected by
# common patterns in the local-part of the email address.
_NOREPLY_LOCAL_PATTERNS = re.compile(
    r"(^|[._\-+])"
    r"("
    r"noreply|no\-reply|donotreply|do\-not\-reply|"
    r"notifications?|notify|"
    r"mailer\-daemon|postmaster|"
    r"automated|auto\-?confirm|auto\-?reply|"
    r"bounce|bounces|"
    r"system|notice|alerts?"
    r")"
    r"($|[._\-+@])",
    re.IGNORECASE,
)


def is_noreply_sender(sender_email: str | None) -> bool:
    """True for senders we should never auto-draft a reply to: noreply/donotreply
    style addresses, mailer-daemons, automated alerts. The check is on the
    local part (before the @) so 'ads-account-noreply@google.com' matches."""
    if not sender_email:
        return False
    addr = sender_email.strip().lower()
    if "@" not in addr:
        return False
    local = addr.split("@", 1)[0]
    return bool(_NOREPLY_LOCAL_PATTERNS.search(local))


# OfficeFlow's own briefing / weekly recap / silence radar mails are sent
# from the user's own mailbox to themselves. They always start with these
# subject prefixes. We detect them so they always get the same label (FYI),
# never get a draft reply, and don't burn classifier tokens.
_OFFICEFLOW_BRIEFING_SUBJECT_PREFIXES = (
    "officeflow briefing",
    "officeflow weekrecap",
    "officeflow weekly recap",
    "officeflow —",   # daily summary "OfficeFlow — 15 min bespaard..."
    "officeflow -",   # ASCII variant
)


def is_self_sent_officeflow_mail(
    sender_email: str | None,
    mailbox_email: str | None,
    subject: str | None,
) -> bool:
    """True for mails OfficeFlow sends to the user's own inbox (briefings,
    recaps, system summaries). Detected by from == mailbox_email AND subject
    matching a known OfficeFlow prefix. We always classify these as FYI to
    keep them out of Marketing / Notification / Priority noise."""
    if not sender_email or not mailbox_email:
        return False
    if sender_email.strip().lower() != mailbox_email.strip().lower():
        return False
    subj = (subject or "").strip().lower()
    return any(subj.startswith(p) for p in _OFFICEFLOW_BRIEFING_SUBJECT_PREFIXES)


def normalize_subject_for_reply(subject: str | None) -> str:
    if not subject:
        return "Re:"

    stripped = subject.strip()
    if stripped.lower().startswith("re:"):
        return stripped

    return f"Re: {stripped}"


def extract_plain_text_from_payload(payload: dict[str, Any]) -> str | None:
    if not payload:
        return None

    if payload.get("mimeType") == "text/plain":
        return decode_base64(payload.get("body", {}).get("data"))

    for part in payload.get("parts", []):
        if part.get("mimeType") == "text/plain":
            return decode_base64(part.get("body", {}).get("data"))

        for nested_part in part.get("parts", []):
            if nested_part.get("mimeType") == "text/plain":
                return decode_base64(nested_part.get("body", {}).get("data"))

    return decode_base64(payload.get("body", {}).get("data"))


def build_pricing_redirect(reason: str) -> str:
    return f"{FRONTEND_PRICING_URL}?reason={quote(reason)}"


def parse_response_data(response: httpx.Response) -> Any:
    if not response.text:
        return None

    try:
        return response.json()
    except Exception:
        return {"raw": response.text}


def parse_internal_date_ms(value: Any) -> int:
    try:
        return int(value)
    except Exception:
        return 0


def supabase_headers() -> dict[str, str]:
    service_role_key = require_env(SUPABASE_SERVICE_ROLE_KEY, "SUPABASE_SERVICE_ROLE_KEY")
    return {
        "apikey": service_role_key,
        "Authorization": f"Bearer {service_role_key}",
    }


def get_status_label_names() -> list[str]:
    return LABELS + list(LEGACY_LABEL_NAME_MAP.keys())


def get_status_label_ids_from_map(label_name_to_id: dict[str, str]) -> set[str]:
    label_ids: set[str] = set()
    for label_name in get_status_label_names():
        label_id = label_name_to_id.get(label_name)
        if label_id:
            label_ids.add(label_id)
    return label_ids


def is_draft_label_ids(label_ids: set[str]) -> bool:
    return "DRAFT" in label_ids and "SENT" not in label_ids


def normalize_string(value: Any) -> str | None:
    if value is None:
        return None

    if isinstance(value, str):
        cleaned = value.strip()
        return cleaned or None

    cleaned = str(value).strip()
    return cleaned or None


def normalize_phrase_list(value: Any) -> list[str]:
    if value is None:
        return []

    if isinstance(value, list):
        cleaned_items: list[str] = []
        for item in value:
            normalized = normalize_string(item)
            if normalized:
                cleaned_items.append(normalized)
        return cleaned_items

    if isinstance(value, str):
        raw = value.strip()
        if not raw:
            return []

        if "\n" in raw:
            parts = [part.strip() for part in raw.splitlines()]
        elif "," in raw:
            parts = [part.strip() for part in raw.split(",")]
        else:
            parts = [raw]

        return [part for part in parts if part]

    normalized = normalize_string(value)
    return [normalized] if normalized else []


def phrase_list_to_prompt_text(values: list[str]) -> str:
    if not values:
        return ""
    return "; ".join(values)


def sanitize_generated_reply(reply: str | None) -> str:
    if not reply:
        return ""

    text = reply.strip().replace("\r\n", "\n")

    text = re.sub(r"^```(?:text|txt|markdown)?\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s*```$", "", text)

    lines = [line.rstrip() for line in text.split("\n")]

    cleaned_lines: list[str] = []
    skip_patterns = [
        r"^\s*onderwerp\s*:",
        r"^\s*subject\s*:",
        r"^\s*re\s*:",
        r"^\s*\[je naam\]\s*$",
        r"^\s*\[your name\]\s*$",
        r"^\s*officeflow\s*$",
    ]

    for line in lines:
        stripped = line.strip()
        if any(re.match(pattern, stripped, flags=re.IGNORECASE) for pattern in skip_patterns):
            continue
        cleaned_lines.append(line)

    text = "\n".join(cleaned_lines).strip()

    placeholder_blocks = [
        "Groet,\n[je naam]\nOfficeFlow",
        "Groet,\n[je naam]",
        "Met vriendelijke groet,\n[je naam]\nOfficeFlow",
        "Met vriendelijke groet,\n[je naam]",
        "Best,\n[your name]",
        "Kind regards,\n[your name]",
    ]
    for block in placeholder_blocks:
        text = text.replace(block, "").strip()

    text = re.sub(r"\n{3,}", "\n\n", text).strip()

    return text


def strip_trailing_generic_signoff(reply_text: str) -> str:
    if not reply_text:
        return ""

    text = reply_text.strip()

    generic_signoffs = [
        r"(?:\n\s*)?met vriendelijke groet,?\s*$",
        r"(?:\n\s*)?vriendelijke groet,?\s*$",
        r"(?:\n\s*)?groet,?\s*$",
        r"(?:\n\s*)?groeten,?\s*$",
        r"(?:\n\s*)?kind regards,?\s*$",
        r"(?:\n\s*)?best regards,?\s*$",
        r"(?:\n\s*)?regards,?\s*$",
        r"(?:\n\s*)?best,?\s*$",
    ]

    changed = True
    while changed:
        changed = False
        for pattern in generic_signoffs:
            updated = re.sub(pattern, "", text, flags=re.IGNORECASE).strip()
            if updated != text:
                text = updated
                changed = True

    return text


def maybe_apply_signature(reply_text: str, settings: dict[str, Any] | None) -> str:
    if not reply_text:
        return ""

    cleaned_reply = strip_trailing_generic_signoff(reply_text.strip())
    signature_text = normalize_string((settings or {}).get("signature_text"))

    if signature_text:
        if signature_text in cleaned_reply:
            return cleaned_reply
        return f"{cleaned_reply}\n\n{signature_text}"

    return cleaned_reply


def split_language_tokens(value: str | None) -> list[str]:
    if not value:
        return []

    normalized = value.replace("/", ",").replace("|", ",").replace(";", ",")
    parts = [part.strip() for part in normalized.split(",") if part.strip()]
    return parts


def normalize_language_code(value: str | None) -> str | None:
    if not value:
        return None

    cleaned = value.strip().lower()
    return LANGUAGE_NAME_MAP.get(cleaned, cleaned)


def get_allowed_reply_languages(preferred_language: str | None) -> list[str]:
    tokens = split_language_tokens(preferred_language)
    normalized: list[str] = []

    for token in tokens:
        code = normalize_language_code(token)
        if code and code not in normalized:
            normalized.append(code)

    return normalized


def detect_language_from_text(text: str | None) -> str | None:
    if not text:
        return None

    lowered = f" {text.lower()} "

    dutch_markers = [
        " de ", " het ", " een ", " en ", " uw ", " je ", " jij ", " jullie ", " wij ",
        " graag ", " alvast ", " hierbij ", " vriendelijk ", " vriendelijke ", " groet ",
        " bedankt ", " dank ", " kunt ", " kunnen ", " vandaag ", " morgen ", " levering ",
        " offerte ", " prijs ", " aanvraag ", " status ", " momenteel ", " behandeling ",
        " ik ", " u ", " op de hoogte ",
    ]
    english_markers = [
        " the ", " and ", " your ", " you ", " we ", " please ", " thanks ", " thank you ",
        " regards ", " best ", " kindly ", " order ", " delivery ", " quote ", " price ",
        " status ", " currently ", " update ", " today ", " tomorrow ", " i ", " appreciate ",
        " looking forward ", " let me know ",
    ]
    german_markers = [
        " der ", " die ", " das ", " und ", " sie ", " wir ", " danke ", " bitte ", " angebot ",
        " lieferung ", " status ", " freundlichen ",
    ]
    french_markers = [
        " le ", " la ", " les ", " et ", " vous ", " nous ", " merci ", " cordialement ",
        " devis ", " livraison ",
    ]
    spanish_markers = [
        " el ", " la ", " los ", " las ", " y ", " usted ", " ustedes ", " gracias ",
        " presupuesto ", " entrega ", " pedido ", " estado ", " estimado ", " saludos ",
        " me gustaría ", " podrían ",
    ]

    scores = {
        "Dutch": sum(lowered.count(marker) for marker in dutch_markers),
        "English": sum(lowered.count(marker) for marker in english_markers),
        "German": sum(lowered.count(marker) for marker in german_markers),
        "French": sum(lowered.count(marker) for marker in french_markers),
        "Spanish": sum(lowered.count(marker) for marker in spanish_markers),
    }

    best_language = max(scores, key=scores.get)
    best_score = scores[best_language]

    if best_score <= 0:
        return None

    sorted_scores = sorted(scores.values(), reverse=True)
    if len(sorted_scores) > 1 and best_score == sorted_scores[1]:
        return None

    return best_language


def choose_reply_language(
    preferred_language: str | None,
    incoming_text: str | None,
) -> dict[str, Any]:
    allowed_languages = get_allowed_reply_languages(preferred_language)
    detected_incoming_language = detect_language_from_text(incoming_text)

    if not allowed_languages:
        return {
            "allowed_languages": [],
            "incoming_language": detected_incoming_language,
            "reply_language": detected_incoming_language,
            "fallback_language": None,
        }

    fallback_language = allowed_languages[0]

    if detected_incoming_language and detected_incoming_language in allowed_languages:
        reply_language = detected_incoming_language
    else:
        reply_language = fallback_language

    return {
        "allowed_languages": allowed_languages,
        "incoming_language": detected_incoming_language,
        "reply_language": reply_language,
        "fallback_language": fallback_language,
    }


def build_language_instruction_block(settings: dict[str, Any] | None, body_text: str | None) -> str:
    preferred_language = (settings or {}).get("preferred_language")
    language_choice = choose_reply_language(preferred_language=preferred_language, incoming_text=body_text)

    allowed_languages = language_choice["allowed_languages"]
    incoming_language = language_choice["incoming_language"]
    reply_language = language_choice["reply_language"]
    fallback_language = language_choice["fallback_language"]

    if not allowed_languages and not reply_language:
        return "Er zijn geen expliciete taalvoorkeuren ingesteld. Gebruik de meest natuurlijke taal voor deze reply."

    instructions: list[str] = []

    if allowed_languages:
        instructions.append("Toegestane antwoordtalen: " + ", ".join(allowed_languages) + ".")

    if incoming_language:
        instructions.append(f"Gedetecteerde taal van de inkomende e-mail: {incoming_language}.")
    else:
        instructions.append("De taal van de inkomende e-mail kon niet met zekerheid worden vastgesteld.")

    if reply_language:
        instructions.append(f"Schrijf deze reply in {reply_language}.")
    elif fallback_language:
        instructions.append(f"Schrijf deze reply in {fallback_language}.")

    if allowed_languages:
        instructions.append("Gebruik nooit een taal buiten de toegestane antwoordtalen.")
        instructions.append("Als de inkomende e-mail in een toegestane taal is geschreven, gebruik diezelfde taal.")
        instructions.append("Als de inkomende e-mail niet in een toegestane taal is geschreven of onduidelijk is, gebruik de primaire voorkeurstaal.")

    instructions.append("Mix geen talen binnen één reply.")

    return "\n".join(instructions)


def safe_parse_json(content: str) -> dict[str, Any]:
    if not content:
        raise ValueError("Empty content")

    cleaned = content.strip()
    cleaned = re.sub(r"^```json\s*", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"^```\s*", "", cleaned)
    cleaned = re.sub(r"\s*```$", "", cleaned)

    return json.loads(cleaned)


def normalize_confidence(raw: Any) -> str:
    """Coerce classifier confidence to one of CONFIDENCE_VALUES.

    Accepts str (case-insensitive), falls back to 'medium' when missing or
    unrecognized. We pick 'medium' (not 'low') as the default so a missing
    field doesn't flood inboxes with red Review labels.
    """
    if isinstance(raw, str):
        value = raw.strip().lower()
        if value in CONFIDENCE_VALUES:
            return value
    return "medium"


async def supabase_get(path_and_query: str, timeout: float = 30.0) -> Any:
    supabase_url = require_env(SUPABASE_URL, "SUPABASE_URL")

    async with httpx.AsyncClient(timeout=timeout) as client:
        response = await client.get(
            f"{supabase_url}{path_and_query}",
            headers=supabase_headers(),
        )

    data = parse_response_data(response)

    if response.status_code >= 400:
        raise HTTPException(status_code=500, detail=f"Supabase GET failed: {data}")

    return data


async def supabase_get_count(path_and_query: str, timeout: float = 30.0) -> int:
    """Return Content-Range total count for a PostgREST query.

    Uses `Prefer: count=exact` + `HEAD` so we don't download rows — just the
    total. The caller supplies the full `/rest/v1/<table>?filters...&select=id`
    string. Returns 0 on failure so callers can treat it as a soft count.
    """
    supabase_url = require_env(SUPABASE_URL, "SUPABASE_URL")

    headers = supabase_headers()
    headers["Prefer"] = "count=exact"
    headers["Range-Unit"] = "items"
    headers["Range"] = "0-0"

    try:
        async with httpx.AsyncClient(timeout=timeout) as client:
            response = await client.head(
                f"{supabase_url}{path_and_query}",
                headers=headers,
            )
        if response.status_code >= 400:
            return 0
        content_range = response.headers.get("content-range", "")
        # Format: "0-0/123" or "*/123"
        if "/" in content_range:
            total = content_range.split("/", 1)[1].strip()
            if total.isdigit():
                return int(total)
    except Exception as exc:
        print(f"[supabase-count] failed for {path_and_query}: {repr(exc)}")
    return 0


async def supabase_post(
    path_and_query: str,
    payload: Any,
    prefer: str = "return=representation",
    timeout: float = 30.0,
) -> Any:
    supabase_url = require_env(SUPABASE_URL, "SUPABASE_URL")

    headers = supabase_headers()
    headers["Content-Type"] = "application/json"
    headers["Prefer"] = prefer

    async with httpx.AsyncClient(timeout=timeout) as client:
        response = await client.post(
            f"{supabase_url}{path_and_query}",
            headers=headers,
            json=payload,
        )

    data = parse_response_data(response)

    if response.status_code >= 400:
        raise HTTPException(status_code=500, detail=f"Supabase POST failed: {data}")

    return data


async def supabase_patch(
    path_and_query: str,
    payload: Any,
    prefer: str = "return=representation",
    timeout: float = 30.0,
) -> Any:
    supabase_url = require_env(SUPABASE_URL, "SUPABASE_URL")

    headers = supabase_headers()
    headers["Content-Type"] = "application/json"
    headers["Prefer"] = prefer

    async with httpx.AsyncClient(timeout=timeout) as client:
        response = await client.patch(
            f"{supabase_url}{path_and_query}",
            headers=headers,
            json=payload,
        )

    data = parse_response_data(response)

    if response.status_code >= 400:
        raise HTTPException(status_code=500, detail=f"Supabase PATCH failed: {data}")

    return data


async def supabase_delete(path_and_query: str, timeout: float = 30.0) -> Any:
    supabase_url = require_env(SUPABASE_URL, "SUPABASE_URL")

    headers = supabase_headers()
    headers["Prefer"] = "return=representation"

    async with httpx.AsyncClient(timeout=timeout) as client:
        response = await client.delete(
            f"{supabase_url}{path_and_query}",
            headers=headers,
        )

    data = parse_response_data(response)

    if response.status_code >= 400:
        raise HTTPException(status_code=500, detail=f"Supabase DELETE failed: {data}")

    return data


# ----------------------------
# Users helpers
# ----------------------------

async def supabase_get_user_by_email(email: str) -> dict[str, Any] | None:
    data = await supabase_get(
        f"/rest/v1/users?email=eq.{quote(email, safe='')}&select=*"
    )
    return data[0] if isinstance(data, list) and data else None


async def supabase_get_user_by_stripe_customer_id(customer_id: str) -> dict[str, Any] | None:
    data = await supabase_get(
        f"/rest/v1/users?stripe_customer_id=eq.{quote(customer_id, safe='')}&select=*"
    )
    return data[0] if isinstance(data, list) and data else None


async def supabase_insert_user(email: str, full_name: str | None) -> dict[str, Any]:
    data = await supabase_post(
        "/rest/v1/users",
        [{"email": email, "full_name": full_name}],
    )

    if not isinstance(data, list) or not data:
        raise HTTPException(status_code=500, detail="Supabase users insert returned no rows")

    return data[0]


async def supabase_update_user_profile(user_id: str, full_name: str | None = None) -> dict[str, Any] | None:
    payload: dict[str, Any] = {}
    if full_name is not None:
        payload["full_name"] = full_name
    if not payload:
        return None

    data = await supabase_patch(
        f"/rest/v1/users?id=eq.{quote(user_id, safe='')}",
        payload,
    )
    return data[0] if isinstance(data, list) and data else data


async def supabase_update_user_subscription(
    user_id: str,
    subscription_status: str | None,
    access_allowed: bool,
    stripe_customer_id: str | None = None,
    stripe_subscription_id: str | None = None,
) -> dict[str, Any] | None:
    payload: dict[str, Any] = {
        "subscription_status": subscription_status,
        "access_allowed": access_allowed,
    }
    if stripe_customer_id is not None:
        payload["stripe_customer_id"] = stripe_customer_id
    if stripe_subscription_id is not None:
        payload["stripe_subscription_id"] = stripe_subscription_id

    data = await supabase_patch(
        f"/rest/v1/users?id=eq.{quote(user_id, safe='')}",
        payload,
    )
    return data[0] if isinstance(data, list) and data else data


async def ensure_user_has_access(email: str) -> dict[str, Any]:
    if not email:
        raise HTTPException(status_code=401, detail="Missing user email")

    user = await supabase_get_user_by_email(email)
    if not user:
        raise HTTPException(status_code=403, detail="No user record found")

    access_allowed = user.get("access_allowed")
    subscription_status = user.get("subscription_status")
    has_access_columns = ("access_allowed" in user) or ("subscription_status" in user)

    # Solo-pad: directe toegang via subscription_status / access_allowed
    solo_access = True
    if has_access_columns:
        if access_allowed is False:
            solo_access = False
        if subscription_status and subscription_status not in ALLOWED_SUBSCRIPTION_STATUSES:
            solo_access = False

    if solo_access:
        return user

    # Team-pad: user heeft geen eigen solo-abonnement, maar mogelijk wel
    # toegang via een actief team waarvan hij lid is.
    try:
        from app import teams as teams_module
        user_teams = await teams_module._resolve_user_teams(user["id"])
        for t in user_teams:
            team = t.get("team") or {}
            if team.get("access_allowed") and team.get("subscription_status") in ALLOWED_SUBSCRIPTION_STATUSES:
                return user
    except Exception as exc:
        print(f"[ensure_user_has_access] team-check failed for {email}: {repr(exc)}")

    raise HTTPException(status_code=403, detail="Active subscription required")


async def get_user_for_billing(email: str) -> dict[str, Any]:
    if not email:
        raise HTTPException(status_code=401, detail="Missing user email")

    user = await supabase_get_user_by_email(email)
    if not user:
        raise HTTPException(status_code=404, detail="No user record found")

    stripe_customer_id = user.get("stripe_customer_id")
    if not stripe_customer_id:
        raise HTTPException(status_code=400, detail="No Stripe customer found")

    return user


# ----------------------------
# Supabase Auth helpers (JWT + invites)
# ----------------------------

async def _verify_token_via_supabase(token: str) -> dict[str, Any]:
    """
    Fallback verification: ask Supabase Auth API to validate the token.
    Works with any JWT signing method (HS256 legacy OR ES256/RS256
    asymmetric signing keys). Returns the auth user payload.
    """
    supabase_url = require_env(SUPABASE_URL, "SUPABASE_URL")
    # Prefer anon key if present, otherwise fall back to service-role for the apikey header.
    apikey = os.getenv("SUPABASE_ANON_KEY") or SUPABASE_SERVICE_ROLE_KEY
    if not apikey:
        raise HTTPException(status_code=500, detail="SUPABASE_ANON_KEY or SUPABASE_SERVICE_ROLE_KEY missing")

    async with httpx.AsyncClient(timeout=10.0) as client:
        res = await client.get(
            f"{supabase_url}/auth/v1/user",
            headers={
                "Authorization": f"Bearer {token}",
                "apikey": apikey,
            },
        )

    if res.status_code == 401:
        raise HTTPException(status_code=401, detail="Token expired or invalid")
    if res.status_code != 200:
        raise HTTPException(status_code=401, detail=f"Auth verify failed ({res.status_code})")

    try:
        return res.json()
    except Exception:
        raise HTTPException(status_code=401, detail="Auth verify returned non-JSON")


async def get_current_user(
    authorization: Optional[str] = Header(None),
) -> dict[str, Any]:
    """
    FastAPI dependency — validates a Supabase Auth JWT sent as
    `Authorization: Bearer <token>` and returns the matching row from
    public.users. Raises 401 if the token is missing/invalid or 403 if
    there is no corresponding user row.

    Fast path: local HS256 decode with SUPABASE_JWT_SECRET (legacy).
    Fallback: remote verification via Supabase Auth API — needed when
    the project uses asymmetric signing keys (ES256/RS256).
    """
    if not authorization or not authorization.lower().startswith("bearer "):
        raise HTTPException(status_code=401, detail="Missing bearer token")

    token = authorization.split(" ", 1)[1].strip()
    email: str = ""

    # --- Fast path: local HS256 decode ---
    if SUPABASE_JWT_SECRET:
        try:
            decoded = jwt.decode(
                token,
                SUPABASE_JWT_SECRET,
                algorithms=["HS256"],
                audience="authenticated",
                options={"verify_aud": True},
            )
            email = (decoded.get("email") or "").strip().lower()
        except jwt.ExpiredSignatureError:
            raise HTTPException(status_code=401, detail="Token expired")
        except jwt.InvalidTokenError:
            # Likely an ES256/RS256 token from new Supabase signing keys —
            # fall through to Supabase Auth API verification below.
            email = ""

    # --- Fallback: verify via Supabase Auth API ---
    if not email:
        auth_user = await _verify_token_via_supabase(token)
        email = (auth_user.get("email") or "").strip().lower()

    if not email:
        raise HTTPException(status_code=401, detail="Token missing email claim")

    user = await supabase_get_user_by_email(email)
    if not user:
        raise HTTPException(status_code=403, detail="No user record found")

    return user


# Constant-time comparison to avoid timing attacks on the shared secret.
import hmac as _hmac
import hashlib as _hashlib
import secrets as _secrets


# ============ CSRF STATE FOR OAUTH FLOW ============
# Generates a HMAC-signed state token that ties /auth/google/start to its
# matching /auth/google/callback request. Prevents an attacker from forging
# an OAuth callback that links someone else's Google account to a user's
# OfficeFlow session.
#
# State format: <random_nonce>.<hmac_sig>
# We sign with INTERNAL_API_SECRET (or fall back to SUPABASE_JWT_SECRET) so
# tampering with the nonce invalidates the signature.

def _oauth_signing_key() -> bytes:
    """Pick the strongest available secret to sign OAuth state with."""
    key = INTERNAL_API_SECRET or SUPABASE_JWT_SECRET or ""
    if not key:
        raise HTTPException(
            status_code=503,
            detail="OAuth state signing not configured (INTERNAL_API_SECRET missing)",
        )
    return key.encode("utf-8")


def make_oauth_state() -> str:
    """Create a signed state token to attach to /auth/google/start."""
    nonce = _secrets.token_urlsafe(24)
    key = _oauth_signing_key()
    sig = _hmac.new(key, nonce.encode("utf-8"), _hashlib.sha256).hexdigest()[:24]
    return f"{nonce}.{sig}"


def verify_oauth_state(state: str | None) -> bool:
    """Returns True if the state token is well-formed and signed by us."""
    if not state or "." not in state:
        return False
    try:
        nonce, sig = state.rsplit(".", 1)
    except ValueError:
        return False
    if not nonce or not sig:
        return False
    try:
        key = _oauth_signing_key()
    except HTTPException:
        return False
    expected = _hmac.new(key, nonce.encode("utf-8"), _hashlib.sha256).hexdigest()[:24]
    return _hmac.compare_digest(sig, expected)


# ============ STRIPE WEBHOOK IDEMPOTENCY ============
# Stripe delivers webhooks at-least-once — a network blip can cause the same
# event to arrive twice. We dedupe in-memory using an LRU cache keyed by
# stripe event id. For a single-instance Railway deploy this is sufficient;
# for multi-instance you'd swap to a DB or Redis store (see Phase 5+).
from collections import OrderedDict as _OrderedDict

_STRIPE_PROCESSED_EVENTS: "_OrderedDict[str, float]" = _OrderedDict()
_STRIPE_PROCESSED_LIMIT = 5000  # entries — ~24h of events at our volume


def stripe_event_already_processed(event_id: str | None) -> bool:
    """Check + register a Stripe event id. Returns True if it was seen
    before (caller should skip processing). Thread-safe enough for our
    single-worker uvicorn setup."""
    if not event_id:
        return False
    if event_id in _STRIPE_PROCESSED_EVENTS:
        return True
    _STRIPE_PROCESSED_EVENTS[event_id] = _dt_now_ts()
    # Trim oldest entries if we're over the cap
    while len(_STRIPE_PROCESSED_EVENTS) > _STRIPE_PROCESSED_LIMIT:
        _STRIPE_PROCESSED_EVENTS.popitem(last=False)
    return False


def _dt_now_ts() -> float:
    return datetime.now(timezone.utc).timestamp()


async def supabase_record_webhook_event(
    stripe_event_id: str,
    event_type: str | None,
    livemode: bool | None,
) -> bool:
    """Atomically claim a Stripe event for processing. Returns True if this
    is the first time we see the event (caller should process it), False if
    it was already received (caller should skip — it's a duplicate).

    Uses the unique constraint on webhook_events.stripe_event_id for
    persistent dedup that survives backend restarts and works across
    multiple Railway instances. Falls back to in-memory cache if DB call
    fails (best-effort)."""
    if not stripe_event_id:
        return True  # no id = can't dedup, let it through

    payload = {
        "stripe_event_id": stripe_event_id,
        "event_type": event_type,
        "livemode": livemode,
        "processing_status": "pending",
    }
    try:
        result = await supabase_post(
            "/rest/v1/webhook_events",
            [payload],
            prefer="return=minimal",
        )
        # If insert succeeded, this is a new event
        return True
    except HTTPException as exc:
        detail = str(exc.detail) if exc.detail else ""
        if "23505" in detail or "duplicate key" in detail.lower():
            # Unique violation = already processed
            return False
        # Other DB error — log + fall through to in-memory cache
        print(f"[webhook-events-db] insert failed: {repr(exc)}")
        return True
    except Exception as exc:
        print(f"[webhook-events-db] unexpected error: {repr(exc)}")
        return True


async def supabase_mark_webhook_event_processed(
    stripe_event_id: str,
    success: bool,
    error_message: str | None = None,
) -> None:
    """Update the webhook_events row after processing. Best-effort — failure
    here doesn't break the webhook response (it just leaves the row in
    'pending' status and we'd manually clean it up if it ever matters)."""
    if not stripe_event_id:
        return
    payload: dict[str, Any] = {
        "processing_status": "processed" if success else "failed",
        "processed_at": datetime.now(timezone.utc).isoformat(),
    }
    if error_message:
        payload["error_message"] = error_message[:500]
    try:
        await supabase_patch(
            f"/rest/v1/webhook_events?stripe_event_id=eq.{quote(stripe_event_id, safe='')}",
            payload,
        )
    except Exception as exc:
        print(f"[webhook-events-db] update failed for {stripe_event_id}: {repr(exc)}")


# ============ TOKEN ENCRYPTION HELPERS (foundation, backward-compat) ============
# Encrypts OAuth tokens at rest using Fernet (symmetric AES-128-CBC + HMAC).
# The decrypt function is BACKWARD-COMPATIBLE: if input doesn't look encrypted
# (no Fernet prefix), it's returned as-is. This lets us roll out encryption
# gradually — new tokens are encrypted, old ones remain plain text until
# they're refreshed and re-stored.
#
# To enable: set TOKEN_ENCRYPTION_KEY env var (32 url-safe base64 bytes).
# Generate with: python3 -c "from cryptography.fernet import Fernet; print(Fernet.generate_key().decode())"

TOKEN_ENCRYPTION_KEY = os.getenv("TOKEN_ENCRYPTION_KEY")

# Lazy-initialised so missing key doesn't crash startup.
_FERNET = None


def _get_fernet():
    global _FERNET
    if _FERNET is None and TOKEN_ENCRYPTION_KEY:
        try:
            from cryptography.fernet import Fernet
            _FERNET = Fernet(TOKEN_ENCRYPTION_KEY.encode("utf-8"))
        except Exception as exc:
            print(f"[token-encryption] Fernet init failed: {repr(exc)}")
            _FERNET = None
    return _FERNET


def encrypt_token(plain: str | None) -> str | None:
    """Encrypt a token for storage. Returns plain string unchanged when no
    encryption key is configured (graceful no-op). Encrypted output starts
    with 'gAAAAA' (Fernet prefix) so we can detect on decrypt."""
    if not plain:
        return plain
    f = _get_fernet()
    if not f:
        return plain
    try:
        return f.encrypt(plain.encode("utf-8")).decode("utf-8")
    except Exception as exc:
        print(f"[token-encryption] encrypt failed: {repr(exc)}")
        return plain  # never lose data — fall back to plain


def decrypt_token(maybe_encrypted: str | None) -> str | None:
    """Decrypt a token from storage. If input doesn't look encrypted (no
    Fernet prefix), returned as-is. This handles legacy plain-text tokens
    during rollout."""
    if not maybe_encrypted:
        return maybe_encrypted
    # Fernet tokens always start with "gAAAAA" (base64 of version byte 0x80)
    if not maybe_encrypted.startswith("gAAAAA"):
        return maybe_encrypted
    f = _get_fernet()
    if not f:
        # We received an encrypted token but have no key — log and return None
        # so caller can detect failure (don't return a corrupt string).
        print("[token-encryption] decrypt called but no key configured")
        return None
    try:
        return f.decrypt(maybe_encrypted.encode("utf-8")).decode("utf-8")
    except Exception as exc:
        print(f"[token-encryption] decrypt failed: {repr(exc)}")
        return None


async def require_internal_secret(
    x_internal_secret: Optional[str] = Header(None, alias="X-Internal-Secret"),
) -> None:
    """FastAPI dependency for `/internal/*` endpoints. Requires the
    X-Internal-Secret header to match the INTERNAL_API_SECRET env var.

    Used to protect cron / scheduled / admin-only endpoints from being
    called by anonymous internet traffic — those endpoints can trigger
    expensive OpenAI / Gmail API operations on any email address, so
    they must NEVER be open.

    Set INTERNAL_API_SECRET on Railway and pass the same value as the
    X-Internal-Secret header from cron jobs / admin scripts.
    """
    if not INTERNAL_API_SECRET:
        # Fail closed: if the secret isn't configured, refuse all internal
        # calls. This prevents accidental open-by-default behaviour.
        raise HTTPException(
            status_code=503,
            detail="Internal API not configured (INTERNAL_API_SECRET missing)",
        )

    if not x_internal_secret or not _hmac.compare_digest(
        x_internal_secret, INTERNAL_API_SECRET
    ):
        raise HTTPException(status_code=401, detail="Invalid internal secret")


async def send_welcome_invite(email: str) -> bool:
    """
    Sends the Supabase Auth invite email so a user can set their first
    password. Safe to call multiple times — if the auth user already exists
    we log and return False instead of failing onboarding.
    """
    if not email:
        return False

    supabase_url = require_env(SUPABASE_URL, "SUPABASE_URL")
    service_role_key = require_env(SUPABASE_SERVICE_ROLE_KEY, "SUPABASE_SERVICE_ROLE_KEY")

    headers = {
        "apikey": service_role_key,
        "Authorization": f"Bearer {service_role_key}",
        "Content-Type": "application/json",
    }
    body = {
        "email": email,
        "data": {"source": "officeflow_onboarding"},
    }
    # Supabase v2 auth endpoint — "redirect_to" is sent as a query string.
    url = (
        f"{supabase_url}/auth/v1/invite"
        f"?redirect_to={quote(SUPABASE_SET_PASSWORD_URL, safe='')}"
    )

    try:
        async with httpx.AsyncClient(timeout=15.0) as client:
            response = await client.post(url, headers=headers, json=body)

        if response.status_code in (200, 201):
            print(f"[invite] Sent welcome invite to {email}")
            return True

        # 422 = user already exists → not fatal
        print(
            f"[invite] Supabase invite returned {response.status_code} "
            f"for {email}: {response.text[:200]}"
        )
        return False
    except Exception as exc:
        print(f"[invite] Failed to send invite to {email}: {repr(exc)}")
        return False


# ----------------------------
# Mailbox helpers
# ----------------------------

async def supabase_get_mailbox_by_user_and_email(
    user_id: str,
    email_address: str,
    provider: str = "gmail",
) -> dict[str, Any] | None:
    data = await supabase_get(
        (
            "/rest/v1/mailboxes"
            f"?user_id=eq.{quote(user_id, safe='')}"
            f"&provider=eq.{quote(provider, safe='')}"
            f"&email_address=eq.{quote(email_address, safe='')}"
            "&select=*"
        )
    )
    return data[0] if isinstance(data, list) and data else None


async def supabase_get_mailbox_by_user_id(
    user_id: str,
    provider: str = "gmail",
) -> dict[str, Any] | None:
    data = await supabase_get(
        (
            "/rest/v1/mailboxes"
            f"?user_id=eq.{quote(user_id, safe='')}"
            f"&provider=eq.{quote(provider, safe='')}"
            "&select=*"
        )
    )
    return data[0] if isinstance(data, list) and data else None


async def supabase_upsert_mailbox(
    user_id: str,
    provider: str,
    email_address: str,
    status: str = "connected",
    team_id: str | None = None,
) -> dict[str, Any]:
    payload: dict[str, Any] = {
        "user_id": user_id,
        "provider": provider,
        "email_address": email_address,
        "status": status,
    }
    if team_id:
        payload["team_id"] = team_id

    data = await supabase_post(
        "/rest/v1/mailboxes?on_conflict=user_id,provider,email_address",
        [payload],
        prefer="resolution=merge-duplicates,return=representation",
    )

    if not isinstance(data, list) or not data:
        raise HTTPException(status_code=500, detail="Supabase mailboxes upsert returned no rows")

    return data[0]


async def supabase_update_mailbox_status(mailbox_id: str, status: str) -> dict[str, Any] | None:
    data = await supabase_patch(
        f"/rest/v1/mailboxes?id=eq.{quote(mailbox_id, safe='')}",
        {"status": status},
    )
    return data[0] if isinstance(data, list) and data else data


# ============ SCHEDULED SENDS ============
async def supabase_insert_scheduled_send(
    user_id: str,
    mailbox_id: str,
    gmail_draft_id: str,
    gmail_thread_id: str | None,
    subject: str | None,
    to_email: str | None,
    scheduled_at_iso: str,
) -> dict[str, Any]:
    data = await supabase_post(
        "/rest/v1/scheduled_sends",
        [{
            "user_id": user_id,
            "mailbox_id": mailbox_id,
            "gmail_draft_id": gmail_draft_id,
            "gmail_thread_id": gmail_thread_id,
            "subject": subject,
            "to_email": to_email,
            "scheduled_at": scheduled_at_iso,
            "status": "pending",
        }],
        prefer="return=representation",
    )
    if not isinstance(data, list) or not data:
        raise HTTPException(status_code=500, detail="Failed to create scheduled send")
    return data[0]


async def supabase_get_user_scheduled_sends(
    user_id: str,
    statuses: list[str] | None = None,
    limit: int = 100,
) -> list[dict[str, Any]]:
    status_filter = ""
    if statuses:
        joined = ",".join(quote(s, safe="") for s in statuses)
        status_filter = f"&status=in.({joined})"
    data = await supabase_get(
        f"/rest/v1/scheduled_sends"
        f"?user_id=eq.{quote(user_id, safe='')}"
        f"{status_filter}"
        f"&order=scheduled_at.asc"
        f"&limit={limit}"
        f"&select=*"
    )
    return data if isinstance(data, list) else []


async def supabase_get_due_scheduled_sends(limit: int = 50) -> list[dict[str, Any]]:
    now_iso = datetime.now(timezone.utc).isoformat()
    data = await supabase_get(
        f"/rest/v1/scheduled_sends"
        f"?status=eq.pending"
        f"&scheduled_at=lte.{quote(now_iso, safe='')}"
        f"&order=scheduled_at.asc"
        f"&limit={limit}"
        f"&select=*"
    )
    return data if isinstance(data, list) else []


async def supabase_update_scheduled_send(
    scheduled_send_id: str,
    status: str,
    sent_at: str | None = None,
    error_message: str | None = None,
) -> dict[str, Any] | None:
    payload: dict[str, Any] = {"status": status}
    if sent_at is not None:
        payload["sent_at"] = sent_at
    if error_message is not None:
        payload["error_message"] = error_message
    data = await supabase_patch(
        f"/rest/v1/scheduled_sends?id=eq.{quote(scheduled_send_id, safe='')}",
        payload,
    )
    return data[0] if isinstance(data, list) and data else data


async def supabase_get_scheduled_send_by_id(
    scheduled_send_id: str,
    user_id: str,
) -> dict[str, Any] | None:
    """Look up a single row scoped to the requesting user (security)."""
    data = await supabase_get(
        f"/rest/v1/scheduled_sends"
        f"?id=eq.{quote(scheduled_send_id, safe='')}"
        f"&user_id=eq.{quote(user_id, safe='')}"
        f"&select=*"
    )
    if isinstance(data, list) and data:
        return data[0]
    return None


# ============================================================
# SNOOZED EMAILS — Supabase helpers
# ============================================================
async def supabase_insert_snoozed_email(
    *,
    user_id: str,
    mailbox_id: str,
    email_id: str,
    gmail_message_id: str,
    gmail_thread_id: str | None,
    subject: str | None,
    original_label: str | None,
    original_label_id: str | None,
    wake_at_iso: str,
) -> dict[str, Any]:
    data = await supabase_post(
        "/rest/v1/snoozed_emails",
        [{
            "user_id": user_id,
            "mailbox_id": mailbox_id,
            "email_id": email_id,
            "gmail_message_id": gmail_message_id,
            "gmail_thread_id": gmail_thread_id,
            "subject": subject,
            "original_label": original_label,
            "original_label_id": original_label_id,
            "wake_at": wake_at_iso,
            "status": "pending",
        }],
        prefer="return=representation",
    )
    if not isinstance(data, list) or not data:
        raise HTTPException(status_code=500, detail="Failed to create snooze")
    return data[0]


async def supabase_get_user_snoozed_emails(
    user_id: str,
    statuses: list[str] | None = None,
    limit: int = 100,
) -> list[dict[str, Any]]:
    status_filter = ""
    if statuses:
        joined = ",".join(quote(s, safe="") for s in statuses)
        status_filter = f"&status=in.({joined})"
    data = await supabase_get(
        f"/rest/v1/snoozed_emails"
        f"?user_id=eq.{quote(user_id, safe='')}"
        f"{status_filter}"
        f"&order=wake_at.asc"
        f"&limit={limit}"
        f"&select=*"
    )
    return data if isinstance(data, list) else []


async def supabase_get_due_snoozed_emails(limit: int = 50) -> list[dict[str, Any]]:
    now_iso = datetime.now(timezone.utc).isoformat()
    data = await supabase_get(
        f"/rest/v1/snoozed_emails"
        f"?status=eq.pending"
        f"&wake_at=lte.{quote(now_iso, safe='')}"
        f"&order=wake_at.asc"
        f"&limit={limit}"
        f"&select=*"
    )
    return data if isinstance(data, list) else []


async def supabase_update_snoozed_email(
    snooze_id: str,
    *,
    status: str,
    woken_at: str | None = None,
    error_message: str | None = None,
) -> dict[str, Any] | None:
    payload: dict[str, Any] = {"status": status}
    if woken_at is not None:
        payload["woken_at"] = woken_at
    if error_message is not None:
        payload["error_message"] = error_message
    data = await supabase_patch(
        f"/rest/v1/snoozed_emails?id=eq.{quote(snooze_id, safe='')}",
        payload,
    )
    return data[0] if isinstance(data, list) and data else data


async def supabase_get_snoozed_email_by_id(
    snooze_id: str,
    user_id: str,
) -> dict[str, Any] | None:
    """Look up a single snooze row scoped to the requesting user (security)."""
    data = await supabase_get(
        f"/rest/v1/snoozed_emails"
        f"?id=eq.{quote(snooze_id, safe='')}"
        f"&user_id=eq.{quote(user_id, safe='')}"
        f"&select=*"
    )
    if isinstance(data, list) and data:
        return data[0]
    return None


async def supabase_get_pending_snooze_for_email(
    *,
    email_id: str,
    user_id: str,
) -> dict[str, Any] | None:
    """Return the active (pending) snooze row for a given email, if any.

    Used by the snooze-create endpoint as a friendly pre-check (the partial
    unique index ultimately enforces this) and by the cancel-by-email-id path.
    """
    data = await supabase_get(
        f"/rest/v1/snoozed_emails"
        f"?email_id=eq.{quote(email_id, safe='')}"
        f"&user_id=eq.{quote(user_id, safe='')}"
        f"&status=eq.pending"
        f"&select=*"
        f"&limit=1"
    )
    if isinstance(data, list) and data:
        return data[0]
    return None


async def supabase_get_email_by_id(
    email_id: str,
    user_id: str,
) -> dict[str, Any] | None:
    """Look up an emails-table row scoped to the requesting user."""
    data = await supabase_get(
        f"/rest/v1/emails"
        f"?id=eq.{quote(email_id, safe='')}"
        f"&user_id=eq.{quote(user_id, safe='')}"
        f"&select=*"
    )
    if isinstance(data, list) and data:
        return data[0]
    return None


def _mailbox_cutoff_epoch(mailbox: dict[str, Any] | None) -> int | None:
    """Return the mailbox's inbox_cutoff_at as Unix epoch seconds.

    Used to add `after:<epoch>` to the Gmail search query so the classifier
    never sees mail that existed before the mailbox was connected. Returns
    None when the field is missing or unparseable — in that case the caller
    falls back to the unrestricted `in:inbox is:unread` query.
    """
    if not mailbox:
        return None
    raw = mailbox.get("inbox_cutoff_at")
    if not raw:
        return None
    try:
        # Postgres timestamptz serializes as ISO8601 with a trailing 'Z' or
        # a numeric offset. Python's fromisoformat understands both only in
        # 3.11+, and earlier versions choke on 'Z'. Normalise first.
        iso = raw.replace("Z", "+00:00") if isinstance(raw, str) else raw
        dt = datetime.fromisoformat(iso) if isinstance(iso, str) else iso
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        return int(dt.timestamp())
    except Exception as exc:
        print(f"[inbox-cutoff] parse failed for {raw!r}: {repr(exc)}")
        return None


async def ensure_mailbox_inbox_cutoff(mailbox: dict[str, Any]) -> None:
    """Stamp inbox_cutoff_at = now() the first time this mailbox is used.

    Ensures the classifier only ever looks at mail that arrived AFTER the
    mailbox was connected to OfficeFlow. We never overwrite an existing
    cutoff — it's a permanent floor. Mutates the passed mailbox dict so the
    caller's in-memory copy reflects the freshly stamped timestamp.
    """
    if not mailbox or mailbox.get("inbox_cutoff_at"):
        return
    mailbox_id = mailbox.get("id")
    if not mailbox_id:
        return
    stamp = utc_now_iso()
    try:
        await supabase_patch(
            f"/rest/v1/mailboxes?id=eq.{quote(mailbox_id, safe='')}"
            f"&inbox_cutoff_at=is.null",
            {"inbox_cutoff_at": stamp},
        )
        mailbox["inbox_cutoff_at"] = stamp
    except Exception as exc:
        print(f"[inbox-cutoff] mailbox {mailbox_id} stamp failed: {repr(exc)}")


async def get_all_active_mailboxes() -> list[dict[str, Any]]:
    data = await supabase_get(
        "/rest/v1/mailboxes?status=eq.connected&provider=eq.gmail&select=*"
    )
    return data if isinstance(data, list) else []


# ----------------------------
# OAuth account helpers
# ----------------------------

async def supabase_get_oauth_account(
    user_id: str,
    provider: str = "google",
) -> dict[str, Any] | None:
    data = await supabase_get(
        (
            "/rest/v1/oauth_accounts"
            f"?user_id=eq.{quote(user_id, safe='')}"
            f"&provider=eq.{quote(provider, safe='')}"
            "&select=*"
        )
    )
    return data[0] if isinstance(data, list) and data else None


async def supabase_update_oauth_account_tokens(
    user_id: str,
    provider: str,
    access_token: str | None = None,
    refresh_token: str | None = None,
) -> dict[str, Any] | None:
    payload: dict[str, Any] = {}
    if access_token is not None:
        payload["access_token"] = access_token
    if refresh_token is not None:
        payload["refresh_token"] = refresh_token
    if not payload:
        return None

    data = await supabase_patch(
        (
            "/rest/v1/oauth_accounts"
            f"?user_id=eq.{quote(user_id, safe='')}"
            f"&provider=eq.{quote(provider, safe='')}"
        ),
        payload,
    )
    return data[0] if isinstance(data, list) and data else data


async def supabase_upsert_oauth_account(
    user_id: str,
    provider: str,
    provider_account_id: str | None,
    access_token: str | None,
    refresh_token: str | None,
) -> dict[str, Any]:
    existing = await supabase_get_oauth_account(user_id=user_id, provider=provider)
    existing_refresh_token = existing.get("refresh_token") if existing else None
    effective_refresh_token = refresh_token if refresh_token else existing_refresh_token

    data = await supabase_post(
        "/rest/v1/oauth_accounts?on_conflict=user_id,provider",
        [{
            "user_id": user_id,
            "provider": provider,
            "provider_account_id": provider_account_id,
            "access_token": access_token,
            "refresh_token": effective_refresh_token,
        }],
        prefer="resolution=merge-duplicates,return=representation",
    )

    if not isinstance(data, list) or not data:
        raise HTTPException(status_code=500, detail="Supabase oauth_accounts upsert returned no rows")

    return data[0]


# ----------------------------
# Onboarding / email / drafts / settings / style profiles
# ----------------------------

async def supabase_upsert_onboarding_state(
    user_id: str,
    gmail_connected: bool,
    profile_completed: bool,
    initial_sync_completed: bool,
    first_draft_generated: bool,
) -> dict[str, Any] | None:
    data = await supabase_post(
        "/rest/v1/onboarding_state?on_conflict=user_id",
        [{
            "user_id": user_id,
            "gmail_connected": gmail_connected,
            "profile_completed": profile_completed,
            "initial_sync_completed": initial_sync_completed,
            "first_draft_generated": first_draft_generated,
        }],
        prefer="resolution=merge-duplicates,return=representation",
    )
    return data[0] if isinstance(data, list) and data else data


async def supabase_get_onboarding_state(user_id: str) -> dict[str, Any] | None:
    data = await supabase_get(
        f"/rest/v1/onboarding_state?user_id=eq.{quote(user_id, safe='')}&select=*"
    )
    return data[0] if isinstance(data, list) and data else None


async def supabase_mark_first_run_seen(user_id: str) -> None:
    """Stamp first_run_seen_at = now() once — never overwritten."""
    try:
        await supabase_patch(
            f"/rest/v1/onboarding_state"
            f"?user_id=eq.{quote(user_id, safe='')}"
            f"&first_run_seen_at=is.null",
            {"first_run_seen_at": utc_now_iso()},
        )
    except Exception as exc:
        print(f"[first-run] mark seen failed for {user_id}: {repr(exc)}")


async def supabase_insert_email(
    user_id: str,
    mailbox_id: str,
    gmail_message_id: str,
    gmail_thread_id: str | None,
    subject: str | None,
) -> dict[str, Any] | None:
    data = await supabase_post(
        "/rest/v1/emails?on_conflict=gmail_message_id",
        [{
            "user_id": user_id,
            "mailbox_id": mailbox_id,
            "gmail_message_id": gmail_message_id,
            "gmail_thread_id": gmail_thread_id,
            "subject": subject,
        }],
        prefer="resolution=merge-duplicates,return=representation",
    )
    return data[0] if isinstance(data, list) and data else data


async def supabase_get_drafts_by_email_id(email_id: str) -> list[dict[str, Any]]:
    data = await supabase_get(
        f"/rest/v1/drafts?email_id=eq.{quote(email_id, safe='')}&select=*"
    )
    return data if isinstance(data, list) else []


async def supabase_insert_draft(
    user_id: str,
    email_id: str,
    gmail_draft_id: str | None,
    subject: str | None,
    draft_body: str,
    status: str = "generated",
    confidence: str | None = None,
) -> dict[str, Any] | None:
    # Only persist confidence if it matches the allowed set; anything else
    # becomes null so the DB check constraint is never violated.
    safe_confidence: str | None = None
    if isinstance(confidence, str) and confidence in ("high", "medium", "low"):
        safe_confidence = confidence

    try:
        data = await supabase_post(
            "/rest/v1/drafts",
            [{
                "user_id": user_id,
                "email_id": email_id,
                "gmail_draft_id": gmail_draft_id,
                "subject": subject,
                "draft_body": draft_body,
                "status": status,
                "confidence": safe_confidence,
            }],
        )
        return data[0] if isinstance(data, list) and data else data
    except HTTPException as exc:
        # Catch the unique-index violation from drafts_one_active_per_email_idx.
        # This means a concurrent process_inbox_for_user run already created
        # a draft for this email — log it and return None so the caller
        # treats it as "draft already exists, skip".
        detail = str(exc.detail) if exc.detail else ""
        if "23505" in detail or "duplicate key" in detail.lower():
            print(
                f"[draft-race] Concurrent draft already exists for email_id={email_id} "
                f"(user {user_id}) — skipping duplicate insert"
            )
            return None
        raise


async def supabase_upsert_gmail_label(
    user_id: str,
    mailbox_id: str,
    label_name: str,
    label_id: str,
) -> dict[str, Any] | None:
    data = await supabase_post(
        "/rest/v1/gmail_labels?on_conflict=user_id,mailbox_id,label_name",
        [{
            "user_id": user_id,
            "mailbox_id": mailbox_id,
            "label_name": label_name,
            "label_id": label_id,
        }],
        prefer="resolution=merge-duplicates,return=representation",
    )
    return data[0] if isinstance(data, list) and data else data


async def supabase_get_user_settings(user_id: str) -> dict[str, Any] | None:
    data = await supabase_get(
        f"/rest/v1/user_settings?user_id=eq.{quote(user_id, safe='')}&select=*"
    )
    return data[0] if isinstance(data, list) and data else None


async def supabase_upsert_user_settings(
    user_id: str,
    payload: dict[str, Any],
) -> dict[str, Any]:
    data = await supabase_post(
        "/rest/v1/user_settings?on_conflict=user_id",
        [{
            "user_id": user_id,
            **payload,
        }],
        prefer="resolution=merge-duplicates,return=representation",
    )

    if not isinstance(data, list) or not data:
        raise HTTPException(status_code=500, detail="Supabase user_settings upsert returned no rows")

    return data[0]


async def supabase_get_user_style_profile(user_id: str) -> dict[str, Any] | None:
    data = await supabase_get(
        f"/rest/v1/user_style_profiles?user_id=eq.{quote(user_id, safe='')}&select=*"
    )
    return data[0] if isinstance(data, list) and data else None


async def supabase_upsert_user_style_profile(
    user_id: str,
    source_sent_count: int,
    style_profile_text: str,
    style_profile_json: dict[str, Any] | None = None,
) -> dict[str, Any]:
    data = await supabase_post(
        "/rest/v1/user_style_profiles?on_conflict=user_id",
        [{
            "user_id": user_id,
            "source_sent_count": source_sent_count,
            "last_trained_at": utc_now_iso(),
            "style_profile_text": style_profile_text,
            "style_profile_json": style_profile_json or {},
        }],
        prefer="resolution=merge-duplicates,return=representation",
    )

    if not isinstance(data, list) or not data:
        raise HTTPException(status_code=500, detail="Supabase user_style_profiles upsert returned no rows")

    return data[0]


def build_reply_style_instructions(settings: dict[str, Any] | None) -> str:
    if not settings:
        return ""

    instructions: list[str] = []

    tone_preference = settings.get("tone_preference")
    formality = settings.get("formality")
    length_preference = settings.get("length_preference")
    emoji_preference = settings.get("emoji_preference")
    cta_preference = settings.get("cta_preference")
    signature_mode = settings.get("signature_mode")
    signature_text = normalize_string(settings.get("signature_text"))
    forbidden_phrases = normalize_phrase_list(settings.get("forbidden_phrases"))
    preferred_phrases = normalize_phrase_list(settings.get("preferred_phrases"))
    custom_instructions = normalize_string(settings.get("custom_instructions"))

    if tone_preference:
        instructions.append(f"Gewenste toon: {tone_preference}.")
    if formality:
        instructions.append(f"Gewenste formaliteit: {formality}.")

    if length_preference == "kort":
        instructions.append(
            "Gewenste lengte: kort. Houd het bij 1 tot 2 korte zinnen als dat voldoende is."
        )
    elif length_preference == "gemiddeld":
        instructions.append(
            "Gewenste lengte: gemiddeld. Geef een volledig maar compact antwoord, meestal 2 tot 4 zinnen."
        )
    elif length_preference == "uitgebreid":
        instructions.append(
            "Gewenste lengte: uitgebreid. Je mag uitgebreider antwoorden als dat nuttig is, zolang het relevant en concreet blijft."
        )
    elif length_preference:
        instructions.append(f"Gewenste lengte: {length_preference}.")

    if emoji_preference is False:
        instructions.append("Gebruik geen emoji.")
    elif emoji_preference is True:
        instructions.append("Gebruik alleen spaarzaam emoji als dat natuurlijk voelt.")

    if cta_preference:
        instructions.append(f"Call-to-action voorkeur: {cta_preference}.")

    if signature_text:
        instructions.append(
            "Er is een vaste handtekening opgeslagen. Voeg zelf geen alternatieve handtekening, losse afsluitgroet of placeholdernaam toe; de vaste handtekening wordt later automatisch toegevoegd."
        )
    elif signature_mode in {"none", None, ""}:
        instructions.append("Voeg geen handtekening toe.")
    elif signature_mode == "include_name":
        instructions.append("Houd de afsluiting minimaal. Voeg geen placeholdernaam toe.")
    elif signature_mode == "full_signature":
        instructions.append("Gebruik alleen een handtekening als die expliciet bekend is. Gebruik nooit placeholders.")

    instructions.append(
        "Verwijs nooit naar websites, pagina's of externe informatie tenzij dat expliciet gevraagd wordt in de e-mail of expliciet is ingesteld door de gebruiker."
    )
    instructions.append("Voeg nooit extra informatie toe die niet gevraagd is.")

    if forbidden_phrases:
        instructions.append(
            f"Gebruik deze formuleringen niet: {phrase_list_to_prompt_text(forbidden_phrases)}."
        )
    if preferred_phrases:
        instructions.append(
            f"Gebruik bij voorkeur deze stijl of formuleringen: {phrase_list_to_prompt_text(preferred_phrases)}."
        )
    if custom_instructions:
        instructions.append(f"Extra instructies van de gebruiker: {custom_instructions}")

    return "\n".join(instructions)


def build_style_profile_instructions(style_profile: dict[str, Any] | None) -> str:
    if not style_profile:
        return ""

    style_profile_text = style_profile.get("style_profile_text")
    if not style_profile_text:
        return ""

    return f"Geleerd stijlprofiel van eerdere echte verzonden mails:\n{style_profile_text}"


def clean_reply_training_text(text: str | None) -> str:
    if not text:
        return ""

    cleaned = text.strip()

    split_markers = [
        "\nOp ",
        "\nOn ",
        "\nFrom:",
        "\nVan:",
        "\n-----Original Message-----",
        "\n________________________________",
    ]

    for marker in split_markers:
        if marker in cleaned:
            cleaned = cleaned.split(marker)[0].strip()

    lines = []
    for line in cleaned.splitlines():
        stripped = line.strip()
        if stripped.startswith(">"):
            continue
        lines.append(line)

    cleaned = "\n".join(lines).strip()
    return cleaned


def build_clean_settings_payload(payload: PromptSettingsPayload) -> dict[str, Any]:
    preferred_language = normalize_string(payload.preferred_language)
    tone_preference = normalize_string(payload.tone_preference)
    formality = normalize_string(payload.formality)
    length_preference = normalize_string(payload.length_preference)
    cta_preference = normalize_string(payload.cta_preference)
    signature_mode = normalize_string(payload.signature_mode)
    signature_text = normalize_string(payload.signature_text)
    custom_instructions = normalize_string(payload.custom_instructions)

    forbidden_phrases = normalize_phrase_list(payload.forbidden_phrases)
    preferred_phrases = normalize_phrase_list(payload.preferred_phrases)

    source_limit = payload.style_learning_source_limit
    if source_limit < 1:
        source_limit = 1
    if source_limit > 20:
        source_limit = 20

    return {
        "preferred_language": preferred_language,
        "tone_preference": tone_preference,
        "formality": formality,
        "length_preference": length_preference,
        "emoji_preference": payload.emoji_preference,
        "cta_preference": cta_preference,
        "signature_mode": signature_mode,
        "signature_text": signature_text,
        "forbidden_phrases": forbidden_phrases,
        "preferred_phrases": preferred_phrases,
        "custom_instructions": custom_instructions,
        "style_learning_enabled": payload.style_learning_enabled,
        "style_learning_source_limit": source_limit,
    }


# ----------------------------
# Gmail auth / context helpers
# ----------------------------

async def get_gmail_context_by_email(email: str) -> dict[str, Any]:
    user = await ensure_user_has_access(email)
    user_id = user["id"]

    mailbox = await supabase_get_mailbox_by_user_and_email(
        user_id=user_id,
        email_address=email,
        provider="gmail",
    )
    if not mailbox:
        mailbox = await supabase_get_mailbox_by_user_id(
            user_id=user_id,
            provider="gmail",
        )
    if not mailbox:
        raise HTTPException(status_code=404, detail="Gmail mailbox not found")

    oauth = await supabase_get_oauth_account(
        user_id=user_id,
        provider="google",
    )
    if not oauth:
        raise HTTPException(status_code=400, detail="Google account not connected")

    return {
        "user": user,
        "mailbox": mailbox,
        "oauth": oauth,
    }


async def refresh_google_access_token(user_id: str, refresh_token: str) -> str:
    client_id = require_env(GOOGLE_CLIENT_ID, "GOOGLE_CLIENT_ID")
    client_secret = require_env(GOOGLE_CLIENT_SECRET, "GOOGLE_CLIENT_SECRET")

    async with httpx.AsyncClient(timeout=60.0) as client:
        response = await client.post(
            "https://oauth2.googleapis.com/token",
            data={
                "client_id": client_id,
                "client_secret": client_secret,
                "refresh_token": refresh_token,
                "grant_type": "refresh_token",
            },
        )

    data = parse_response_data(response)
    new_access_token = data.get("access_token") if isinstance(data, dict) else None

    if not new_access_token:
        # invalid_grant = refresh token is permanently revoked/expired (user
        # disconnected the app, scope changed, 6mo inactivity, etc). Mark the
        # mailbox as needs_reconnect so the polling loop skips it and the
        # dashboard can surface a "reconnect Gmail" CTA. Re-OAuth will flip
        # the status back to 'connected' automatically.
        error_code = data.get("error") if isinstance(data, dict) else None
        if error_code == "invalid_grant":
            try:
                mailbox = await supabase_get_mailbox_by_user_id(
                    user_id=user_id,
                    provider="gmail",
                )
                if mailbox and mailbox.get("status") != "needs_reconnect":
                    await supabase_update_mailbox_status(
                        mailbox_id=mailbox["id"],
                        status="needs_reconnect",
                    )
                    print(
                        f"[oauth-revoked] Marked mailbox {mailbox.get('email_address')} "
                        f"(user {user_id}) as needs_reconnect — refresh token invalid"
                    )
            except Exception as exc:
                print(f"[oauth-revoked] Failed to mark mailbox for user {user_id}: {repr(exc)}")
        raise HTTPException(status_code=401, detail=f"Google token refresh failed: {data}")

    await supabase_update_oauth_account_tokens(
        user_id=user_id,
        provider="google",
        access_token=new_access_token,
    )
    return new_access_token


async def gmail_get_json_for_user(user_id: str, url: str, params: dict[str, Any] | None = None) -> Any:
    oauth = await supabase_get_oauth_account(user_id=user_id, provider="google")
    if not oauth:
        raise HTTPException(status_code=400, detail="Google account not connected")

    access_token = oauth.get("access_token")
    refresh_token = oauth.get("refresh_token")
    if not access_token:
        raise HTTPException(status_code=400, detail="Missing Google access token")

    async with httpx.AsyncClient(timeout=60.0) as client:
        response = await client.get(
            url,
            headers={"Authorization": f"Bearer {access_token}"},
            params=params,
        )

        if response.status_code == 401 and refresh_token:
            access_token = await refresh_google_access_token(user_id=user_id, refresh_token=refresh_token)
            response = await client.get(
                url,
                headers={"Authorization": f"Bearer {access_token}"},
                params=params,
            )

    data = parse_response_data(response)
    if response.status_code >= 400:
        raise HTTPException(status_code=response.status_code, detail=f"Gmail API error: {data}")
    return data


async def gmail_post_json_for_user(user_id: str, url: str, payload: dict[str, Any]) -> Any:
    oauth = await supabase_get_oauth_account(user_id=user_id, provider="google")
    if not oauth:
        raise HTTPException(status_code=400, detail="Google account not connected")

    access_token = oauth.get("access_token")
    refresh_token = oauth.get("refresh_token")
    if not access_token:
        raise HTTPException(status_code=400, detail="Missing Google access token")

    async with httpx.AsyncClient(timeout=60.0) as client:
        response = await client.post(
            url,
            headers={
                "Authorization": f"Bearer {access_token}",
                "Content-Type": "application/json",
            },
            json=payload,
        )

        if response.status_code == 401 and refresh_token:
            access_token = await refresh_google_access_token(user_id=user_id, refresh_token=refresh_token)
            response = await client.post(
                url,
                headers={
                    "Authorization": f"Bearer {access_token}",
                    "Content-Type": "application/json",
                },
                json=payload,
            )

    data = parse_response_data(response)
    if response.status_code >= 400:
        raise HTTPException(status_code=response.status_code, detail=f"Gmail API error: {data}")
    return data


async def gmail_patch_json_for_user(user_id: str, url: str, payload: dict[str, Any]) -> Any:
    oauth = await supabase_get_oauth_account(user_id=user_id, provider="google")
    if not oauth:
        raise HTTPException(status_code=400, detail="Google account not connected")

    access_token = oauth.get("access_token")
    refresh_token = oauth.get("refresh_token")
    if not access_token:
        raise HTTPException(status_code=400, detail="Missing Google access token")

    async with httpx.AsyncClient(timeout=60.0) as client:
        response = await client.patch(
            url,
            headers={
                "Authorization": f"Bearer {access_token}",
                "Content-Type": "application/json",
            },
            json=payload,
        )

        if response.status_code == 401 and refresh_token:
            access_token = await refresh_google_access_token(user_id=user_id, refresh_token=refresh_token)
            response = await client.patch(
                url,
                headers={
                    "Authorization": f"Bearer {access_token}",
                    "Content-Type": "application/json",
                },
                json=payload,
            )

    data = parse_response_data(response)
    if response.status_code >= 400:
        raise HTTPException(status_code=response.status_code, detail=f"Gmail API error: {data}")
    return data


async def gmail_delete_for_user(user_id: str, url: str) -> Any:
    oauth = await supabase_get_oauth_account(user_id=user_id, provider="google")
    if not oauth:
        raise HTTPException(status_code=400, detail="Google account not connected")

    access_token = oauth.get("access_token")
    refresh_token = oauth.get("refresh_token")
    if not access_token:
        raise HTTPException(status_code=400, detail="Missing Google access token")

    async with httpx.AsyncClient(timeout=60.0) as client:
        response = await client.delete(
            url,
            headers={"Authorization": f"Bearer {access_token}"},
        )

        if response.status_code == 401 and refresh_token:
            access_token = await refresh_google_access_token(user_id=user_id, refresh_token=refresh_token)
            response = await client.delete(
                url,
                headers={"Authorization": f"Bearer {access_token}"},
            )

    if response.status_code == 404:
        return {"status": "not_found"}

    data = parse_response_data(response)
    if response.status_code >= 400:
        raise HTTPException(status_code=response.status_code, detail=f"Gmail API error: {data}")
    return data or {"status": "deleted"}


async def get_all_gmail_labels(user_id: str) -> dict[str, dict[str, Any]]:
    labels_response = await gmail_get_json_for_user(
        user_id=user_id,
        url=f"{GMAIL_API_BASE}/labels",
    )
    return {label["name"]: label for label in labels_response.get("labels", [])}


async def modify_gmail_message_labels(
    user_id: str,
    gmail_message_id: str,
    add_label_ids: list[str] | None = None,
    remove_label_ids: list[str] | None = None,
) -> Any:
    payload: dict[str, Any] = {}
    if add_label_ids:
        payload["addLabelIds"] = add_label_ids
    if remove_label_ids:
        payload["removeLabelIds"] = remove_label_ids
    if not payload:
        return None

    return await gmail_post_json_for_user(
        user_id=user_id,
        url=f"{GMAIL_API_BASE}/messages/{gmail_message_id}/modify",
        payload=payload,
    )


async def archive_low_value_thread(
    user_id: str,
    thread_id: str | None,
    current_message_id: str,
) -> None:
    """Remove the Gmail INBOX label from every message in a thread.

    The thread remains in All Mail and stays findable via its status label
    (Marketing / Notification / Ignore). Nothing is trashed, no
    other labels are touched. Safe to call multiple times.
    """
    if not thread_id:
        try:
            await modify_gmail_message_labels(
                user_id=user_id,
                gmail_message_id=current_message_id,
                remove_label_ids=["INBOX"],
            )
        except Exception as exc:
            print(f"[auto-archive] msg {current_message_id}: {repr(exc)}")
        return

    try:
        thread_data = await gmail_get_json_for_user(
            user_id=user_id,
            url=f"{GMAIL_API_BASE}/threads/{thread_id}",
        )
    except Exception as exc:
        print(f"[auto-archive] thread fetch {thread_id}: {repr(exc)}")
        return

    for thread_message in thread_data.get("messages", []) or []:
        thread_message_id = thread_message.get("id")
        if not thread_message_id:
            continue
        if "INBOX" not in set(thread_message.get("labelIds", []) or []):
            continue
        try:
            await modify_gmail_message_labels(
                user_id=user_id,
                gmail_message_id=thread_message_id,
                remove_label_ids=["INBOX"],
            )
        except Exception as exc:
            print(f"[auto-archive] msg {thread_message_id}: {repr(exc)}")


async def is_trusted_sender(
    user_id: str,
    sender_email: str | None,
    cache: dict[str, bool] | None = None,
) -> bool:
    """True if the user has ever communicated with this sender or previously
    flagged their mail as important.

    Two safety-net checks (OR-combined):
      1. User has sent at least one mail TO this sender (reply history).
      2. At least one mail FROM this sender was previously labeled
         Priority / To Respond / Follow Up.

    Used as a guard before auto-archiving. Bekende contacten worden nooit
    uit de Inbox gehaald, ook niet als een mail als Marketing gelabeld wordt.
    """
    if not sender_email:
        return False
    normalized = sender_email.lower().strip()
    if not normalized:
        return False
    if cache is not None and normalized in cache:
        return cache[normalized]

    trusted = False

    # Check 1: reply history — have we ever sent to this address?
    try:
        data = await gmail_get_json_for_user(
            user_id=user_id,
            url=f"{GMAIL_API_BASE}/messages",
            params={
                "q": f'in:sent to:"{normalized}"',
                "maxResults": 1,
            },
        )
        if data.get("messages"):
            trusted = True
    except Exception as exc:
        print(f"[trust-list] sent-check failed for {normalized}: {repr(exc)}")

    # Check 2: prior-important history — ever labeled Priority / To Respond / Follow Up?
    if not trusted:
        label_clause = " OR ".join(f'label:"{name}"' for name in TRUSTED_INDICATOR_LABELS)
        try:
            data = await gmail_get_json_for_user(
                user_id=user_id,
                url=f"{GMAIL_API_BASE}/messages",
                params={
                    "q": f'from:"{normalized}" ({label_clause})',
                    "maxResults": 1,
                },
            )
            if data.get("messages"):
                trusted = True
        except Exception as exc:
            print(f"[trust-list] important-label check failed for {normalized}: {repr(exc)}")

    if cache is not None:
        cache[normalized] = trusted
    return trusted


# ============================================================
# SNOOZE — Gmail-level helpers
# ============================================================
async def snooze_thread_in_gmail(
    *,
    user_id: str,
    gmail_message_id: str,
    gmail_thread_id: str | None,
    original_label_id: str | None,
    snoozed_label_id: str,
) -> None:
    """Hide a mail/thread from Inbox by removing INBOX + the original status
    label and adding the OfficeFlow/Snoozed label.

    Applied to every message in the thread so the thread fully disappears
    from Inbox view (Gmail shows a thread in Inbox if any of its messages
    has the INBOX label).

    Best-effort: per-message exceptions are logged but don't abort the loop;
    the wake worker rebuilds state from `original_label_id` regardless.
    """
    remove_ids = ["INBOX"]
    if original_label_id:
        remove_ids.append(original_label_id)
    add_ids = [snoozed_label_id]

    if not gmail_thread_id:
        try:
            await modify_gmail_message_labels(
                user_id=user_id,
                gmail_message_id=gmail_message_id,
                add_label_ids=add_ids,
                remove_label_ids=remove_ids,
            )
        except Exception as exc:
            print(f"[snooze] msg {gmail_message_id}: {repr(exc)}")
        return

    try:
        thread_data = await gmail_get_json_for_user(
            user_id=user_id,
            url=f"{GMAIL_API_BASE}/threads/{gmail_thread_id}",
        )
    except Exception as exc:
        print(f"[snooze] thread fetch {gmail_thread_id}: {repr(exc)}")
        # Fallback: at least snooze the trigger message
        try:
            await modify_gmail_message_labels(
                user_id=user_id,
                gmail_message_id=gmail_message_id,
                add_label_ids=add_ids,
                remove_label_ids=remove_ids,
            )
        except Exception as exc2:
            print(f"[snooze] fallback msg {gmail_message_id}: {repr(exc2)}")
        return

    for thread_message in thread_data.get("messages", []) or []:
        thread_message_id = thread_message.get("id")
        if not thread_message_id:
            continue
        try:
            await modify_gmail_message_labels(
                user_id=user_id,
                gmail_message_id=thread_message_id,
                add_label_ids=add_ids,
                remove_label_ids=remove_ids,
            )
        except Exception as exc:
            print(f"[snooze] msg {thread_message_id}: {repr(exc)}")


async def wake_thread_in_gmail(
    *,
    user_id: str,
    gmail_message_id: str,
    gmail_thread_id: str | None,
    original_label_id: str | None,
    snoozed_label_id: str | None,
) -> None:
    """Restore a snoozed mail/thread to Inbox: add INBOX + original status
    label back, remove the OfficeFlow/Snoozed label.

    Mirror of snooze_thread_in_gmail. Safe to call multiple times — Gmail
    treats add/remove of an already-correct label as a no-op.
    """
    add_ids = ["INBOX"]
    if original_label_id:
        add_ids.append(original_label_id)
    remove_ids = [snoozed_label_id] if snoozed_label_id else []

    if not gmail_thread_id:
        try:
            await modify_gmail_message_labels(
                user_id=user_id,
                gmail_message_id=gmail_message_id,
                add_label_ids=add_ids,
                remove_label_ids=remove_ids,
            )
        except Exception as exc:
            print(f"[wake] msg {gmail_message_id}: {repr(exc)}")
        return

    try:
        thread_data = await gmail_get_json_for_user(
            user_id=user_id,
            url=f"{GMAIL_API_BASE}/threads/{gmail_thread_id}",
        )
    except Exception as exc:
        print(f"[wake] thread fetch {gmail_thread_id}: {repr(exc)}")
        try:
            await modify_gmail_message_labels(
                user_id=user_id,
                gmail_message_id=gmail_message_id,
                add_label_ids=add_ids,
                remove_label_ids=remove_ids,
            )
        except Exception as exc2:
            print(f"[wake] fallback msg {gmail_message_id}: {repr(exc2)}")
        return

    for thread_message in thread_data.get("messages", []) or []:
        thread_message_id = thread_message.get("id")
        if not thread_message_id:
            continue
        try:
            await modify_gmail_message_labels(
                user_id=user_id,
                gmail_message_id=thread_message_id,
                add_label_ids=add_ids,
                remove_label_ids=remove_ids,
            )
        except Exception as exc:
            print(f"[wake] msg {thread_message_id}: {repr(exc)}")


async def detect_current_status_label(
    *,
    user_id: str,
    gmail_message_id: str,
) -> tuple[str | None, str | None]:
    """Inspect the message's current Gmail labels and return the (name, id)
    of the OfficeFlow status label currently on it, if any.

    Returns (None, None) if the mail is unclassified (e.g. brand-new mail
    that hasn't been processed yet, or one that was already manually moved
    out of all status labels).
    """
    try:
        msg = await gmail_get_json_for_user(
            user_id=user_id,
            url=f"{GMAIL_API_BASE}/messages/{gmail_message_id}",
            params={"format": "minimal"},
        )
    except Exception as exc:
        print(f"[snooze] detect-status fetch {gmail_message_id}: {repr(exc)}")
        return (None, None)

    current_label_ids = set(msg.get("labelIds", []) or [])
    if not current_label_ids:
        return (None, None)

    # Resolve user's full label map once. We compare BY name so we tolerate
    # "Priority" and "OfficeFlow/Priority" as the same logical label.
    label_map = await get_all_gmail_labels(user_id)
    # Build inverse: label_id -> label_name
    id_to_name: dict[str, str] = {}
    for name, label in label_map.items():
        lbl_id = label.get("id") if isinstance(label, dict) else None
        if lbl_id:
            id_to_name[lbl_id] = name

    canonical_set = set(LABELS)
    for lbl_id in current_label_ids:
        name = id_to_name.get(lbl_id)
        if not name:
            continue
        # Strip the legacy "OfficeFlow/" prefix when matching.
        canonical = LEGACY_LABEL_NAME_MAP.get(name, name)
        if canonical in canonical_set:
            return (canonical, lbl_id)

    return (None, None)


async def ensure_label_exists(user_id: str, label_name: str) -> str:
    current_labels = await get_all_gmail_labels(user_id)

    # Exact match first, then case-insensitive fallback
    existing = current_labels.get(label_name)
    if not existing:
        for name, label in current_labels.items():
            if name.lower() == label_name.lower():
                existing = label
                break

    if existing:
        label_id = existing["id"]
        color = LABEL_COLORS.get(label_name)
        if color:
            try:
                await gmail_patch_json_for_user(
                    user_id=user_id,
                    url=f"{GMAIL_API_BASE}/labels/{label_id}",
                    payload={
                        "color": {
                            "textColor": color["textColor"],
                            "backgroundColor": color["backgroundColor"],
                        }
                    },
                )
            except Exception:
                pass
        return label_id

    payload: dict[str, Any] = {
        "name": label_name,
        "labelListVisibility": "labelShow",
        "messageListVisibility": "show",
    }
    color = LABEL_COLORS.get(label_name)
    if color:
        payload["color"] = {
            "textColor": color["textColor"],
            "backgroundColor": color["backgroundColor"],
        }

    def is_conflict_error(exc: Exception) -> bool:
        return "409" in str(exc) or "Label name exists" in str(exc)

    async def find_label_after_conflict() -> str:
        refreshed = await get_all_gmail_labels(user_id)
        for name, label in refreshed.items():
            if name.lower() == label_name.lower():
                return label["id"]
        raise HTTPException(status_code=500, detail=f"Label '{label_name}' conflict but not found after refresh")

    try:
        created = await gmail_post_json_for_user(
            user_id=user_id,
            url=f"{GMAIL_API_BASE}/labels",
            payload=payload,
        )
        return created["id"]
    except Exception as e:
        if is_conflict_error(e):
            return await find_label_after_conflict()
        # Not a conflict — retry without color
        payload.pop("color", None)
        try:
            created = await gmail_post_json_for_user(
                user_id=user_id,
                url=f"{GMAIL_API_BASE}/labels",
                payload=payload,
            )
            return created["id"]
        except Exception as e2:
            if is_conflict_error(e2):
                return await find_label_after_conflict()
            raise


async def delete_gmail_label_if_exists(user_id: str, label_name: str) -> dict[str, Any]:
    label_map = await get_all_gmail_labels(user_id)
    existing = label_map.get(label_name)

    if not existing:
        return {
            "label_name": label_name,
            "gmail_deleted": False,
            "reason": "not_found_in_gmail",
        }

    label_id = existing["id"]
    await gmail_delete_for_user(
        user_id=user_id,
        url=f"{GMAIL_API_BASE}/labels/{label_id}",
    )

    return {
        "label_name": label_name,
        "label_id": label_id,
        "gmail_deleted": True,
    }


async def cleanup_legacy_labels_for_mailbox(user_id: str, mailbox_id: str) -> list[dict[str, Any]]:
    results: list[dict[str, Any]] = []

    for legacy_label_name in LEGACY_LABEL_NAME_MAP.keys():
        gmail_result = await delete_gmail_label_if_exists(
            user_id=user_id,
            label_name=legacy_label_name,
        )

        db_deleted = await supabase_delete(
            (
                "/rest/v1/gmail_labels"
                f"?user_id=eq.{quote(user_id, safe='')}"
                f"&mailbox_id=eq.{quote(mailbox_id, safe='')}"
                f"&label_name=eq.{quote(legacy_label_name, safe='')}"
            )
        )

        results.append(
            {
                "label_name": legacy_label_name,
                "gmail": gmail_result,
                "db_deleted_rows": db_deleted if isinstance(db_deleted, list) else [],
            }
        )

    return results


async def sync_single_label(
    user_id: str,
    gmail_message_id: str,
    current_label_ids: set[str],
    label_name_to_id: dict[str, str],
    target_label_name: str,
) -> set[str]:
    target_label_id = label_name_to_id.get(target_label_name)
    if not target_label_id:
        raise HTTPException(status_code=500, detail=f"Missing Gmail label id for {target_label_name}")

    remove_label_ids: list[str] = []

    for label_name in get_status_label_names():
        label_id = label_name_to_id.get(label_name)
        if label_id and label_id in current_label_ids and label_id != target_label_id:
            remove_label_ids.append(label_id)

    add_label_ids: list[str] = []
    if target_label_id not in current_label_ids:
        add_label_ids.append(target_label_id)

    if add_label_ids or remove_label_ids:
        await modify_gmail_message_labels(
            user_id=user_id,
            gmail_message_id=gmail_message_id,
            add_label_ids=add_label_ids or None,
            remove_label_ids=remove_label_ids or None,
        )

        updated = set(current_label_ids)
        for label_id in remove_label_ids:
            updated.discard(label_id)
        for label_id in add_label_ids:
            updated.add(label_id)
        return updated

    return current_label_ids


async def sync_thread_status(
    user_id: str,
    thread_id: str | None,
    current_message_id: str,
    current_label_ids: set[str],
    label_name_to_id: dict[str, str],
    target_label_name: str,
) -> set[str]:
    if not thread_id:
        return await sync_single_label(
            user_id=user_id,
            gmail_message_id=current_message_id,
            current_label_ids=current_label_ids,
            label_name_to_id=label_name_to_id,
            target_label_name=target_label_name,
        )

    target_label_id = label_name_to_id.get(target_label_name)
    if not target_label_id:
        raise HTTPException(status_code=500, detail=f"Missing Gmail label id for {target_label_name}")

    removable_label_ids = get_status_label_ids_from_map(label_name_to_id)

    thread_data = await gmail_get_json_for_user(
        user_id=user_id,
        url=f"{GMAIL_API_BASE}/threads/{thread_id}",
    )

    updated_current_label_ids = set(current_label_ids)

    for thread_message in thread_data.get("messages", []):
        thread_message_id = thread_message.get("id")
        if not thread_message_id:
            continue

        thread_message_label_ids = set(thread_message.get("labelIds", []))

        remove_label_ids = [
            label_id
            for label_id in removable_label_ids
            if label_id in thread_message_label_ids and label_id != target_label_id
        ]

        add_label_ids: list[str] = []
        if target_label_id not in thread_message_label_ids:
            add_label_ids.append(target_label_id)

        if add_label_ids or remove_label_ids:
            await modify_gmail_message_labels(
                user_id=user_id,
                gmail_message_id=thread_message_id,
                add_label_ids=add_label_ids or None,
                remove_label_ids=remove_label_ids or None,
            )

        if thread_message_id == current_message_id:
            updated_current_label_ids = set(thread_message_label_ids)
            for label_id in remove_label_ids:
                updated_current_label_ids.discard(label_id)
            for label_id in add_label_ids:
                updated_current_label_ids.add(label_id)

    return updated_current_label_ids


def build_threaded_reply_raw(
    to_email: str,
    subject: str | None,
    body: str,
    original_message_id_header: str | None = None,
    references_header: str | None = None,
) -> str:
    message = MIMEText(body, "plain", "utf-8")
    message["To"] = to_email
    message["Subject"] = normalize_subject_for_reply(subject)

    if original_message_id_header:
        message["In-Reply-To"] = original_message_id_header
        if references_header:
            message["References"] = f"{references_header} {original_message_id_header}".strip()
        else:
            message["References"] = original_message_id_header

    raw_message = base64.urlsafe_b64encode(message.as_bytes()).decode()
    return raw_message


async def create_gmail_threaded_draft(
    user_id: str,
    to_email: str,
    subject: str | None,
    body: str,
    thread_id: str | None = None,
    original_message_id_header: str | None = None,
    references_header: str | None = None,
) -> dict[str, Any]:
    raw_message = build_threaded_reply_raw(
        to_email=to_email,
        subject=subject,
        body=body,
        original_message_id_header=original_message_id_header,
        references_header=references_header,
    )

    payload: dict[str, Any] = {"message": {"raw": raw_message}}
    if thread_id:
        payload["message"]["threadId"] = thread_id

    return await gmail_post_json_for_user(
        user_id=user_id,
        url=f"{GMAIL_API_BASE}/drafts",
        payload=payload,
    )


async def send_gmail_draft(user_id: str, gmail_draft_id: str) -> dict[str, Any]:
    """Send an existing Gmail draft. Whatever content is currently in the
    draft (including any user edits in Gmail) is what gets sent. The draft
    is consumed and removed from the drafts folder afterwards."""
    return await gmail_post_json_for_user(
        user_id=user_id,
        url=f"{GMAIL_API_BASE}/drafts/send",
        payload={"id": gmail_draft_id},
    )


# ----------------------------
# Attachment AI helpers
# ----------------------------

# Limits — keep cost predictable and avoid hammering OpenAI on huge files.
ATTACHMENT_MAX_SIZE_BYTES = int(os.getenv("ATTACHMENT_MAX_SIZE_BYTES", "5242880"))  # 5 MB
ATTACHMENT_MAX_PER_MESSAGE = int(os.getenv("ATTACHMENT_MAX_PER_MESSAGE", "3"))
ATTACHMENT_MAX_TEXT_CHARS = int(os.getenv("ATTACHMENT_MAX_TEXT_CHARS", "30000"))  # ~ first ~7-8 pages

SUPPORTED_ATTACHMENT_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # .docx
    "application/msword",  # legacy .doc — pypdf/docx won't extract; will mark unsupported
}

# Only run attachment summarisation when the AI classified the mail into one
# of these labels. Saves cost on Marketing / Notification / Ignore noise.
ATTACHMENT_SUMMARY_LABELS = {"To Respond", "Priority", "Follow Up"}


def _walk_message_parts(payload: dict[str, Any]) -> list[dict[str, Any]]:
    """Flatten Gmail message payload into a list of all parts (recursive)."""
    parts: list[dict[str, Any]] = []
    if not payload:
        return parts
    parts.append(payload)
    for child in (payload.get("parts") or []):
        parts.extend(_walk_message_parts(child))
    return parts


def find_attachments_in_message(message_data: dict[str, Any]) -> list[dict[str, Any]]:
    """Return a list of {filename, mime_type, size, attachment_id} for every
    attached file in the Gmail message. Inline parts (no filename) are skipped."""
    out: list[dict[str, Any]] = []
    payload = message_data.get("payload") or {}
    for part in _walk_message_parts(payload):
        filename = (part.get("filename") or "").strip()
        body = part.get("body") or {}
        attachment_id = body.get("attachmentId")
        if not filename or not attachment_id:
            continue
        out.append({
            "filename": filename,
            "mime_type": part.get("mimeType") or "",
            "size_bytes": int(body.get("size") or 0),
            "attachment_id": attachment_id,
        })
    return out


async def download_gmail_attachment_bytes(
    user_id: str,
    gmail_message_id: str,
    attachment_id: str,
) -> bytes:
    """Fetch the raw bytes of a Gmail attachment via the API."""
    data = await gmail_get_json_for_user(
        user_id=user_id,
        url=f"{GMAIL_API_BASE}/messages/{gmail_message_id}/attachments/{attachment_id}",
    )
    raw_b64 = (data or {}).get("data") or ""
    # Gmail uses URL-safe base64 with no padding; pad to multiple of 4.
    padding = "=" * (-len(raw_b64) % 4)
    return base64.urlsafe_b64decode(raw_b64 + padding)


def extract_text_from_pdf_bytes(blob: bytes) -> str:
    """Extract text content from a PDF. Returns empty string if extraction fails."""
    try:
        from pypdf import PdfReader  # local import to keep startup fast
        from io import BytesIO
        reader = PdfReader(BytesIO(blob))
        chunks: list[str] = []
        for page in reader.pages:
            try:
                chunks.append(page.extract_text() or "")
            except Exception:
                continue
        return "\n\n".join(c for c in chunks if c).strip()
    except Exception as exc:
        print(f"[attachment-pdf] extract failed: {repr(exc)}")
        return ""


def extract_text_from_docx_bytes(blob: bytes) -> str:
    """Extract text content from a .docx file. Returns empty string on failure."""
    try:
        from docx import Document  # python-docx
        from io import BytesIO
        doc = Document(BytesIO(blob))
        paragraphs = [p.text for p in doc.paragraphs if (p.text or "").strip()]
        # Tables — flatten into pipe-separated rows so the LLM can read them.
        for table in doc.tables:
            for row in table.rows:
                cells = [(c.text or "").strip() for c in row.cells]
                if any(cells):
                    paragraphs.append(" | ".join(cells))
        return "\n".join(paragraphs).strip()
    except Exception as exc:
        print(f"[attachment-docx] extract failed: {repr(exc)}")
        return ""


def extract_text_from_attachment(filename: str, mime_type: str, blob: bytes) -> tuple[str, str]:
    """Returns (extracted_text, status). Status is 'success' or 'unsupported'."""
    name_lower = (filename or "").lower()
    mt = (mime_type or "").lower()
    if mt == "application/pdf" or name_lower.endswith(".pdf"):
        text = extract_text_from_pdf_bytes(blob)
        return text, ("success" if text else "failed")
    if (
        mt == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or name_lower.endswith(".docx")
    ):
        text = extract_text_from_docx_bytes(blob)
        return text, ("success" if text else "failed")
    return "", "unsupported"


async def summarize_attachment_text(filename: str, text: str) -> dict[str, Any]:
    """Send extracted text to OpenAI and ask for a summary + structured data.
    Returns {summary: str, key_data: dict} — never raises; falls back to
    {summary: '', key_data: {}} on error."""
    if not text:
        return {"summary": "", "key_data": {}}

    api_key = require_env(OPENAI_API_KEY, "OPENAI_API_KEY")

    # Truncate aggressively — cost control. The first N chars of a business
    # document almost always contain the substance (date, amount, parties).
    snippet = text[:ATTACHMENT_MAX_TEXT_CHARS]

    prompt = f"""Je bent een zakelijke assistent. Je krijgt de tekst van een bijlage uit een e-mail. Geef een zo concreet mogelijke samenvatting plus de belangrijkste feiten in een gestructureerd formaat.

Bestandsnaam: {filename}

Inhoud:
\"\"\"
{snippet}
\"\"\"

Antwoord ALLEEN met JSON in dit formaat:
{{
  "summary": "2-3 zinnen die de essentie pakken in zakelijk Nederlands",
  "key_data": {{
    "amounts": ["bedragen incl. valuta + waarvoor, bijv. '€18.500 ex BTW — totaal offerte'"],
    "dates": ["belangrijke datums + waarvoor"],
    "deadlines": ["uiterste data + waarvoor"],
    "parties": ["genoemde organisaties/personen"],
    "key_clauses": ["belangrijke voorwaarden of bepalingen die de lezer moet weten"]
  }}
}}

Regels:
- Lege array als veld niet van toepassing is — niet verzinnen.
- Hou het concreet en kort.
- Geen interpretatie, alleen wat in de tekst staat.
- Geen markdown, geen uitleg, alleen valide JSON."""

    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            response = await client.post(
                "https://api.openai.com/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json",
                },
                json={
                    "model": "gpt-4o-mini",
                    "messages": [{"role": "user", "content": prompt}],
                    "temperature": 0.1,
                    "response_format": {"type": "json_object"},
                },
            )
        if response.status_code != 200:
            print(f"[attachment-ai] OpenAI {response.status_code}: {response.text[:200]}")
            return {"summary": "", "key_data": {}}
        data = response.json()
        content = data["choices"][0]["message"]["content"] or "{}"
        parsed = json.loads(content)
        return {
            "summary": (parsed.get("summary") or "").strip(),
            "key_data": parsed.get("key_data") or {},
        }
    except Exception as exc:
        print(f"[attachment-ai] summarize failed: {repr(exc)}")
        return {"summary": "", "key_data": {}}


async def supabase_insert_attachment_summary(
    user_id: str,
    email_id: str | None,
    gmail_message_id: str,
    gmail_attachment_id: str | None,
    filename: str,
    mime_type: str | None,
    size_bytes: int | None,
    summary: str,
    key_data: dict[str, Any] | None,
    status: str,
    error_message: str | None = None,
) -> dict[str, Any] | None:
    payload = {
        "user_id": user_id,
        "email_id": email_id,
        "gmail_message_id": gmail_message_id,
        "gmail_attachment_id": gmail_attachment_id,
        "filename": filename,
        "mime_type": mime_type,
        "size_bytes": size_bytes,
        "summary": summary or None,
        "key_data": key_data or None,
        "status": status,
        "error_message": error_message,
    }
    try:
        data = await supabase_post(
            "/rest/v1/attachment_summaries"
            "?on_conflict=gmail_message_id,gmail_attachment_id",
            [payload],
            prefer="resolution=merge-duplicates,return=representation",
        )
        if isinstance(data, list) and data:
            return data[0]
    except Exception as exc:
        print(f"[attachment-summary-db] insert failed for {filename}: {repr(exc)}")
    return None


async def supabase_get_attachment_summaries_for_emails(
    user_id: str,
    email_ids: list[str],
) -> list[dict[str, Any]]:
    if not email_ids:
        return []
    in_clause = ",".join(quote(eid, safe="") for eid in email_ids if eid)
    if not in_clause:
        return []
    data = await supabase_get(
        "/rest/v1/attachment_summaries"
        f"?user_id=eq.{quote(user_id, safe='')}"
        f"&email_id=in.({in_clause})"
        f"&select=*"
    )
    return data if isinstance(data, list) else []


async def supabase_get_attachment_summaries_for_message(
    user_id: str,
    gmail_message_id: str,
) -> list[dict[str, Any]]:
    """Fetch all (successful or failed) attachment summaries for a single
    Gmail message. Used by generate_ai_reply to inject concrete facts
    from PDFs/DOCX into the draft."""
    data = await supabase_get(
        "/rest/v1/attachment_summaries"
        f"?user_id=eq.{quote(user_id, safe='')}"
        f"&gmail_message_id=eq.{quote(gmail_message_id, safe='')}"
        f"&select=*"
    )
    return data if isinstance(data, list) else []


async def process_message_attachments(
    user_id: str,
    email_id: str | None,
    gmail_message_id: str,
    message_data: dict[str, Any],
    target_label: str,
) -> int:
    """Find, extract, summarize and persist attachments for a single message.
    Returns the number of attachments actually processed (success or fail)."""
    print(f"[attachment-ai] called for msg {gmail_message_id} label={target_label!r}")

    if target_label not in ATTACHMENT_SUMMARY_LABELS:
        print(f"[attachment-ai] skip msg {gmail_message_id}: label {target_label!r} not in supported set")
        return 0

    attachments = find_attachments_in_message(message_data)
    print(f"[attachment-ai] msg {gmail_message_id}: found {len(attachments)} attachment(s)")
    if not attachments:
        # Diagnostic: show the top-level mime type so we know if Gmail gave us
        # the right payload structure.
        top_mime = (message_data.get("payload") or {}).get("mimeType")
        has_parts = bool((message_data.get("payload") or {}).get("parts"))
        print(f"[attachment-ai] msg {gmail_message_id}: top mime={top_mime!r} has_parts={has_parts}")
        return 0

    # Cap to N attachments per message — anything beyond is unusual and costly.
    attachments = attachments[:ATTACHMENT_MAX_PER_MESSAGE]
    processed = 0

    for att in attachments:
        filename = att["filename"]
        mime_type = att["mime_type"]
        size_bytes = att["size_bytes"]
        attachment_id = att["attachment_id"]

        # Size guard
        if size_bytes and size_bytes > ATTACHMENT_MAX_SIZE_BYTES:
            await supabase_insert_attachment_summary(
                user_id=user_id,
                email_id=email_id,
                gmail_message_id=gmail_message_id,
                gmail_attachment_id=attachment_id,
                filename=filename,
                mime_type=mime_type,
                size_bytes=size_bytes,
                summary="",
                key_data=None,
                status="too_large",
                error_message=f"File exceeds {ATTACHMENT_MAX_SIZE_BYTES} bytes",
            )
            processed += 1
            continue

        # Type guard — only spend OpenAI cost on supported types
        mt = (mime_type or "").lower()
        name_lower = filename.lower()
        is_supported = (
            mt in SUPPORTED_ATTACHMENT_MIME_TYPES
            or name_lower.endswith((".pdf", ".docx"))
        )
        if not is_supported:
            await supabase_insert_attachment_summary(
                user_id=user_id,
                email_id=email_id,
                gmail_message_id=gmail_message_id,
                gmail_attachment_id=attachment_id,
                filename=filename,
                mime_type=mime_type,
                size_bytes=size_bytes,
                summary="",
                key_data=None,
                status="unsupported",
            )
            processed += 1
            continue

        # Download + extract + summarize. Each step has its own try/except
        # so a single bad attachment never poisons the rest of the message.
        try:
            blob = await download_gmail_attachment_bytes(
                user_id=user_id,
                gmail_message_id=gmail_message_id,
                attachment_id=attachment_id,
            )
        except Exception as exc:
            await supabase_insert_attachment_summary(
                user_id=user_id, email_id=email_id,
                gmail_message_id=gmail_message_id, gmail_attachment_id=attachment_id,
                filename=filename, mime_type=mime_type, size_bytes=size_bytes,
                summary="", key_data=None, status="failed",
                error_message=f"download failed: {repr(exc)[:240]}",
            )
            processed += 1
            continue

        text, extract_status = extract_text_from_attachment(filename, mime_type, blob)
        if extract_status != "success" or not text:
            await supabase_insert_attachment_summary(
                user_id=user_id, email_id=email_id,
                gmail_message_id=gmail_message_id, gmail_attachment_id=attachment_id,
                filename=filename, mime_type=mime_type, size_bytes=size_bytes,
                summary="", key_data=None,
                status=("unsupported" if extract_status == "unsupported" else "failed"),
                error_message=("text extraction returned empty" if extract_status == "success" else None),
            )
            processed += 1
            continue

        summary_obj = await summarize_attachment_text(filename, text)
        final_status = "success" if summary_obj.get("summary") else "failed"
        print(
            f"[attachment-ai] msg {gmail_message_id} file={filename!r} "
            f"text_len={len(text)} status={final_status}"
        )
        await supabase_insert_attachment_summary(
            user_id=user_id, email_id=email_id,
            gmail_message_id=gmail_message_id, gmail_attachment_id=attachment_id,
            filename=filename, mime_type=mime_type, size_bytes=size_bytes,
            summary=summary_obj.get("summary") or "",
            key_data=summary_obj.get("key_data") or {},
            status=final_status,
            error_message=(None if summary_obj.get("summary") else "OpenAI returned empty summary"),
        )
        processed += 1

    print(f"[attachment-ai] msg {gmail_message_id}: processed {processed} attachment(s)")
    return processed


# ----------------------------
# Thread state helpers
# ----------------------------

def get_message_direction(message_data: dict[str, Any], mailbox_email: str | None) -> str:
    label_ids = set(message_data.get("labelIds", []))

    if is_draft_label_ids(label_ids):
        return "draft"

    payload = message_data.get("payload", {})
    headers = payload.get("headers", [])
    from_header = get_header_value(headers, "From")
    from_email = extract_email_address(from_header)

    if "SENT" in label_ids:
        return "sent_by_user"

    if mailbox_email and from_email and from_email.lower() == mailbox_email.lower():
        return "sent_by_user"

    return "incoming"


async def get_thread_reply_state(
    user_id: str,
    thread_id: str | None,
    mailbox_email: str | None,
) -> dict[str, Any]:
    if not thread_id:
        return {
            "has_user_reply": False,
            "user_is_latest_sender": False,
            "needs_response_after_reply": False,
            "latest_message": None,
            "latest_user_message": None,
            "latest_incoming_message": None,
            "has_open_draft": False,
            "latest_draft_message": None,
        }

    thread_data = await gmail_get_json_for_user(
        user_id=user_id,
        url=f"{GMAIL_API_BASE}/threads/{thread_id}",
    )

    latest_message = None
    latest_message_date = 0

    latest_user_message = None
    latest_user_sent_date = 0

    latest_incoming_message = None
    latest_incoming_date = 0

    latest_draft_message = None
    latest_draft_date = 0
    has_open_draft = False

    for thread_message in thread_data.get("messages", []):
        internal_date = parse_internal_date_ms(thread_message.get("internalDate"))
        direction = get_message_direction(thread_message, mailbox_email)

        if direction == "draft":
            has_open_draft = True
            if internal_date >= latest_draft_date:
                latest_draft_message = thread_message
                latest_draft_date = internal_date
            continue

        if internal_date >= latest_message_date:
            latest_message = thread_message
            latest_message_date = internal_date

        if direction == "sent_by_user":
            if internal_date >= latest_user_sent_date:
                latest_user_message = thread_message
                latest_user_sent_date = internal_date
        else:
            if internal_date >= latest_incoming_date:
                latest_incoming_message = thread_message
                latest_incoming_date = internal_date

    has_user_reply = latest_user_sent_date > 0
    user_is_latest_sender = has_user_reply and latest_user_sent_date > latest_incoming_date
    needs_response_after_reply = has_user_reply and latest_incoming_date > latest_user_sent_date

    return {
        "has_user_reply": has_user_reply,
        "user_is_latest_sender": user_is_latest_sender,
        "needs_response_after_reply": needs_response_after_reply,
        "latest_message": latest_message,
        "latest_user_message": latest_user_message,
        "latest_incoming_message": latest_incoming_message,
        "has_open_draft": has_open_draft,
        "latest_draft_message": latest_draft_message,
    }


# ----------------------------
# Labels
# ----------------------------

async def setup_gmail_labels_for_mailbox(user_id: str, mailbox_id: str) -> list[dict[str, Any]]:
    results: list[dict[str, Any]] = []

    # Status labels only (Priority / To Respond / ...). Confidence is tracked
    # internally and surfaced via the API response — no Gmail labels, to
    # keep the inbox clean.
    for label_name in LABELS:
        label_id = await ensure_label_exists(user_id=user_id, label_name=label_name)

        saved = await supabase_upsert_gmail_label(
            user_id=user_id,
            mailbox_id=mailbox_id,
            label_name=label_name,
            label_id=label_id,
        )

        results.append(
            {
                "label_name": label_name,
                "label_id": label_id,
                "db_row": saved,
            }
        )

    # One-time cleanup of legacy AI confidence labels (AI · Sure / Check /
    # Review). Deleting the Gmail label strips it from every message at
    # once. Idempotent — after the first deletion every subsequent call is
    # a no-op.
    for legacy_name in CONFIDENCE_LEGACY_LABEL_NAMES:
        try:
            await delete_gmail_label_if_exists(user_id=user_id, label_name=legacy_name)
            await supabase_delete(
                "/rest/v1/gmail_labels"
                f"?user_id=eq.{quote(user_id, safe='')}"
                f"&mailbox_id=eq.{quote(mailbox_id, safe='')}"
                f"&label_name=eq.{quote(legacy_name, safe='')}"
            )
        except Exception as exc:
            print(f"[cleanup-confidence-label {legacy_name}] {repr(exc)}")

    return results


# ----------------------------
# AI
# ----------------------------

# Unsubscribe / opt-out detection.
# A mail asking to be removed from a list MUST be classified as To Respond,
# never Ignore/Marketing/Notification — for GDPR opt-out compliance. We
# enforce this at two levels: (1) explicit rule in the classifier prompt,
# (2) post-classification override if the AI gets it wrong. The override
# is the safety net.
_UNSUBSCRIBE_PATTERNS = re.compile(
    r"\b("
    r"unsubscribe|uitschrijven|uitschrijving|afmelden|afmelding|"
    r"opt[\s\-]?out|opted[\s\-]?out|"
    r"verwijder\s+mij|verwijder\s+me|remove\s+me|please\s+remove|"
    r"haal\s+mij\s+(uit|van)|haal\s+me\s+(uit|van)|"
    r"graag\s+afmelden|graag\s+uitschrijven|wil\s+(mij|me)\s+afmelden"
    r")\b",
    re.IGNORECASE,
)


def detect_unsubscribe_request(subject: str | None, body_text: str | None) -> bool:
    """True when the mail clearly is a request to be unsubscribed from a
    mailing list. Scoped to subject hits OR short bodies so we don't flag
    newsletter footers that also contain the word 'unsubscribe'."""
    subject_lower = (subject or "").lower()
    body_text_str = body_text or ""
    body_lower = body_text_str.lower()

    # Strongest signal: subject says it directly
    if _UNSUBSCRIBE_PATTERNS.search(subject_lower):
        return True

    # Medium signal: short body where unsubscribe is the main intent.
    # 800 chars covers polite Dutch unsubscribe requests including a
    # short polite signature, but excludes newsletter footers (those are
    # always inside long marketing bodies).
    if len(body_text_str) <= 800 and _UNSUBSCRIBE_PATTERNS.search(body_lower):
        return True

    return False


async def classify_email(
    subject: str | None,
    sender: str | None,
    body_text: str | None,
    mailbox_email: str | None = None,
) -> dict[str, Any]:
    api_key = require_env(OPENAI_API_KEY, "OPENAI_API_KEY")

    # ============ EARLY OVERRIDE — OFFICEFLOW SELF-SENT MAILS ============
    # Briefings, weekly recaps, silence-radar summaries that OfficeFlow sends
    # from the user's mailbox to themselves are always FYI. Skip the OpenAI
    # call entirely — saves tokens and guarantees consistent labelling.
    sender_addr_for_self = extract_email_address(sender)
    if is_self_sent_officeflow_mail(sender_addr_for_self, mailbox_email, subject):
        return {
            "label": "FYI",
            "reason": "OfficeFlow self-sent briefing/recap mail — always FYI",
            "confidence": "high",
            "generate_draft": False,
        }

    # Safeguard: lege body -> niet classificeren, FYI als veilig default
    if not body_text or len(body_text.strip()) < 10:
        return {
            "label": "FYI",
            "reason": "Body text too short or empty to classify",
            "confidence": "low",
            "generate_draft": False,
        }

    # Trunc body voor classifier (full body wordt alleen bij draft gebruikt)
    body_for_prompt = (body_text or "")[:3000]

    prompt = f"""
Je bent een e-mail classifier voor OfficeFlow.

Kies exact 1 label uit deze lijst:
- Priority
- To Respond
- FYI
- Notification
- Marketing
- Ignore

Regels:
- Priority: ALLEEN bij duidelijke urgentie of directe business impact. Bijvoorbeeld: harde deadline binnen 24-48 uur, klant met klacht of blokkerend probleem, spoedopdracht, uitval, verlies van omzet. Een normale klantmail is GEEN Priority.
- To Respond: mail waar de gebruiker inhoudelijk op moet antwoorden. Dit is de standaard voor normale klant- en zakelijke mails zonder acute urgentie.
- FYI: informatief, geen reactie nodig, maar wel relevant genoeg om te lezen.
- Notification: automatische systeemmelding, statusupdate, bevestiging, receipt, log.
- Marketing: nieuwsbrief, promotie, sales outreach, cold outreach, aanbieding.
- Ignore: duidelijk irrelevant, lage kwaliteit, scraper, phishing-achtig, of rommel.

Belangrijk:
- Bij twijfel tussen Priority en To Respond -> kies To Respond.
- Bij twijfel tussen Marketing en Ignore -> kies Marketing (veiliger voor gebruiker).

KRITISCHE REGEL — Uitschrijf / afmelding / opt-out requests:
- Als de mail expliciet vraagt om afmelding, uitschrijving, opt-out, "remove me", "unsubscribe", "verwijder mij", "haal me uit jullie lijst" of vergelijkbaar -> ALTIJD label "To Respond".
- Dit overschrijft alle andere signalen. NOOIT Ignore, Marketing, Notification of FYI voor expliciete uitschrijf-verzoeken.
- Reden: AVG/GDPR opt-out compliance — afzender moet binnen redelijke termijn uit alle mailing lists verwijderd worden.

Geef ook een confidence score terug (hoe zeker ben je van de classificatie):
- "high": signalen in de mail zijn eenduidig; geen significante twijfel tussen labels.
- "medium": waarschijnlijk correct, maar er zijn 1-2 signalen die ook een ander label zouden kunnen rechtvaardigen.
- "low": onduidelijk, korte of dubbelzinnige mail, mogelijke misclassificatie.

Geef alleen geldige JSON terug in exact dit formaat:
{{
  "label": "To Respond",
  "reason": "Korte reden",
  "confidence": "high"
}}

Van: {sender}
Onderwerp: {subject}

E-mail:
{body_for_prompt}
""".strip()

    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            response = await client.post(
                "https://api.openai.com/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json",
                },
                json={
                    "model": "gpt-4o-mini",
                    "messages": [
                        {
                            "role": "system",
                            "content": "Je classificeert zakelijke e-mails en geeft alleen geldige JSON terug.",
                        },
                        {
                            "role": "user",
                            "content": prompt,
                        },
                    ],
                    "temperature": 0,
                },
            )
    except Exception as exc:
        print(f"[classify_email] OpenAI call failed: {repr(exc)} -- falling back to 'To Respond'")
        return {
            "label": "To Respond",
            "reason": f"Classifier unavailable: {type(exc).__name__}",
            "confidence": "low",
            "generate_draft": LABEL_RULES["To Respond"]["generate_draft"],
        }

    data = parse_response_data(response)
    if response.status_code >= 400:
        print(f"[classify_email] OpenAI {response.status_code}: {data} -- falling back to 'To Respond'")
        return {
            "label": "To Respond",
            "reason": f"Classifier HTTP {response.status_code}",
            "confidence": "low",
            "generate_draft": LABEL_RULES["To Respond"]["generate_draft"],
        }

    try:
        content = data["choices"][0]["message"]["content"]
        parsed = safe_parse_json(content)
    except Exception:
        print(f"[classify_email] Invalid response shape: {data} -- falling back to 'To Respond'")
        return {
            "label": "To Respond",
            "reason": "Classifier returned invalid JSON",
            "confidence": "low",
            "generate_draft": LABEL_RULES["To Respond"]["generate_draft"],
        }

    label = parsed.get("label")
    reason = parsed.get("reason") or "No reason given"
    confidence = normalize_confidence(parsed.get("confidence"))

    if label not in CLASSIFIER_LABELS:
        print(f"[classify_email] Unknown label '{label}' -- falling back to 'To Respond'")
        return {
            "label": "To Respond",
            "reason": f"Classifier returned unknown label '{label}'",
            "confidence": "low",
            "generate_draft": LABEL_RULES["To Respond"]["generate_draft"],
        }

    # ============ POST-CLASSIFICATION OVERRIDES ============
    # These are belt-and-braces safety nets that apply AFTER the AI classifier.
    # Order matters: noreply check first (a noreply "unsubscribe confirmed"
    # is a notification, not a real opt-out request from a person).
    sender_addr = extract_email_address(sender)
    sender_is_noreply = is_noreply_sender(sender_addr)

    # 1) NOREPLY → never get To Respond / Priority. You can't reply to them.
    if sender_is_noreply and label in ("To Respond", "Priority"):
        print(
            f"[classify_email] Noreply override: '{label}' -> 'Notification' "
            f"because sender '{sender_addr}' is a noreply / system address"
        )
        return {
            "label": "Notification",
            "reason": (
                f"Noreply override: sender '{sender_addr}' cannot receive replies. "
                f"Classifier originally said '{label}'."
            ),
            "confidence": "high",
            "generate_draft": LABEL_RULES["Notification"]["generate_draft"],
        }

    # 2) UNSUBSCRIBE / OPT-OUT — but only when sender is a real person.
    # A noreply system saying "you have been unsubscribed" is just a
    # notification of an event, not a request the user must process.
    if not sender_is_noreply and label in ("Ignore", "Marketing", "Notification", "FYI"):
        if detect_unsubscribe_request(subject, body_text):
            print(
                f"[classify_email] Compliance override: '{label}' -> 'To Respond' "
                f"because mail looks like an unsubscribe / opt-out request "
                f"(subject={subject!r})"
            )
            return {
                "label": "To Respond",
                "reason": (
                    f"Compliance override: detected unsubscribe / opt-out "
                    f"request (GDPR — must be processed by user). "
                    f"Classifier originally said '{label}'."
                ),
                "confidence": "high",
                "generate_draft": LABEL_RULES["To Respond"]["generate_draft"],
            }

    return {
        "label": label,
        "reason": reason,
        "confidence": confidence,
        "generate_draft": LABEL_RULES[label]["generate_draft"],
    }


async def classify_follow_up_email(
    subject: str | None,
    sender: str | None,
    body_text: str | None,
    mailbox_email: str | None = None,
) -> dict[str, Any]:
    api_key = require_env(OPENAI_API_KEY, "OPENAI_API_KEY")

    # Same self-sent OfficeFlow override as classify_email.
    sender_addr_for_self = extract_email_address(sender)
    if is_self_sent_officeflow_mail(sender_addr_for_self, mailbox_email, subject):
        return {
            "label": "FYI",
            "reason": "OfficeFlow self-sent briefing/recap mail — always FYI",
            "confidence": "high",
            "generate_draft": False,
        }

    # Safeguard: lege body -> niet classificeren, FYI als veilig default
    if not body_text or len(body_text.strip()) < 10:
        return {
            "label": "FYI",
            "reason": "Body text too short or empty to classify",
            "confidence": "low",
            "generate_draft": False,
        }

    body_for_prompt = (body_text or "")[:3000]

    prompt = f"""
Je bent een e-mail classifier voor OfficeFlow.

Context:
- Dit is een bestaand gesprek.
- De gebruiker heeft eerder al echt gereageerd in deze thread.
- Er is nu weer een nieuw inkomend bericht binnengekomen.
- Jij moet bepalen of deze thread opnieuw open moet staan als actiepunt, of juist inhoudelijk klaar is.

Kies exact 1 label uit deze lijst:
- Priority
- To Respond
- Waiting On Reply
- Done
- FYI
- Notification
- Ignore

Regels:
- Priority: er is nu duidelijke urgentie, business impact of snelle actie nodig
- To Respond: de gebruiker moet nu weer inhoudelijk reageren of actie nemen
- Waiting On Reply: het gesprek is nog actief/open, maar dit laatste bericht vraagt niet direct om een reactie; gebruik dit spaarzaam
- Done: het gesprek is inhoudelijk afgerond, bevestigd of afgesloten; er is geen verdere actie nodig
- FYI: informatief, geen actie nodig
- Notification: automatische melding of systeemupdate
- Ignore: irrelevant of ongewenst

Belangrijke voorkeur:
- Als het laatste inkomende bericht dingen zegt als "bedankt", "duidelijk", "meer hoef ik niet te weten", "helemaal goed", "is prima", "alles is geregeld", kies dan Done.
- Kies alleen To Respond als de gebruiker nu echt weer iets moet doen.
- Kies Waiting On Reply alleen als het gesprek nog open voelt maar niet echt klaar is.
- Bij twijfel tussen To Respond en een andere label -> kies To Respond.

KRITISCHE REGEL — Uitschrijf / afmelding / opt-out requests:
- Als de afzender expliciet vraagt om afmelding, uitschrijving, opt-out, "remove me", "unsubscribe", "verwijder mij", "haal me uit jullie lijst" of vergelijkbaar -> ALTIJD label "To Respond".
- Dit overschrijft alle andere signalen. NOOIT Ignore, Notification, FYI of Done voor expliciete uitschrijf-verzoeken.
- Reden: AVG/GDPR opt-out compliance — afzender moet binnen redelijke termijn uit alle mailing lists verwijderd worden.

Geef ook een confidence score terug (hoe zeker ben je van de classificatie):
- "high": signalen in de mail zijn eenduidig; geen significante twijfel tussen labels.
- "medium": waarschijnlijk correct, maar er zijn 1-2 signalen die ook een ander label zouden kunnen rechtvaardigen.
- "low": onduidelijk, korte of dubbelzinnige mail, mogelijke misclassificatie.

Geef alleen geldige JSON terug in exact dit formaat:
{{
  "label": "Done",
  "reason": "Korte reden",
  "confidence": "high"
}}

Van: {sender}
Onderwerp: {subject}

E-mail:
{body_for_prompt}
""".strip()

    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            response = await client.post(
                "https://api.openai.com/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json",
                },
                json={
                    "model": "gpt-4o-mini",
                    "messages": [
                        {
                            "role": "system",
                            "content": "Je classificeert vervolgberichten in bestaande e-mailthreads en geeft alleen geldige JSON terug.",
                        },
                        {
                            "role": "user",
                            "content": prompt,
                        },
                    ],
                    "temperature": 0,
                },
            )
    except Exception as exc:
        print(f"[classify_follow_up_email] OpenAI call failed: {repr(exc)} -- fallback 'To Respond'")
        return {
            "label": "To Respond",
            "reason": f"Classifier unavailable: {type(exc).__name__}",
            "confidence": "low",
            "generate_draft": LABEL_RULES["To Respond"]["generate_draft"],
        }

    data = parse_response_data(response)
    if response.status_code >= 400:
        print(f"[classify_follow_up_email] OpenAI {response.status_code}: {data} -- fallback 'To Respond'")
        return {
            "label": "To Respond",
            "reason": f"Classifier HTTP {response.status_code}",
            "confidence": "low",
            "generate_draft": LABEL_RULES["To Respond"]["generate_draft"],
        }

    try:
        content = data["choices"][0]["message"]["content"]
        parsed = safe_parse_json(content)
    except Exception:
        print(f"[classify_follow_up_email] Invalid response shape: {data} -- fallback 'To Respond'")
        return {
            "label": "To Respond",
            "reason": "Classifier returned invalid JSON",
            "confidence": "low",
            "generate_draft": LABEL_RULES["To Respond"]["generate_draft"],
        }

    label = parsed.get("label")
    reason = parsed.get("reason") or "No reason given"
    confidence = normalize_confidence(parsed.get("confidence"))

    if label not in FOLLOW_UP_CLASSIFIER_LABELS:
        print(f"[classify_follow_up_email] Unknown label '{label}' -- fallback 'To Respond'")
        return {
            "label": "To Respond",
            "reason": f"Classifier returned unknown label '{label}'",
            "confidence": "low",
            "generate_draft": LABEL_RULES["To Respond"]["generate_draft"],
        }

    # ============ POST-CLASSIFICATION OVERRIDES ============
    # Same two safety nets as classify_email. Order matters.
    sender_addr = extract_email_address(sender)
    sender_is_noreply = is_noreply_sender(sender_addr)

    # 1) NOREPLY → never To Respond / Priority. You can't reply to them.
    if sender_is_noreply and label in ("To Respond", "Priority"):
        print(
            f"[classify_follow_up_email] Noreply override: '{label}' -> 'Notification' "
            f"because sender '{sender_addr}' is a noreply / system address"
        )
        return {
            "label": "Notification",
            "reason": (
                f"Noreply override: sender '{sender_addr}' cannot receive replies. "
                f"Classifier originally said '{label}'."
            ),
            "confidence": "high",
            "generate_draft": LABEL_RULES["Notification"]["generate_draft"],
        }

    # 2) UNSUBSCRIBE — only for real human senders (not noreply confirmations).
    # Marketing included to mirror initial classifier (classify_email) so
    # opt-out / unsubscribe mails from real senders never get buried.
    if not sender_is_noreply and label in ("Done", "Ignore", "Marketing", "Notification", "FYI", "Waiting On Reply"):
        if detect_unsubscribe_request(subject, body_text):
            print(
                f"[classify_follow_up_email] Compliance override: '{label}' -> 'To Respond' "
                f"because mail looks like an unsubscribe / opt-out request "
                f"(subject={subject!r})"
            )
            return {
                "label": "To Respond",
                "reason": (
                    f"Compliance override: detected unsubscribe / opt-out "
                    f"request (GDPR — must be processed by user). "
                    f"Classifier originally said '{label}'."
                ),
                "confidence": "high",
                "generate_draft": LABEL_RULES["To Respond"]["generate_draft"],
            }

    return {
        "label": label,
        "reason": reason,
        "confidence": confidence,
        "generate_draft": LABEL_RULES[label]["generate_draft"],
    }


async def classify_latest_sent_reply_status(subject: str | None, body_text: str | None) -> dict[str, Any]:
    api_key = require_env(OPENAI_API_KEY, "OPENAI_API_KEY")

    # Safeguard: lege body -> thread conservatief op Waiting On Reply zetten
    if not body_text or len(body_text.strip()) < 10:
        return {
            "label": "Waiting On Reply",
            "reason": "Sent body too short or empty to classify",
            "generate_draft": False,
        }

    body_for_prompt = (body_text or "")[:3000]

    prompt = f"""
Je bent een e-mail classifier voor OfficeFlow.

Context:
- Dit is het LAATSTE ECHT VERZONDEN bericht dat de gebruiker zelf heeft verstuurd in een thread.
- Jij moet bepalen of de thread na dit verzonden bericht moet staan op Waiting On Reply, Follow Up of Done.
- Wees CONSERVATIEF met Done.
- Als er nog enige openheid, onzekerheid, vervolgstap of impliciete opvolging in de tekst zit, kies NIET Done.

Kies exact 1 label uit deze lijst:
- Waiting On Reply
- Follow Up
- Done

Strikte definities:
- Waiting On Reply:
  de gebruiker wacht op een reactie, bevestiging, antwoord of actie van de andere partij.
  Er staat een open vraag, verzoek, check of reactie-vraag richting de ander.

- Follow Up:
  de gebruiker hoeft nu niet direct een antwoord van de ander te krijgen,
  maar heeft zelf nog een vervolgstap, update, check, terugkoppeling of opvolging openstaan.

- Done:
  de thread is inhoudelijk echt afgerond.
  Er is geen antwoord van de andere partij meer nodig
  EN de gebruiker heeft ook zelf geen vervolgstap meer openstaan.

ZEER BELANGRIJKE BESLISREGELS:
- Kies Done ALLEEN als het gesprek duidelijk volledig afgerond is.
- Als de gebruiker zegt of impliceert dat hij later nog iets laat weten, terugkoppelt, uitzoekt, checkt, opvolgt of op de hoogte houdt, kies Follow Up.
- Als de gebruiker wacht op antwoord, bevestiging, akkoord, reactie of informatie van de ander, kies Waiting On Reply.
- Bij twijfel NOOIT Done kiezen.
- Bij twijfel tussen Done en Follow Up -> kies Follow Up.
- Bij twijfel tussen Done en Waiting On Reply -> kies Waiting On Reply.
- Bij twijfel tussen Waiting On Reply en Follow Up:
  - kies Waiting On Reply als de ander nu aan zet is
  - kies Follow Up als de gebruiker zelf nog aan zet is

Voorbeelden die bijna altijd Follow Up zijn:
- "Ik kom hier later op terug"
- "Ik laat je dit nog weten"
- "Ik check het en kom erop terug"
- "Ik neem contact op en laat het je weten"
- "We houden je op de hoogte"
- "Ik stuur later nog een update"
- "Ik zoek het uit"
- "Ik pak dit op"

Voorbeelden die bijna altijd Waiting On Reply zijn:
- "Laat je weten of dit lukt?"
- "Kun je dit bevestigen?"
- "Ik hoor graag van je"
- "Kun je aangeven wanneer..."
- "Graag ontvang ik je reactie"

Voorbeelden die bijna altijd Done zijn:
- "Top, dank"
- "Helemaal goed"
- "Dank, hiermee kan ik verder"
- "Bedankt, voor nu is alles duidelijk"
- "Het is afgehandeld"
- "Niet meer nodig"

Extra nuance:
- Informatieve updates zijn NIET automatisch Done.
- Een statusupdate zoals "de bestelling is onderweg" is alleen Done als er géén open vervolg meer is.
- Zinnen zoals "we houden je op de hoogte" of "ik laat het weten" betekenen expliciet dat de thread NIET klaar is -> Follow Up.
- Een nette afsluiting of beleefdheid betekent niet automatisch Done.

Geef alleen geldige JSON terug in exact dit formaat:
{{
  "label": "Follow Up",
  "reason": "Korte reden"
}}

Onderwerp: {subject}

Laatste verzonden bericht:
{body_for_prompt}
""".strip()

    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            response = await client.post(
                "https://api.openai.com/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json",
                },
                json={
                    "model": "gpt-4o-mini",
                    "messages": [
                        {
                            "role": "system",
                            "content": (
                                "Je bepaalt de juiste threadstatus na een echt verzonden reply. "
                                "Je bent streng met Done. "
                                "Je kiest alleen Done als de thread echt volledig afgerond is. "
                                "Je geeft alleen geldige JSON terug."
                            ),
                        },
                        {
                            "role": "user",
                            "content": prompt,
                        },
                    ],
                    "temperature": 0,
                },
            )
    except Exception as exc:
        print(f"[classify_sent_reply] OpenAI call failed: {repr(exc)} -- fallback 'Waiting On Reply'")
        return {
            "label": "Waiting On Reply",
            "reason": f"Classifier unavailable: {type(exc).__name__}",
            "generate_draft": False,
        }

    data = parse_response_data(response)
    if response.status_code >= 400:
        print(f"[classify_sent_reply] OpenAI {response.status_code}: {data} -- fallback 'Waiting On Reply'")
        return {
            "label": "Waiting On Reply",
            "reason": f"Classifier HTTP {response.status_code}",
            "generate_draft": False,
        }

    try:
        content = data["choices"][0]["message"]["content"]
        parsed = safe_parse_json(content)
    except Exception:
        print(f"[classify_sent_reply] Invalid response shape: {data} -- fallback 'Waiting On Reply'")
        return {
            "label": "Waiting On Reply",
            "reason": "Classifier returned invalid JSON",
            "generate_draft": False,
        }

    label = parsed.get("label")
    reason = parsed.get("reason") or "No reason given"

    if label not in SENT_REPLY_STATUS_LABELS:
        print(f"[classify_sent_reply] Unknown label '{label}' -- fallback 'Waiting On Reply'")
        return {
            "label": "Waiting On Reply",
            "reason": f"Classifier returned unknown label '{label}'",
            "generate_draft": False,
        }

    return {
        "label": label,
        "reason": reason,
        "generate_draft": False,
    }


async def get_recent_sent_reply_samples(user_id: str, limit: int = 30) -> list[str]:
    gmail_data = await gmail_get_json_for_user(
        user_id=user_id,
        url=f"{GMAIL_API_BASE}/messages",
        params={
            "maxResults": min(max(limit, 1), 50),
            "labelIds": "SENT",
        },
    )

    messages = gmail_data.get("messages", [])
    samples: list[str] = []

    for message in messages:
        message_id = message.get("id")
        if not message_id:
            continue

        message_data = await gmail_get_json_for_user(
            user_id=user_id,
            url=f"{GMAIL_API_BASE}/messages/{message_id}",
        )

        label_ids = set(message_data.get("labelIds", []))
        if is_draft_label_ids(label_ids):
            continue

        payload = message_data.get("payload", {})
        body_text = extract_plain_text_from_payload(payload)
        cleaned = clean_reply_training_text(body_text)

        if not cleaned or len(cleaned) < 20:
            continue

        lowered = cleaned.lower().strip()
        if lowered in {"thanks", "thank you", "top", "prima", "ok", "oke"}:
            continue

        samples.append(cleaned)

    return samples


async def train_style_profile_from_sent_messages(
    user_id: str,
    source_limit: int = 30,
) -> dict[str, Any]:
    api_key = require_env(OPENAI_API_KEY, "OPENAI_API_KEY")

    samples = await get_recent_sent_reply_samples(user_id=user_id, limit=source_limit)

    if len(samples) < 3:
        raise HTTPException(status_code=400, detail="Not enough usable sent emails to build style profile")

    joined_samples = "\n\n---EMAIL SAMPLE---\n\n".join(samples[:30])

    prompt = f"""
Je analyseert echte eerder verzonden zakelijke e-mails van één gebruiker.

Doel:
- vat de schrijfstijl compact samen
- maak een bruikbaar stijlprofiel voor toekomstige draft replies
- beschrijf toon, lengte, directheid, formaliteit, typische afsluiting en opvallende stijlkenmerken
- verzin niets dat niet uit de voorbeelden blijkt
- gebruik duidelijke, praktische taal

Geef alleen geldige JSON terug in exact dit formaat:
{{
  "style_profile_text": "Compacte stijlomschrijving",
  "style_profile_json": {{
    "language": "nl",
    "tone": "friendly_professional",
    "length": "short",
    "formality": "neutral",
    "directness": "direct",
    "closing_style": "practical",
    "emoji_usage": "rare",
    "key_traits": ["trait 1", "trait 2"]
  }}
}}

VOORBEELDEN:
{joined_samples}
""".strip()

    async with httpx.AsyncClient(timeout=90.0) as client:
        response = await client.post(
            "https://api.openai.com/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            },
            json={
                "model": "gpt-4o-mini",
                "messages": [
                    {
                        "role": "system",
                        "content": "Je analyseert schrijfstijl van zakelijke e-mails en geeft alleen geldige JSON terug.",
                    },
                    {
                        "role": "user",
                        "content": prompt,
                    },
                ],
                "temperature": 0,
            },
        )

    data = parse_response_data(response)
    if response.status_code >= 400:
        raise HTTPException(status_code=500, detail=f"OpenAI style training error: {data}")

    try:
        content = data["choices"][0]["message"]["content"]
        parsed = safe_parse_json(content)
    except Exception:
        raise HTTPException(status_code=500, detail=f"Invalid style training response: {data}")

    style_profile_text = parsed.get("style_profile_text")
    style_profile_json = parsed.get("style_profile_json") or {}

    if not style_profile_text:
        raise HTTPException(status_code=500, detail="Style training returned no style_profile_text")

    saved = await supabase_upsert_user_style_profile(
        user_id=user_id,
        source_sent_count=len(samples),
        style_profile_text=style_profile_text,
        style_profile_json=style_profile_json,
    )

    return {
        "saved_profile": saved,
        "source_sent_count": len(samples),
    }


def _format_attachment_context(summaries: list[dict[str, Any]] | None) -> str:
    """Build a prompt-ready block from attachment summaries. Returns empty
    string if there's nothing useful to add (no summaries, all failed/unsupported)."""
    if not summaries:
        return ""
    blocks: list[str] = []
    for s in summaries:
        if (s.get("status") or "") != "success":
            continue
        summary_text = (s.get("summary") or "").strip()
        if not summary_text:
            continue
        filename = s.get("filename") or "Bijlage"
        kd = s.get("key_data") or {}

        lines = [f"{filename}:", f"- Samenvatting: {summary_text}"]
        for label, key in [
            ("Bedragen", "amounts"),
            ("Datums", "dates"),
            ("Deadlines", "deadlines"),
            ("Partijen", "parties"),
            ("Belangrijke clausules", "key_clauses"),
        ]:
            values = kd.get(key) or []
            if isinstance(values, list) and values:
                joined = "; ".join(str(v) for v in values if v)
                if joined:
                    lines.append(f"- {label}: {joined}")
        blocks.append("\n".join(lines))

    if not blocks:
        return ""

    return (
        "Bijlage-context (door OfficeFlow geanalyseerd uit de bijgevoegde bestanden):\n\n"
        + "\n\n".join(blocks)
        + "\n\nGebruik deze concrete feiten in je reply waar relevant. "
        "Verwijs naar specifieke bedragen, data en voorwaarden in plaats van een "
        "generieke \"ik bekijk het en kom terug\" — tenzij de afzender niets vraagt "
        "wat met deze feiten te maken heeft."
    )


async def generate_ai_reply(
    user_id: str,
    subject: str | None,
    sender: str | None,
    body_text: str | None,
    attachment_summaries: list[dict[str, Any]] | None = None,
) -> str:
    api_key = require_env(OPENAI_API_KEY, "OPENAI_API_KEY")

    settings = await supabase_get_user_settings(user_id)
    style_profile = await supabase_get_user_style_profile(user_id)

    style_instructions = build_reply_style_instructions(settings)
    learned_style_instructions = build_style_profile_instructions(style_profile)
    language_instruction_block = build_language_instruction_block(settings, body_text)
    attachment_context_block = _format_attachment_context(attachment_summaries)

    prompt = f"""
Je bent de persoonlijke e-mailassistent van de gebruiker.

Jouw taak:
- schrijf ALLEEN de body van een reply
- schrijf alsof de gebruiker zelf antwoordt
- maak het bruikbaar als echte conceptmail in Gmail

Harde regels:
- schrijf GEEN onderwerpregel
- schrijf NOOIT "Onderwerp:" of "Subject:"
- gebruik GEEN placeholders zoals "[je naam]" of "[your name]"
- noem NOOIT "OfficeFlow", tenzij dit expliciet nodig is vanuit de e-mail of gebruikersinstructies
- verzin geen details, prijzen, data, deadlines of beloftes die niet in de mail, instructies of bijlage-context staan
- doe geen concrete toezeggingen over timing, planning, prijs of oplevering tenzij die expliciet bekend zijn (bijv. uit een bijlage)
- voeg NOOIT extra informatie toe die niet expliciet gevraagd wordt
- verwijs NOOIT naar websites, pagina's of externe informatie tenzij dat expliciet gevraagd wordt in de e-mail of expliciet is ingesteld door de gebruiker
- schrijf natuurlijk, menselijk, geloofwaardig en direct
- als een korte bevestiging genoeg is, houd het kort
- als er meer context nodig is en de gebruiker heeft voorkeur voor langere antwoorden, mag je uitgebreider zijn zolang het relevant blijft
- als er nog iets moet volgen, formuleer dat concreet maar zonder iets te beloven dat niet vaststaat
- als er een vaste handtekening bekend is, schrijf zelf geen losse afsluitgroet; die wordt later toegevoegd
- als er geen handtekening bekend is, laat die weg

Taalregels:
{language_instruction_block}

Vermijd expliciet dit soort formuleringen:
- "Bedankt voor je vraag"
- "Dank voor je bericht" tenzij het echt natuurlijk voelt
- "Ik hoop dat het goed met je gaat"
- "Graag wil ik je informeren"
- "Ik kom hier zo snel mogelijk op terug"
- "Neem gerust contact op"
- "Laat het gerust weten"
- "Mocht je nog vragen hebben"
- "Hierbij"

Voorkeursstijl:
- direct
- zakelijk menselijk
- geen overdreven beleefdheidsvulling
- geen marketingtaal
- geen opvulzinnen

Specifieke instructie voor inhoud:
- als iemand om een offerte, prijs of indicatie vraagt, geef een korte of passende directe terugkoppeling zonder overbodige woorden
- doe NOOIT concrete tijdsbeloftes zoals "vandaag", "straks", "vanmiddag", "binnen een uur" of andere deadlines tenzij die expliciet vaststaan in de input of gebruikersinstructies
- verzin nooit prijzen, timings, oplevertermijnen of toezeggingen
- als de gebruiker voorkeur voor korte antwoorden heeft, houd het compact
- als de gebruiker voorkeur voor langere antwoorden heeft, mag je vollediger antwoorden, maar blijf relevant en to the point
- goede stijl is bijvoorbeeld: "Ik kom bij je terug met een prijsindicatie en een inschatting van de doorlooptijd."
- dus liever 1 sterk relevant antwoord dan extra zinnen die niets toevoegen

Gebruikersvoorkeuren:
{style_instructions if style_instructions else "Geen extra voorkeuren ingesteld."}

Geleerde schrijfstijl:
{learned_style_instructions if learned_style_instructions else "Nog geen stijlprofiel beschikbaar."}

{attachment_context_block if attachment_context_block else ""}

Belangrijk:
- expliciete gebruikersinstellingen gaan boven het geleerde stijlprofiel
- output moet direct bruikbaar zijn als draft body
- schrijf alleen de uiteindelijke tekst, zonder uitleg of toelichting
- bij bijlage-context: gebruik concrete feiten (bedragen, datums) waar de afzender ernaar vraagt of daarmee samenhangt

Van: {sender}
Onderwerp: {subject}

Originele e-mail:
{body_text}
""".strip()

    system_message = (
        "Je schrijft natuurlijke zakelijke replies namens de gebruiker. "
        "Je verzint geen details. "
        "Je doet geen ongegronde beloftes over tijd, prijs, planning of oplevering. "
        "Je voegt geen irrelevante extra informatie toe. "
        "Je verwijst niet naar websites tenzij dat expliciet gevraagd of ingesteld is. "
        "Je output bevat alleen de uiteindelijke mailtekst."
    )

    async with httpx.AsyncClient(timeout=60.0) as client:
        response = await client.post(
            "https://api.openai.com/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            },
            json={
                "model": "gpt-4o-mini",
                "messages": [
                    {
                        "role": "system",
                        "content": system_message,
                    },
                    {
                        "role": "user",
                        "content": prompt,
                    },
                ],
                "temperature": 0.2,
            },
        )

    data = parse_response_data(response)
    if response.status_code >= 400:
        raise HTTPException(status_code=500, detail=f"OpenAI error: {data}")
    if not isinstance(data, dict) or "choices" not in data:
        raise HTTPException(status_code=500, detail=f"OpenAI error: {data}")

    raw_reply = data["choices"][0]["message"]["content"]
    cleaned_reply = sanitize_generated_reply(raw_reply)
    final_reply = maybe_apply_signature(cleaned_reply, settings)

    if not final_reply:
        raise HTTPException(status_code=500, detail="Generated reply was empty after cleanup")

    return final_reply


# ----------------------------
# Processing engine
# ----------------------------

async def process_open_to_respond_threads(
    user_id: str,
    mailbox_id: str,
    mailbox_email: str | None,
    label_name_to_id: dict[str, str],
    max_results: int = 25,
) -> list[dict[str, Any]]:
    label_ids_to_scan = [
        label_name_to_id[name]
        for name in ("To Respond", "Priority")
        if label_name_to_id.get(name)
    ]
    if not label_ids_to_scan:
        return []

    seen_message_ids: set[str] = set()
    messages: list[dict[str, Any]] = []
    for label_id in label_ids_to_scan:
        gmail_data = await gmail_get_json_for_user(
            user_id=user_id,
            url=f"{GMAIL_API_BASE}/messages",
            params={
                "maxResults": max_results,
                "labelIds": label_id,
            },
        )
        for message in gmail_data.get("messages", []):
            message_id = message.get("id")
            if not message_id or message_id in seen_message_ids:
                continue
            seen_message_ids.add(message_id)
            messages.append(message)

    processed_thread_ids: set[str] = set()
    results: list[dict[str, Any]] = []

    for message in messages:
        message_id = message.get("id")
        if not message_id:
            continue

        try:
            message_data = await gmail_get_json_for_user(
                user_id=user_id,
                url=f"{GMAIL_API_BASE}/messages/{message_id}",
            )

            thread_id = message_data.get("threadId")
            if not thread_id or thread_id in processed_thread_ids:
                continue

            processed_thread_ids.add(thread_id)

            current_label_ids = set(message_data.get("labelIds", []))
            if not any(lid in current_label_ids for lid in label_ids_to_scan):
                continue

            thread_reply_state = await get_thread_reply_state(
                user_id=user_id,
                thread_id=thread_id,
                mailbox_email=mailbox_email,
            )

            latest_user_message = thread_reply_state.get("latest_user_message")
            has_open_draft = thread_reply_state.get("has_open_draft", False)

            if not latest_user_message or has_open_draft:
                continue

            if not thread_reply_state["user_is_latest_sender"]:
                continue

            latest_user_payload = latest_user_message.get("payload", {})
            latest_user_headers = latest_user_payload.get("headers", [])
            latest_user_subject = get_header_value(latest_user_headers, "Subject")
            latest_user_body_text = extract_plain_text_from_payload(latest_user_payload)

            sent_status = await classify_latest_sent_reply_status(
                subject=latest_user_subject,
                body_text=latest_user_body_text,
            )
            target_label = sent_status["label"]

            updated_label_ids = await sync_thread_status(
                user_id=user_id,
                thread_id=thread_id,
                current_message_id=message_id,
                current_label_ids=current_label_ids,
                label_name_to_id=label_name_to_id,
                target_label_name=target_label,
            )

            # --- FOLLOW-UP RADAR: detect commitments in this just-sent mail ---
            commitment_info = {"detected": False, "forced_follow_up": False}
            try:
                to_header = get_header_value(latest_user_headers, "To") or ""
                recipient_email = None
                recipient_name = None
                if to_header:
                    import re as _re_cm
                    m = _re_cm.match(r'^\s*(?:"?([^"<]+?)"?\s*)?<?([^<>\s]+@[^<>\s]+)>?\s*$', to_header)
                    if m:
                        recipient_name = (m.group(1) or "").strip() or None
                        recipient_email = (m.group(2) or "").strip().lower() or None

                sent_at_iso = None
                try:
                    internal_ts = int(latest_user_message.get("internalDate", "0"))
                    if internal_ts:
                        sent_at_iso = datetime.fromtimestamp(internal_ts / 1000, tz=timezone.utc).isoformat()
                except Exception:
                    pass

                commitment_info = await process_sent_mail_for_commitment(
                    user_id=user_id,
                    mailbox_id=mailbox_id,
                    thread_id=thread_id,
                    message_id=message_id,
                    subject=latest_user_subject,
                    body_text=latest_user_body_text,
                    recipient_email=recipient_email,
                    recipient_name=recipient_name,
                    sent_at_iso=sent_at_iso,
                    label_name_to_id=label_name_to_id,
                    current_label_ids=set(updated_label_ids),
                )
                if commitment_info.get("forced_follow_up"):
                    target_label = "Follow Up"
            except Exception as exc:
                print(f"[commitment-hook] Failed for thread {thread_id}: {repr(exc)}")

            # --- SILENCE RADAR: track / cancel "Waiting On Reply" state ---
            try:
                if target_label == "Waiting On Reply":
                    to_header_s = get_header_value(latest_user_headers, "To") or ""
                    recipient_email_s = None
                    recipient_name_s = None
                    if to_header_s:
                        import re as _re_sr
                        ms = _re_sr.match(
                            r'^\s*(?:"?([^"<]+?)"?\s*)?<?([^<>\s]+@[^<>\s]+)>?\s*$',
                            to_header_s,
                        )
                        if ms:
                            recipient_name_s = (ms.group(1) or "").strip() or None
                            recipient_email_s = (ms.group(2) or "").strip().lower() or None

                    sent_at_iso_s = None
                    try:
                        internal_ts_s = int(latest_user_message.get("internalDate", "0"))
                        if internal_ts_s:
                            sent_at_iso_s = datetime.fromtimestamp(
                                internal_ts_s / 1000, tz=timezone.utc
                            ).isoformat()
                    except Exception:
                        pass

                    message_id_header_s = get_header_value(latest_user_headers, "Message-ID")

                    if recipient_email_s:
                        await supabase_upsert_awaiting_reply(
                            user_id=user_id,
                            mailbox_id=mailbox_id,
                            gmail_thread_id=thread_id,
                            last_user_message_id=message_id_header_s or message_id,
                            last_user_sent_at=sent_at_iso_s or _dt.now(_tz.utc).isoformat(),
                            recipient_email=recipient_email_s,
                            recipient_name=recipient_name_s,
                            subject=latest_user_subject,
                        )
                else:
                    # Thread moved away from Waiting On Reply — cancel any active
                    # awaiting_reply so we don't nudge later.
                    await supabase_cancel_awaiting_reply_for_thread(thread_id)
            except Exception as exc:
                print(f"[silence-hook] Failed for thread {thread_id}: {repr(exc)}")

            results.append(
                {
                    "gmail_message_id": message_id,
                    "gmail_thread_id": thread_id,
                    "label": target_label,
                    "classification_reason": sent_status["reason"],
                    "draft_created": False,
                    "gmail_draft_id": None,
                    "label_ids": list(updated_label_ids),
                    "thread_status": "updated_after_sent_reply",
                    "commitment_detected": commitment_info.get("detected", False),
                }
            )
        except Exception as exc:
            print(f"[open-to-respond] Thread {message_id} failed: {repr(exc)}")
            continue

    return results


async def process_inbox_for_user(
    email: str,
    max_results: int = 10,
    initial_run: bool = False,
) -> dict[str, Any]:
    context = await get_gmail_context_by_email(email)
    user = context["user"]
    mailbox = context["mailbox"]
    mailbox_email = mailbox.get("email_address")

    if max_results < 1:
        max_results = 1
    if max_results > 20:
        max_results = 20

    await setup_gmail_labels_for_mailbox(
        user_id=user["id"],
        mailbox_id=mailbox["id"],
    )

    # Ensure this mailbox has an inbox_cutoff_at — first time only. After
    # this, the classifier ignores everything with internalDate < cutoff.
    # Mutates `mailbox` in place so cutoff_epoch below sees the stamp.
    await ensure_mailbox_inbox_cutoff(mailbox)

    label_map = await get_all_gmail_labels(user["id"])
    label_name_to_id = {name: label["id"] for name, label in label_map.items()}
    custom_label_ids = get_status_label_ids_from_map(label_name_to_id)

    results: list[dict[str, Any]] = []

    # Build the Gmail search query.
    # - Normal runs (cron): UNREAD inbox mail AFTER cutoff — ignores old
    #   marketing/notifications already lying around when the user connected.
    # - Initial run (onboarding): the latest N inbox mails regardless of
    #   read-status / cutoff. Gives every new user an immediate "wow" moment
    #   with labels and drafts on their actual recent inbox, instead of an
    #   empty dashboard while they wait for new mail to arrive.
    if initial_run:
        # Initial onboarding run: skip the cutoff filter entirely so brand-new
        # users see labels + drafts immediately on the latest inbox mails.
        cutoff_epoch = None
        gmail_query = "in:inbox"
    else:
        cutoff_epoch = _mailbox_cutoff_epoch(mailbox)
        base_query = "in:inbox is:unread"
        gmail_query = f"{base_query} after:{cutoff_epoch}" if cutoff_epoch else base_query

    gmail_data = await gmail_get_json_for_user(
        user_id=user["id"],
        url=f"{GMAIL_API_BASE}/messages",
        params={
            "maxResults": max_results,
            "q": gmail_query,
        },
    )

    messages = gmail_data.get("messages", [])

    for message in messages:
        message_id = message.get("id")
        if not message_id:
            continue

        message_data = await gmail_get_json_for_user(
            user_id=user["id"],
            url=f"{GMAIL_API_BASE}/messages/{message_id}",
        )

        # Defensive cutoff check — Gmail's `after:` operator is usually
        # reliable, but we belt-and-brace with internalDate (ms since epoch)
        # so historic mail can never slip through even if the query is off.
        if cutoff_epoch is not None:
            try:
                internal_ms = int(message_data.get("internalDate") or 0)
                if internal_ms and internal_ms // 1000 < cutoff_epoch:
                    # Skip silently — this mail predates the OfficeFlow connect.
                    continue
            except Exception:
                pass

        payload = message_data.get("payload", {})
        headers = payload.get("headers", [])

        from_header = get_header_value(headers, "From")
        subject = get_header_value(headers, "Subject")
        body_text = extract_plain_text_from_payload(payload)
        original_message_id_header = get_header_value(headers, "Message-ID")
        references_header = get_header_value(headers, "References")
        thread_id = message_data.get("threadId")
        current_label_ids = set(message_data.get("labelIds", []))

        has_any_custom_label = any(label_id in current_label_ids for label_id in custom_label_ids)

        email_row = await supabase_insert_email(
            user_id=user["id"],
            mailbox_id=mailbox["id"],
            gmail_message_id=message_id,
            gmail_thread_id=thread_id,
            subject=subject,
        )

        existing_drafts = await supabase_get_drafts_by_email_id(email_row["id"]) if email_row else []

        # Als deze unread mail al eerder verwerkt is, sla hem over
        if has_any_custom_label or existing_drafts:
            results.append(
                {
                    "gmail_message_id": message_data.get("id"),
                    "gmail_thread_id": thread_id,
                    "subject": subject,
                    "from_name": from_header,
                    "from_email": extract_email_address(from_header),
                    "snippet": message_data.get("snippet"),
                    "body_text": body_text[:500] if body_text else None,
                    "label_ids": list(current_label_ids),
                    "label": None,
                    "generate_draft": False,
                    "classification_reason": "Skipped: already processed",
                    "draft_created": False,
                    "gmail_draft_id": None,
                    "already_had_label": has_any_custom_label,
                    "thread_status": "skipped_already_processed",
                    "has_open_draft": bool(existing_drafts),
                }
            )
            continue

        thread_reply_state = await get_thread_reply_state(
            user_id=user["id"],
            thread_id=thread_id,
            mailbox_email=mailbox_email,
        )

        latest_incoming_message = thread_reply_state.get("latest_incoming_message")
        has_open_draft = thread_reply_state.get("has_open_draft", False)

        latest_incoming_payload = latest_incoming_message.get("payload", {}) if latest_incoming_message else {}
        latest_incoming_headers = latest_incoming_payload.get("headers", []) if latest_incoming_message else []
        latest_incoming_subject = get_header_value(latest_incoming_headers, "Subject")
        latest_incoming_from = get_header_value(latest_incoming_headers, "From")
        latest_incoming_body_text = extract_plain_text_from_payload(latest_incoming_payload) if latest_incoming_message else None

        is_marketing_tab = "CATEGORY_PROMOTIONS" in current_label_ids
        is_social_tab = "CATEGORY_SOCIAL" in current_label_ids
        is_updates_tab = "CATEGORY_UPDATES" in current_label_ids

        # Alleen follow-up classificatie als er echt eerder een user reply in de thread zat
        if thread_reply_state["needs_response_after_reply"] and latest_incoming_message is not None:
            classification = await classify_follow_up_email(
                subject=latest_incoming_subject or subject,
                sender=latest_incoming_from or from_header,
                body_text=latest_incoming_body_text or body_text,
                mailbox_email=mailbox_email,
            )
            target_label = classification["label"]
            classification_reason = classification["reason"]
            classification_confidence = normalize_confidence(classification.get("confidence"))
            thread_status = "reopened"
        else:
            classification = await classify_email(
                subject=subject,
                sender=from_header,
                body_text=body_text,
                mailbox_email=mailbox_email,
            )
            target_label = classification["label"]
            classification_reason = classification["reason"]
            classification_confidence = normalize_confidence(classification.get("confidence"))
            thread_status = "classified"

        # --- FOLLOW-UP RADAR: close any active commitment on this thread
        # now that an external reply has arrived ---
        try:
            if thread_reply_state.get("needs_response_after_reply") and latest_incoming_message is not None:
                await mark_commitments_completed_on_reply(thread_id)
        except Exception as exc:
            print(f"[radar-reply-hook] Failed for thread {thread_id}: {repr(exc)}")

        # --- SILENCE RADAR: close any active awaiting_reply now that reply arrived ---
        try:
            if thread_reply_state.get("needs_response_after_reply") and latest_incoming_message is not None:
                await mark_awaiting_replies_replied_on_reply(thread_id)
        except Exception as exc:
            print(f"[silence-reply-hook] Failed for thread {thread_id}: {repr(exc)}")

        current_label_ids = await sync_thread_status(
            user_id=user["id"],
            thread_id=thread_id,
            current_message_id=message_id,
            current_label_ids=current_label_ids,
            label_name_to_id=label_name_to_id,
            target_label_name=target_label,
        )

        # --- AUTO-ARCHIVE LOW-VALUE MAIL ---
        # Safe set (Marketing / Ignore) is archived by default when
        # auto_archive_low_value is on. Notification is opt-in via
        # notification_auto_archive because notifications are often transactional.
        # Guarded by the trust-list: bekende contacten worden nooit gearchiveerd.
        try:
            auto_archive_on = bool(mailbox.get("auto_archive_low_value", True))
            in_safe_set = target_label in LOW_VALUE_LABELS
            in_notification_optin = (
                target_label == NOTIFICATION_AUTO_ARCHIVE_LABEL
                and bool(mailbox.get("notification_auto_archive", False))
            )
            if auto_archive_on and (in_safe_set or in_notification_optin):
                sender_email = extract_email_address(from_header) or ""
                trusted = await is_trusted_sender(
                    user_id=user["id"],
                    sender_email=sender_email,
                )
                if trusted:
                    classification_reason = (
                        (classification_reason or "").rstrip(" ·")
                        + " · Trusted sender — kept in Inbox"
                    ).lstrip(" ·").strip()
                else:
                    await archive_low_value_thread(
                        user_id=user["id"],
                        thread_id=thread_id,
                        current_message_id=message_id,
                    )
                    current_label_ids.discard("INBOX")
                    classification_reason = (
                        (classification_reason or "").rstrip(" ·")
                        + " · Auto-archived (low-value label)"
                    ).lstrip(" ·").strip()
        except Exception as exc:
            print(f"[auto-archive-hook] thread {thread_id}: {repr(exc)}")

        # ============ ATTACHMENT AI SUMMARIES ============
        # Process attachments BEFORE draft generation so the AI reply can
        # reference concrete facts (amounts, dates, deadlines) from PDF/DOCX.
        # Best-effort: failure here must not block the rest of the flow.
        # Only runs for To Respond / Priority / Follow Up labels (cost guard
        # is inside process_message_attachments).
        attachments_processed = 0
        attachment_summaries_for_draft: list[dict[str, Any]] = []
        try:
            attachments_processed = await process_message_attachments(
                user_id=user["id"],
                email_id=email_row["id"] if email_row else None,
                gmail_message_id=message_id,
                message_data=message_data,
                target_label=target_label,
            )
            if attachments_processed > 0:
                attachment_summaries_for_draft = (
                    await supabase_get_attachment_summaries_for_message(
                        user_id=user["id"],
                        gmail_message_id=message_id,
                    )
                )
        except Exception as exc:
            print(f"[attachment-ai] outer error for message {message_id}: {repr(exc)}")

        # Hard guard: never draft a reply to a noreply / donotreply / system
        # sender. You literally cannot reply to them and any AI draft would
        # echo their notice back at /dev/null. Common case: Google Ads
        # ads-account-noreply@google.com, Stripe receipts, etc.
        sender_is_noreply = is_noreply_sender(extract_email_address(from_header))
        if sender_is_noreply:
            print(
                f"[draft-skip] Skipping draft for message {message_id}: "
                f"sender '{from_header}' is a noreply / system address."
            )

        should_generate_draft = (
            LABEL_RULES[target_label]["generate_draft"]
            and not is_marketing_tab
            and not is_social_tab
            and not is_updates_tab
            and not has_open_draft
            and not sender_is_noreply
        )

        draft_created = False
        gmail_draft_id = None
        draft_error: str | None = None

        if should_generate_draft and email_row and not existing_drafts:
            try:
                ai_reply = await generate_ai_reply(
                    user_id=user["id"],
                    subject=subject,
                    sender=from_header,
                    body_text=body_text,
                    attachment_summaries=attachment_summaries_for_draft or None,
                )

                to_email = extract_email_address(from_header)

                if to_email and ai_reply:
                    draft_result = await create_gmail_threaded_draft(
                        user_id=user["id"],
                        to_email=to_email,
                        subject=subject,
                        body=ai_reply,
                        thread_id=thread_id,
                        original_message_id_header=original_message_id_header,
                        references_header=references_header,
                    )

                    gmail_draft_id = draft_result.get("id")

                    if gmail_draft_id:
                        await supabase_insert_draft(
                            user_id=user["id"],
                            email_id=email_row["id"],
                            gmail_draft_id=gmail_draft_id,
                            subject=subject,
                            draft_body=ai_reply,
                            status="generated",
                            confidence=classification_confidence,
                        )
                        draft_created = True

                        await supabase_upsert_onboarding_state(
                            user_id=user["id"],
                            gmail_connected=True,
                            profile_completed=True,
                            initial_sync_completed=False,
                            first_draft_generated=True,
                        )
            except Exception as exc:
                draft_error = f"{type(exc).__name__}: {str(exc)[:200]}"
                print(f"[draft-generation] Failed for message {message_id}: {repr(exc)}")
                # Mail blijft gelabeld. Volgende auto-process run probeert opnieuw
                # omdat existing_drafts nog steeds leeg is.

        results.append(
            {
                "gmail_message_id": message_data.get("id"),
                "gmail_thread_id": message_data.get("threadId"),
                "subject": subject,
                "from_name": from_header,
                "from_email": extract_email_address(from_header),
                "snippet": message_data.get("snippet"),
                "body_text": body_text[:500] if body_text else None,
                "label_ids": list(current_label_ids),
                "label": target_label,
                "confidence": classification_confidence,
                "generate_draft": should_generate_draft,
                "classification_reason": classification_reason,
                "draft_created": draft_created,
                "gmail_draft_id": gmail_draft_id,
                "draft_error": draft_error,
                "already_had_label": has_any_custom_label,
                "thread_status": thread_status,
                "has_open_draft": has_open_draft,
                "attachments_processed": attachments_processed,
            }
        )

    # 2) Alleen open To Respond threads opnieuw beoordelen als gebruiker echt gereageerd heeft
    updated_open_threads = await process_open_to_respond_threads(
        user_id=user["id"],
        mailbox_id=mailbox["id"],
        mailbox_email=mailbox_email,
        label_name_to_id=label_name_to_id,
        max_results=max(max_results, 20),
    )
    results.extend(updated_open_threads)

    return {
        "status": "ok",
        "count": len(results),
        "messages": results,
    }


async def auto_process_loop():
    await asyncio.sleep(8)

    sem = asyncio.Semaphore(AUTO_PROCESS_MAX_CONCURRENCY)

    async def _process_one(mailbox: dict[str, Any]) -> None:
        email = mailbox.get("email_address")
        if not email:
            return
        async with sem:
            try:
                await process_inbox_for_user(
                    email=email,
                    max_results=AUTO_PROCESS_MAX_RESULTS,
                )
                print(f"Processed mailbox: {email}")
            except Exception as exc:
                print(f"Failed mailbox {email}: {repr(exc)}")

    while True:
        try:
            if not AUTO_PROCESS_ENABLED:
                await asyncio.sleep(AUTO_PROCESS_INTERVAL_SECONDS)
                continue

            mailboxes = await get_all_active_mailboxes()
            print(
                f"Auto processing {len(mailboxes)} connected mailbox(es) "
                f"(concurrent, max {AUTO_PROCESS_MAX_CONCURRENCY} in flight)..."
            )

            if mailboxes:
                await asyncio.gather(*(_process_one(mb) for mb in mailboxes))

        except Exception as exc:
            print(f"Auto processor loop error: {repr(exc)}")

        await asyncio.sleep(AUTO_PROCESS_INTERVAL_SECONDS)


SCHEDULED_SEND_INTERVAL_SECONDS = int(os.getenv("SCHEDULED_SEND_INTERVAL_SECONDS", "30"))


async def scheduled_send_loop():
    """Background dispatcher for user-scheduled Gmail sends. Runs every
    SCHEDULED_SEND_INTERVAL_SECONDS, picks up rows where status='pending'
    AND scheduled_at <= now(), and sends each via Gmail's drafts/send API.

    Failure modes (all marked 'failed' with error_message — never retried):
      - draft was deleted in Gmail before scheduled time
      - mailbox token revoked (status moves to needs_reconnect upstream)
      - any other Gmail API error
    """
    await asyncio.sleep(12)

    while True:
        try:
            due_rows = await supabase_get_due_scheduled_sends(limit=50)
            if due_rows:
                print(f"[scheduled-send] Dispatching {len(due_rows)} due send(s)")

            for row in due_rows:
                row_id = row.get("id")
                user_id = row.get("user_id")
                gmail_draft_id = row.get("gmail_draft_id")
                if not (row_id and user_id and gmail_draft_id):
                    continue

                try:
                    await send_gmail_draft(
                        user_id=user_id,
                        gmail_draft_id=gmail_draft_id,
                    )
                    await supabase_update_scheduled_send(
                        scheduled_send_id=row_id,
                        status="sent",
                        sent_at=datetime.now(timezone.utc).isoformat(),
                    )
                    print(f"[scheduled-send] Sent draft {gmail_draft_id} for user {user_id}")
                except Exception as exc:
                    err = repr(exc)[:480]
                    try:
                        await supabase_update_scheduled_send(
                            scheduled_send_id=row_id,
                            status="failed",
                            error_message=err,
                        )
                    except Exception:
                        pass
                    print(f"[scheduled-send] FAILED draft {gmail_draft_id} (user {user_id}): {err}")

        except Exception as exc:
            print(f"[scheduled-send] loop error: {repr(exc)}")

        await asyncio.sleep(SCHEDULED_SEND_INTERVAL_SECONDS)


SNOOZE_WAKE_INTERVAL_SECONDS = int(os.getenv("SNOOZE_WAKE_INTERVAL_SECONDS", "60"))


async def snooze_wake_loop():
    """Background dispatcher for waking snoozed mails. Runs every
    SNOOZE_WAKE_INTERVAL_SECONDS, picks up rows where status='pending'
    AND wake_at <= now(), and restores each mail's INBOX + original status
    label via Gmail API.

    Failure modes (all marked 'failed' with error_message — never retried):
      - mailbox token revoked
      - mail was permanently deleted in Gmail
      - any other Gmail API error
    """
    await asyncio.sleep(15)

    while True:
        try:
            due_rows = await supabase_get_due_snoozed_emails(limit=50)
            if due_rows:
                print(f"[snooze-wake] Waking {len(due_rows)} due snooze(s)")

            for row in due_rows:
                row_id = row.get("id")
                user_id = row.get("user_id")
                gmail_message_id = row.get("gmail_message_id")
                gmail_thread_id = row.get("gmail_thread_id")
                original_label_id = row.get("original_label_id")
                if not (row_id and user_id and gmail_message_id):
                    continue

                try:
                    # Resolve Snoozed label id fresh — could have been
                    # deleted by the user between snooze and wake.
                    snoozed_label_id = None
                    try:
                        snoozed_label_id = await ensure_label_exists(
                            user_id=user_id,
                            label_name=SNOOZED_LABEL,
                        )
                    except Exception as exc:
                        print(f"[snooze-wake] could not resolve Snoozed label "
                              f"for user {user_id}: {repr(exc)}")

                    await wake_thread_in_gmail(
                        user_id=user_id,
                        gmail_message_id=gmail_message_id,
                        gmail_thread_id=gmail_thread_id,
                        original_label_id=original_label_id,
                        snoozed_label_id=snoozed_label_id,
                    )
                    await supabase_update_snoozed_email(
                        snooze_id=row_id,
                        status="woken",
                        woken_at=datetime.now(timezone.utc).isoformat(),
                    )
                    print(f"[snooze-wake] Woke mail {gmail_message_id} for user {user_id}")
                except Exception as exc:
                    err = repr(exc)[:480]
                    try:
                        await supabase_update_snoozed_email(
                            snooze_id=row_id,
                            status="failed",
                            error_message=err,
                        )
                    except Exception:
                        pass
                    print(f"[snooze-wake] FAILED mail {gmail_message_id} "
                          f"(user {user_id}): {err}")

        except Exception as exc:
            print(f"[snooze-wake] loop error: {repr(exc)}")

        await asyncio.sleep(SNOOZE_WAKE_INTERVAL_SECONDS)


@app.on_event("startup")
async def start_background_tasks():
    print("Multi-tenant auto processor started")
    asyncio.create_task(auto_process_loop())
    print("Scheduled-send dispatcher started")
    asyncio.create_task(scheduled_send_loop())
    print("Snooze wake dispatcher started")
    asyncio.create_task(snooze_wake_loop())


# ----------------------------
# Routes
# ----------------------------

@app.get("/")
def home():
    return {
        "name": "AI Mail Assistant API",
        "status": "online",
        "message": "API is running.",
    }


@app.get("/privacy")
def privacy():
    return {
        "title": "Privacy Policy",
        "content": (
            "We use Google account data only to authenticate users, read selected Gmail messages, "
            "generate AI-based email drafts, and create Gmail draft replies. "
            "We do not sell user data or share Gmail content with third parties except where required "
            "to provide the AI drafting service."
        ),
    }


@app.get("/terms")
def terms():
    return {
        "title": "Terms of Service",
        "content": (
            "This service helps generate AI email drafts. Users are responsible for reviewing all drafts "
            "before sending. We do not guarantee correctness, completeness, or suitability of generated content."
        ),
    }


@app.get("/health")
def health():
    return {"status": "ok"}


@app.get("/billing/status")
async def billing_status(email: str):
    user = await supabase_get_user_by_email(email)

    if not user:
        return JSONResponse(
            status_code=404,
            content={
                "email": email,
                "found": False,
                "subscription_status": None,
                "access_allowed": False,
            },
        )

    return {
        "email": user.get("email"),
        "found": True,
        "subscription_status": user.get("subscription_status"),
        "access_allowed": user.get("access_allowed", True),
        "stripe_customer_id": user.get("stripe_customer_id"),
        "stripe_subscription_id": user.get("stripe_subscription_id"),
    }


@app.post("/billing/portal")
async def create_billing_portal(email: str = Body(...)):
    user = await get_user_for_billing(email)

    try:
        session = stripe.billing_portal.Session.create(
            customer=user["stripe_customer_id"],
            return_url=STRIPE_BILLING_PORTAL_RETURN_URL,
        )
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to create billing portal session: {str(exc)}")

    return {
        "status": "ok",
        "url": session.url,
    }


@app.get("/auth/google/start")
def google_login():
    client_id = require_env(GOOGLE_CLIENT_ID, "GOOGLE_CLIENT_ID")
    redirect_uri = require_env(GOOGLE_REDIRECT_URI, "GOOGLE_REDIRECT_URI")

    # CSRF protection: signed state parameter prevents an attacker from
    # forging an OAuth callback that links someone else's Google account
    # to the wrong OfficeFlow user.
    state = make_oauth_state()

    url = (
        "https://accounts.google.com/o/oauth2/v2/auth"
        f"?client_id={client_id}"
        f"&redirect_uri={redirect_uri}"
        "&response_type=code"
        f"&scope={quote(GMAIL_SCOPE, safe=':/')}"
        "&access_type=offline"
        "&prompt=consent"
        f"&state={state}"
    )
    return RedirectResponse(url)


@app.get("/auth/google/callback")
async def google_callback(code: str, state: str | None = None):
    # CSRF check: state must be present and signed by us
    if not verify_oauth_state(state):
        print(f"[oauth-callback] Invalid OAuth state: {state!r}")
        raise HTTPException(
            status_code=400,
            detail="Invalid OAuth state — possible CSRF attempt. Restart the connection flow.",
        )

    client_id = require_env(GOOGLE_CLIENT_ID, "GOOGLE_CLIENT_ID")
    client_secret = require_env(GOOGLE_CLIENT_SECRET, "GOOGLE_CLIENT_SECRET")
    redirect_uri = require_env(GOOGLE_REDIRECT_URI, "GOOGLE_REDIRECT_URI")

    async with httpx.AsyncClient(timeout=60.0) as client:
        token_response = await client.post(
            "https://oauth2.googleapis.com/token",
            data={
                "code": code,
                "client_id": client_id,
                "client_secret": client_secret,
                "redirect_uri": redirect_uri,
                "grant_type": "authorization_code",
            },
        )
        token_data = parse_response_data(token_response)

        access_token = token_data.get("access_token") if isinstance(token_data, dict) else None
        refresh_token = token_data.get("refresh_token") if isinstance(token_data, dict) else None

        if not access_token:
            raise HTTPException(status_code=400, detail=f"Google token error: {token_data}")

        user_response = await client.get(
            "https://www.googleapis.com/oauth2/v2/userinfo",
            headers={"Authorization": f"Bearer {access_token}"},
        )
        user_info = parse_response_data(user_response)

    user_email = user_info.get("email") if isinstance(user_info, dict) else None
    user_name = user_info.get("name") if isinstance(user_info, dict) else None
    provider_account_id = user_info.get("id") if isinstance(user_info, dict) else None

    if not user_email:
        raise HTTPException(status_code=400, detail=f"Google userinfo error: {user_info}")

    try:
        existing_user = await ensure_user_has_access(user_email)
    except HTTPException as exc:
        if exc.status_code == 403:
            reason = "subscription_required"
            if exc.detail == "No user record found":
                reason = "no_subscription_record"

            return RedirectResponse(url=build_pricing_redirect(reason), status_code=302)
        raise

    user_id = existing_user["id"]

    if user_name and user_name != existing_user.get("full_name"):
        await supabase_update_user_profile(user_id=user_id, full_name=user_name)

    # --- Team seat-enforcement ---------------------------------
    # Als deze user lid is van een team, check of er nog een seat beschikbaar is.
    # Zo niet → redirect naar team-pagina met upgrade-hint.
    team_id_for_mailbox: str | None = None
    try:
        from app import teams as teams_module
        user_teams = await teams_module._resolve_user_teams(user_id)
        if user_teams:
            # Pak eerste team waar user lid van is met active access
            active_team = next(
                (t for t in user_teams if t["team"].get("access_allowed")),
                None,
            )
            if active_team:
                # Check of deze mailbox al gekoppeld is (herkoppeling mag altijd)
                already_connected = any(
                    mb.get("email_address", "").lower() == user_email.lower()
                    for mb in active_team["mailboxes"]
                )
                if not already_connected and active_team["seats_used"] >= active_team["seats_total"]:
                    # Seat-limit bereikt → redirect naar team-page
                    return RedirectResponse(
                        url="https://officeflowcompany.com/team.html?error=seat_limit",
                        status_code=302,
                    )
                team_id_for_mailbox = active_team["team"]["id"]
    except Exception as exc:
        print(f"[oauth-callback] team seat check failed: {repr(exc)}")
    # ------------------------------------------------------------

    await supabase_upsert_oauth_account(
        user_id=user_id,
        provider="google",
        provider_account_id=provider_account_id,
        access_token=access_token,
        refresh_token=refresh_token,
    )

    mailbox_payload: dict[str, Any] = {
        "user_id": user_id,
        "provider": "gmail",
        "email_address": user_email,
        "status": "pending_setup",
    }
    if team_id_for_mailbox:
        mailbox_payload["team_id"] = team_id_for_mailbox

    await supabase_upsert_mailbox(**mailbox_payload)

    await supabase_upsert_onboarding_state(
        user_id=user_id,
        gmail_connected=True,
        profile_completed=False,
        initial_sync_completed=False,
        first_draft_generated=False,
    )

    # Voor team-mailboxen: terug naar team-dashboard zodat admin/lid de
    # nieuwe mailbox direct in de sidebar ziet. Solo blijft via onboarding.
    if team_id_for_mailbox:
        return RedirectResponse(
            url="https://officeflowcompany.com/team.html?welcome=connected",
            status_code=302,
        )

    return RedirectResponse(
        url=f"https://officeflowcompany.com/onboarding/preferences.html?email={quote(user_email)}",
        status_code=302,
    )


@app.get("/test/protected")
async def test_protected(email: str):
    user = await ensure_user_has_access(email)
    mailbox = await supabase_get_mailbox_by_user_id(user_id=user["id"], provider="gmail")

    return {
        "status": "allowed",
        "user_id": user["id"],
        "mailbox_id": mailbox["id"] if mailbox else None,
    }


@app.get("/test/protected-ui")
async def test_protected_ui(email: str):
    user = await ensure_user_has_access(email)
    mailbox = await supabase_get_mailbox_by_user_id(user_id=user["id"], provider="gmail")

    return {
        "status": "allowed",
        "email": email,
        "user_id": user["id"],
        "mailbox_id": mailbox["id"] if mailbox else None,
        "message": "Protected route accessible",
    }


@app.post("/ai/reply")
async def ai_reply_route(
    email: str = Body(...),
    subject: str | None = Body(default=None),
    sender: str | None = Body(default=None),
    body_text: str | None = Body(default=None),
):
    user = await ensure_user_has_access(email)

    reply = await generate_ai_reply(
        user_id=user["id"],
        subject=subject,
        sender=sender,
        body_text=body_text,
    )

    return {"status": "ok", "reply": reply}


@app.post("/gmail/classify")
async def gmail_classify_route(
    email: str = Body(...),
    subject: str | None = Body(default=None),
    sender: str | None = Body(default=None),
    body_text: str | None = Body(default=None),
):
    await ensure_user_has_access(email)

    result = await classify_email(
        subject=subject,
        sender=sender,
        body_text=body_text,
    )

    return {"status": "ok", "classification": result}


@app.get("/gmail/inbox")
async def gmail_inbox(email: str, max_results: int = 10):
    return await process_inbox_for_user(email=email, max_results=max_results)


@app.post("/internal/process-inbox")
async def process_inbox_route(
    email: str = Body(...),
    max_results: int = Body(default=10),
    _: None = Depends(require_internal_secret),
):
    result = await process_inbox_for_user(email=email, max_results=max_results)

    return {
        "status": "ok",
        "processed_email": email,
        "count": result["count"],
        "messages": result["messages"],
    }


@app.post("/gmail/draft")
async def gmail_draft_route(
    email: str = Body(...),
    to_email: str = Body(...),
    subject: str | None = Body(default=None),
    body: str = Body(...),
    email_id: str | None = Body(default=None),
    gmail_thread_id: str | None = Body(default=None),
    original_message_id: str | None = Body(default=None),
    references: str | None = Body(default=None),
):
    context = await get_gmail_context_by_email(email)
    user = context["user"]

    draft_result = await create_gmail_threaded_draft(
        user_id=user["id"],
        to_email=to_email,
        subject=subject,
        body=body,
        thread_id=gmail_thread_id,
        original_message_id_header=original_message_id,
        references_header=references,
    )

    gmail_draft_id = draft_result.get("id")

    saved_draft = None
    if email_id and gmail_draft_id:
        saved_draft = await supabase_insert_draft(
            user_id=user["id"],
            email_id=email_id,
            gmail_draft_id=gmail_draft_id,
            subject=subject,
            draft_body=body,
            status="generated",
        )

    return {
        "status": "ok",
        "gmail_draft_id": gmail_draft_id,
        "draft": draft_result,
        "saved_draft": saved_draft,
    }


@app.post("/gmail/mark-done")
async def gmail_mark_done(
    email: str = Body(...),
    gmail_message_id: str = Body(...),
    archive: bool = Body(default=False),
):
    context = await get_gmail_context_by_email(email)
    user = context["user"]

    label_map = await get_all_gmail_labels(user["id"])
    label_name_to_id = {name: label["id"] for name, label in label_map.items()}

    message_data = await gmail_get_json_for_user(
        user_id=user["id"],
        url=f"{GMAIL_API_BASE}/messages/{gmail_message_id}",
    )
    current_label_ids = set(message_data.get("labelIds", []))
    thread_id = message_data.get("threadId")

    updated_label_ids = await sync_thread_status(
        user_id=user["id"],
        thread_id=thread_id,
        current_message_id=gmail_message_id,
        current_label_ids=current_label_ids,
        label_name_to_id=label_name_to_id,
        target_label_name="Done",
    )

    if archive and "INBOX" in updated_label_ids:
        await modify_gmail_message_labels(
            user_id=user["id"],
            gmail_message_id=gmail_message_id,
            remove_label_ids=["INBOX"],
        )
        updated_label_ids.discard("INBOX")

    return {
        "status": "ok",
        "gmail_message_id": gmail_message_id,
        "label": "Done",
        "label_ids": list(updated_label_ids),
        "archived": archive,
    }


@app.post("/internal/setup-labels")
async def setup_labels(
    email: str = Body(...),
    _: None = Depends(require_internal_secret),
):
    context = await get_gmail_context_by_email(email)
    user = context["user"]
    mailbox = context["mailbox"]

    result = await setup_gmail_labels_for_mailbox(
        user_id=user["id"],
        mailbox_id=mailbox["id"],
    )

    return {
        "status": "ok",
        "user_id": user["id"],
        "mailbox_id": mailbox["id"],
        "labels": result,
    }


@app.post("/internal/cleanup-legacy-labels")
async def cleanup_legacy_labels(
    email: str = Body(...),
    _: None = Depends(require_internal_secret),
):
    context = await get_gmail_context_by_email(email)
    user = context["user"]
    mailbox = context["mailbox"]

    result = await cleanup_legacy_labels_for_mailbox(
        user_id=user["id"],
        mailbox_id=mailbox["id"],
    )

    return {
        "status": "ok",
        "user_id": user["id"],
        "mailbox_id": mailbox["id"],
        "cleaned_labels": result,
    }


@app.post("/settings/prompt")
async def save_prompt_settings(payload: PromptSettingsPayload):
    user = await ensure_user_has_access(payload.email)
    clean_payload = build_clean_settings_payload(payload)

    saved = await supabase_upsert_user_settings(
        user_id=user["id"],
        payload=clean_payload,
    )

    return {
        "status": "ok",
        "email": payload.email,
        "settings": saved,
    }


# ---------------------------------------------------------------------------
# JWT-auth variants for the dashboard Voorkeuren-tab (no email in body)
# ---------------------------------------------------------------------------

@app.get("/api/preferences/me")
async def api_get_my_preferences(
    user: dict[str, Any] = Depends(get_current_user),
):
    """Return current user_settings for the logged-in user."""
    settings = await supabase_get_user_settings(user["id"])
    return {"status": "ok", "settings": settings or {}}


@app.get("/api/mailbox/status")
async def api_get_mailbox_status(
    user: dict[str, Any] = Depends(get_current_user),
):
    """Return mailbox connection status for the logged-in user. Used by the
    dashboard to show a "Reconnect Gmail" banner when the refresh token has
    been revoked/expired (status = 'needs_reconnect')."""
    mailbox = await supabase_get_mailbox_by_user_id(
        user_id=user["id"],
        provider="gmail",
    )
    if not mailbox:
        return {"status": "ok", "mailbox_status": "no_mailbox", "email": None}
    return {
        "status": "ok",
        "mailbox_status": mailbox.get("status") or "unknown",
        "email": mailbox.get("email_address"),
    }


# ============ SCHEDULED SENDS ============
class ScheduledSendCreate(BaseModel):
    gmail_draft_id: str
    scheduled_at: str  # ISO 8601 timestamp with timezone
    gmail_thread_id: str | None = None
    subject: str | None = None
    to_email: str | None = None


@app.post("/api/scheduled-sends")
async def api_create_scheduled_send(
    body: ScheduledSendCreate,
    user: dict[str, Any] = Depends(get_current_user),
):
    """Schedule an existing Gmail draft to be sent at a future time."""
    # Validate scheduled_at is parseable and in the future (with 30s grace).
    try:
        when = datetime.fromisoformat(body.scheduled_at.replace("Z", "+00:00"))
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid scheduled_at — expected ISO 8601")
    if when.tzinfo is None:
        when = when.replace(tzinfo=timezone.utc)
    now = datetime.now(timezone.utc)
    if when < now - timedelta(seconds=30):
        raise HTTPException(status_code=400, detail="scheduled_at must be in the future")

    mailbox = await supabase_get_mailbox_by_user_id(user_id=user["id"], provider="gmail")
    if not mailbox:
        raise HTTPException(status_code=400, detail="No connected Gmail mailbox")
    if mailbox.get("status") != "connected":
        raise HTTPException(
            status_code=400,
            detail=f"Mailbox not connected (status={mailbox.get('status')})",
        )

    row = await supabase_insert_scheduled_send(
        user_id=user["id"],
        mailbox_id=mailbox["id"],
        gmail_draft_id=body.gmail_draft_id,
        gmail_thread_id=body.gmail_thread_id,
        subject=body.subject,
        to_email=body.to_email,
        scheduled_at_iso=when.isoformat(),
    )
    return {"status": "ok", "scheduled_send": row}


@app.get("/api/scheduled-sends")
async def api_list_scheduled_sends(
    user: dict[str, Any] = Depends(get_current_user),
):
    """List the user's scheduled sends. By default returns pending + recently
    sent/failed/cancelled rows so the dashboard can show full history."""
    rows = await supabase_get_user_scheduled_sends(
        user_id=user["id"],
        statuses=["pending", "sent", "failed", "cancelled"],
        limit=100,
    )
    return {"status": "ok", "items": rows}


@app.delete("/api/scheduled-sends/{scheduled_send_id}")
async def api_cancel_scheduled_send(
    scheduled_send_id: str,
    user: dict[str, Any] = Depends(get_current_user),
):
    """Cancel a pending scheduled send. No-op for already-sent rows."""
    row = await supabase_get_scheduled_send_by_id(
        scheduled_send_id=scheduled_send_id,
        user_id=user["id"],
    )
    if not row:
        raise HTTPException(status_code=404, detail="Scheduled send not found")
    if row.get("status") != "pending":
        return {
            "status": "ok",
            "scheduled_send": row,
            "note": f"Cannot cancel — already {row.get('status')}",
        }
    updated = await supabase_update_scheduled_send(
        scheduled_send_id=scheduled_send_id,
        status="cancelled",
    )
    return {"status": "ok", "scheduled_send": updated}


# ============ SNOOZE ============
class SnoozeCreate(BaseModel):
    wake_at: str  # ISO 8601, with or without trailing Z


@app.post("/api/emails/{email_id}/snooze")
async def api_snooze_email(
    email_id: str,
    body: SnoozeCreate,
    user: dict[str, Any] = Depends(get_current_user),
):
    """Snooze a mail: remove from Inbox + status label, add OfficeFlow/Snoozed
    label. Wake worker restores it at wake_at.
    """
    # 1) Validate wake_at is parseable and in the future (with 30s grace).
    try:
        when = datetime.fromisoformat(body.wake_at.replace("Z", "+00:00"))
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid wake_at — expected ISO 8601")
    if when.tzinfo is None:
        when = when.replace(tzinfo=timezone.utc)
    now = datetime.now(timezone.utc)
    if when < now + timedelta(seconds=30):
        raise HTTPException(status_code=400, detail="wake_at must be at least 30s in the future")

    # 2) Resolve email + verify ownership.
    email_row = await supabase_get_email_by_id(email_id=email_id, user_id=user["id"])
    if not email_row:
        raise HTTPException(status_code=404, detail="Email not found")

    gmail_message_id = email_row.get("gmail_message_id")
    gmail_thread_id = email_row.get("gmail_thread_id")
    mailbox_id = email_row.get("mailbox_id")
    if not (gmail_message_id and mailbox_id):
        raise HTTPException(status_code=400, detail="Email missing Gmail/mailbox link")

    # 3) Mailbox must be connected.
    mailbox = await supabase_get_mailbox_by_user_id(user_id=user["id"], provider="gmail")
    if not mailbox or mailbox.get("status") != "connected":
        raise HTTPException(
            status_code=400,
            detail=f"Mailbox not connected (status={mailbox.get('status') if mailbox else 'missing'})",
        )

    # 4) Reject double-snooze proactively (the partial unique index also enforces this).
    existing = await supabase_get_pending_snooze_for_email(
        email_id=email_id, user_id=user["id"],
    )
    if existing:
        raise HTTPException(
            status_code=409,
            detail="Email is already snoozed",
        )

    # 5) Detect current OfficeFlow status label so we can restore it on wake.
    original_label, original_label_id = await detect_current_status_label(
        user_id=user["id"], gmail_message_id=gmail_message_id,
    )

    # 6) Resolve (or create) the Snoozed label id.
    try:
        snoozed_label_id = await ensure_label_exists(
            user_id=user["id"], label_name=SNOOZED_LABEL,
        )
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"Could not create Snoozed label: {exc}")

    # 7) Apply the Gmail label change FIRST. If it fails we don't want a
    #    pending DB row that doesn't reflect Gmail state.
    try:
        await snooze_thread_in_gmail(
            user_id=user["id"],
            gmail_message_id=gmail_message_id,
            gmail_thread_id=gmail_thread_id,
            original_label_id=original_label_id,
            snoozed_label_id=snoozed_label_id,
        )
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"Gmail snooze failed: {exc}")

    # 8) Persist the snooze row. If this fails we have a Gmail-snoozed mail
    #    with no DB record — the wake worker won't restore it. Surface clearly.
    try:
        row = await supabase_insert_snoozed_email(
            user_id=user["id"],
            mailbox_id=mailbox_id,
            email_id=email_id,
            gmail_message_id=gmail_message_id,
            gmail_thread_id=gmail_thread_id,
            subject=email_row.get("subject"),
            original_label=original_label,
            original_label_id=original_label_id,
            wake_at_iso=when.isoformat(),
        )
    except Exception as exc:
        # Best-effort rollback: put the mail back in Inbox so user isn't
        # left with a "ghost snooze" we can't track.
        try:
            await wake_thread_in_gmail(
                user_id=user["id"],
                gmail_message_id=gmail_message_id,
                gmail_thread_id=gmail_thread_id,
                original_label_id=original_label_id,
                snoozed_label_id=snoozed_label_id,
            )
        except Exception:
            pass
        raise HTTPException(status_code=500, detail=f"Snooze persist failed: {exc}")

    return {"status": "ok", "snooze": row}


@app.get("/api/emails/snoozed")
async def api_list_snoozed(
    user: dict[str, Any] = Depends(get_current_user),
):
    """List the user's snoozes. Returns pending + recently woken/cancelled/failed
    so the dashboard can show full history."""
    rows = await supabase_get_user_snoozed_emails(
        user_id=user["id"],
        statuses=["pending", "woken", "cancelled", "failed"],
        limit=100,
    )
    return {"status": "ok", "items": rows}


@app.delete("/api/emails/{email_id}/snooze")
async def api_cancel_snooze_for_email(
    email_id: str,
    user: dict[str, Any] = Depends(get_current_user),
):
    """Cancel a pending snooze for a given email — wake the mail right now."""
    row = await supabase_get_pending_snooze_for_email(
        email_id=email_id, user_id=user["id"],
    )
    if not row:
        raise HTTPException(status_code=404, detail="No pending snooze for this email")

    snoozed_label_id = None
    try:
        snoozed_label_id = await ensure_label_exists(
            user_id=user["id"], label_name=SNOOZED_LABEL,
        )
    except Exception as exc:
        print(f"[snooze-cancel] could not resolve Snoozed label "
              f"for user {user['id']}: {repr(exc)}")

    try:
        await wake_thread_in_gmail(
            user_id=user["id"],
            gmail_message_id=row.get("gmail_message_id"),
            gmail_thread_id=row.get("gmail_thread_id"),
            original_label_id=row.get("original_label_id"),
            snoozed_label_id=snoozed_label_id,
        )
    except Exception as exc:
        # Mark failed so we don't keep the row in pending forever, but still
        # surface the error to the caller.
        try:
            await supabase_update_snoozed_email(
                snooze_id=row["id"], status="failed", error_message=repr(exc)[:480],
            )
        except Exception:
            pass
        raise HTTPException(status_code=502, detail=f"Gmail wake failed: {exc}")

    updated = await supabase_update_snoozed_email(
        snooze_id=row["id"],
        status="cancelled",
        woken_at=datetime.now(timezone.utc).isoformat(),
    )
    return {"status": "ok", "snooze": updated}


# ============ ATTACHMENT SUMMARIES ============
@app.get("/api/attachment-summaries")
async def api_attachment_summaries(
    email_ids: str = "",
    user: dict[str, Any] = Depends(get_current_user),
):
    """Return AI summaries for attachments on the given draft email_ids.
    Frontend passes a comma-separated list (max 50). Empty list returns []."""
    if not email_ids:
        return {"status": "ok", "items": []}
    raw = [eid.strip() for eid in email_ids.split(",") if eid.strip()]
    raw = raw[:50]
    rows = await supabase_get_attachment_summaries_for_emails(
        user_id=user["id"],
        email_ids=raw,
    )
    return {"status": "ok", "items": rows}


@app.patch("/api/preferences/me")
async def api_update_my_preferences(
    body: PromptSettingsBody,
    user: dict[str, Any] = Depends(get_current_user),
):
    """Update user_settings for the logged-in user. Same shape as the
    onboarding endpoint, just authenticated via Supabase JWT instead of email."""
    full_payload = PromptSettingsPayload(
        email=user.get("email", ""),
        **body.model_dump(),
    )
    clean = build_clean_settings_payload(full_payload)
    saved = await supabase_upsert_user_settings(
        user_id=user["id"],
        payload=clean,
    )
    return {"status": "ok", "settings": saved}


@app.post("/onboarding/complete")
async def onboarding_complete(payload: OnboardingCompletePayload):
    user = await ensure_user_has_access(payload.email)

    mailbox = await supabase_get_mailbox_by_user_and_email(
        user_id=user["id"],
        email_address=payload.email,
        provider="gmail",
    )

    if not mailbox:
        mailbox = await supabase_get_mailbox_by_user_id(
            user_id=user["id"],
            provider="gmail",
        )

    if not mailbox:
        raise HTTPException(status_code=404, detail="Gmail mailbox not found")

    await supabase_update_mailbox_status(
        mailbox_id=mailbox["id"],
        status="connected",
    )

    # Send Supabase Auth invite email so the user can set a password and
    # log in at /login.html. Best-effort: failure here must not block
    # onboarding, because Gmail is already connected successfully.
    try:
        invite_already_sent = bool(user.get("invite_sent_at"))
        if not invite_already_sent:
            invited = await send_welcome_invite(payload.email)
            if invited:
                await supabase_patch(
                    f"/rest/v1/users?id=eq.{quote(user['id'], safe='')}",
                    {"invite_sent_at": utc_now_iso()},
                )
    except Exception as exc:
        print(f"[onboarding] Invite send skipped: {repr(exc)}")

    onboarding_state = await supabase_get_onboarding_state(user["id"])
    existing_first_draft_generated = False
    if onboarding_state:
        existing_first_draft_generated = bool(onboarding_state.get("first_draft_generated", False))

    await supabase_upsert_onboarding_state(
        user_id=user["id"],
        gmail_connected=True,
        profile_completed=True,
        initial_sync_completed=False,
        first_draft_generated=existing_first_draft_generated,
    )

    # Train style profile op basis van SENT mail, indien aangezet door user
    settings = await supabase_get_user_settings(user["id"])
    if settings and settings.get("style_learning_enabled"):
        try:
            await train_style_profile_from_sent_messages(
                user_id=user["id"],
                source_limit=settings.get("style_learning_source_limit") or 20,
            )
        except HTTPException:
            pass
        except Exception as exc:
            print(f"Style profile training skipped: {repr(exc)}")

    labels_result = await setup_gmail_labels_for_mailbox(
        user_id=user["id"],
        mailbox_id=mailbox["id"],
    )

    cleanup_result = await cleanup_legacy_labels_for_mailbox(
        user_id=user["id"],
        mailbox_id=mailbox["id"],
    )

    # Initial onboarding run: pak de laatste 10 inbox-mails ongeacht read-status
    # of cutoff. Geeft elke nieuwe gebruiker direct labels + drafts om te zien
    # i.p.v. een lege dashboard te zien tot er nieuwe mail binnenkomt.
    #
    # Bewust try/except: als de classificatie/draft-stap faalt (OpenAI flakey,
    # Gmail API hiccup, etc.) mag de gebruiker nog steeds door naar success.html.
    # De cron pakt 'm bij de volgende run alsnog op.
    process_result: dict[str, Any] = {"count": 0, "messages": []}
    try:
        process_result = await process_inbox_for_user(
            email=payload.email,
            max_results=AUTO_PROCESS_MAX_RESULTS,
            initial_run=True,
        )
    except Exception as exc:
        print(f"[onboarding] Initial inbox processing failed: {repr(exc)}")

    first_draft_generated = existing_first_draft_generated or any(
        message.get("draft_created") is True
        for message in process_result.get("messages", [])
    )

    await supabase_upsert_onboarding_state(
        user_id=user["id"],
        gmail_connected=True,
        profile_completed=True,
        initial_sync_completed=True,
        first_draft_generated=first_draft_generated,
    )

    return {
        "status": "ok",
        "email": payload.email,
        "mailbox_status": "connected",
        "profile_completed": True,
        "initial_sync_completed": True,
        "first_draft_generated": first_draft_generated,
        "labels_created_count": len(labels_result),
        "legacy_labels_cleaned_count": len(cleanup_result),
        "processed_count": process_result.get("count", 0),
        "messages": process_result.get("messages", []),
    }


@app.get("/settings/prompt")
async def get_prompt_settings(email: str):
    user = await ensure_user_has_access(email)
    settings = await supabase_get_user_settings(user["id"])

    return {
        "status": "ok",
        "email": email,
        "settings": settings,
    }


@app.get("/settings/style-profile")
async def get_style_profile(email: str):
    user = await ensure_user_has_access(email)
    profile = await supabase_get_user_style_profile(user["id"])

    return {
        "status": "ok",
        "email": email,
        "style_profile": profile,
    }


@app.post("/internal/train-style-profile")
async def train_style_profile(
    email: str = Body(...),
    source_limit: int = Body(default=30),
    _: None = Depends(require_internal_secret),
):
    user = await ensure_user_has_access(email)

    result = await train_style_profile_from_sent_messages(
        user_id=user["id"],
        source_limit=source_limit,
    )

    return {
        "status": "ok",
        "email": email,
        "source_sent_count": result["source_sent_count"],
        "style_profile": result["saved_profile"],
    }


# ---------------------------------------------------------------------------
# CRON: hertrainen van stijlprofielen (maandelijks)
# ---------------------------------------------------------------------------

@app.post("/internal/style-profile/retrain-stale")
async def retrain_stale_style_profiles(
    age_days: int = 30,
    max_per_run: int = 50,
    _: None = Depends(require_internal_secret),
):
    """
    Hertraint stijlprofielen die ouder zijn dan `age_days` dagen, voor users
    met style_learning_enabled=True. Bedoeld om periodiek (maandelijks) door
    een cron-job te worden aangeroepen.

    Veiligheden:
    - max_per_run cap voorkomt dat 1 batch alle Gmail/OpenAI quota opmaakt
    - per-user try/except: 1 fout breekt de batch niet
    - users zonder bestaand profiel worden ook hertraind (eerste keer)
    """
    try:
        users_settings = await supabase_get(
            "/rest/v1/user_settings"
            "?style_learning_enabled=eq.true"
            "&select=user_id,style_learning_source_limit"
            f"&limit={max_per_run * 5}"
        )
    except Exception as exc:
        print(f"[retrain-stale] settings fetch failed: {repr(exc)}")
        return {"status": "error", "error": "settings_fetch_failed"}

    if not isinstance(users_settings, list):
        return {"status": "ok", "retrained": 0, "skipped": 0, "total_users_checked": 0}

    threshold_iso = (_dt.now(_tz.utc) - _td(days=age_days)).isoformat()
    retrained = 0
    skipped_fresh = 0
    failed = 0

    for s in users_settings:
        if retrained >= max_per_run:
            break
        user_id = s.get("user_id")
        if not user_id:
            continue
        try:
            existing = await supabase_get_user_style_profile(user_id)
            updated_at = None
            if existing:
                updated_at = existing.get("updated_at") or existing.get("created_at")

            if existing and updated_at and updated_at > threshold_iso:
                skipped_fresh += 1
                continue

            await train_style_profile_from_sent_messages(
                user_id=user_id,
                source_limit=s.get("style_learning_source_limit") or 20,
            )
            retrained += 1
            print(f"[retrain-stale] retrained user_id={user_id}")
        except HTTPException as http_exc:
            # bv. 'Not enough usable sent emails' — geen kritiek, log + door
            failed += 1
            print(f"[retrain-stale] user {user_id} skipped (HTTP): {http_exc.detail}")
            continue
        except Exception as exc:
            failed += 1
            print(f"[retrain-stale] user {user_id} failed: {repr(exc)}")
            continue

    summary = {
        "status": "ok",
        "retrained": retrained,
        "skipped_fresh": skipped_fresh,
        "failed": failed,
        "total_users_checked": len(users_settings),
        "age_threshold_days": age_days,
    }
    print(f"[retrain-stale] done: {summary}")
    return summary


# ---------------------------------------------------------------------------
# Subscription management (dashboard)
# ---------------------------------------------------------------------------

def _stripe_sub_period_end_iso(subscription: Any) -> Optional[str]:
    try:
        period_end = subscription.get("current_period_end")
        if period_end:
            return datetime.fromtimestamp(int(period_end), tz=timezone.utc).isoformat()
    except Exception:
        return None
    return None


def _stripe_sub_price_amount(subscription: Any) -> Optional[str]:
    try:
        items = subscription.get("items", {}).get("data", [])
        if not items:
            return None
        price = items[0].get("price") or {}
        unit_amount = price.get("unit_amount")
        currency = (price.get("currency") or "eur").upper()
        if unit_amount is None:
            return None
        return f"{currency} {unit_amount / 100:.2f}"
    except Exception:
        return None


@app.get("/api/subscription/status")
async def get_subscription_status(user: dict[str, Any] = Depends(get_current_user)):
    stripe_customer_id = user.get("stripe_customer_id")
    stripe_subscription_id = user.get("stripe_subscription_id")

    payload: dict[str, Any] = {
        "subscription_status": user.get("subscription_status"),
        "access_allowed": user.get("access_allowed"),
        "cancel_at_period_end": False,
        "current_period_end": None,
        "plan_price": user.get("plan_price"),
        "plan_name": user.get("plan_name"),
        "cancels_at": user.get("cancels_at"),
        "cancelled_at": user.get("cancelled_at"),
        "deletion_scheduled_at": user.get("deletion_scheduled_at"),
    }

    if not STRIPE_SECRET_KEY or not stripe_subscription_id:
        return payload

    try:
        stripe.api_key = STRIPE_SECRET_KEY
        sub = stripe.Subscription.retrieve(stripe_subscription_id)
        sub_dict = sub if isinstance(sub, dict) else sub.to_dict()
        payload["cancel_at_period_end"] = bool(sub_dict.get("cancel_at_period_end"))
        payload["current_period_end"] = _stripe_sub_period_end_iso(sub_dict)
        price_label = _stripe_sub_price_amount(sub_dict)
        if price_label and not payload["plan_price"]:
            payload["plan_price"] = price_label
    except Exception as exc:
        print(f"[subscription/status] Stripe lookup failed: {repr(exc)}")

    return payload


@app.post("/api/subscription/cancel")
async def cancel_subscription(
    body: CancelSubscriptionPayload,
    user: dict[str, Any] = Depends(get_current_user),
):
    stripe_subscription_id = user.get("stripe_subscription_id")
    if not stripe_subscription_id:
        raise HTTPException(status_code=400, detail="Geen actief abonnement gevonden.")

    require_env(STRIPE_SECRET_KEY, "STRIPE_SECRET_KEY")
    stripe.api_key = STRIPE_SECRET_KEY

    try:
        sub = stripe.Subscription.modify(
            stripe_subscription_id,
            cancel_at_period_end=True,
        )
        sub_dict = sub if isinstance(sub, dict) else sub.to_dict()
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Stripe cancel failed: {exc}")

    period_end_iso = _stripe_sub_period_end_iso(sub_dict)

    patch_payload: dict[str, Any] = {
        "subscription_status": "canceling",
        "access_allowed": True,
        "cancels_at": period_end_iso,
        "cancel_reason": body.reason,
        "cancel_feedback": body.feedback,
    }

    try:
        await supabase_patch(
            f"/rest/v1/users?id=eq.{quote(user['id'], safe='')}",
            patch_payload,
        )
    except Exception as exc:
        # Don't fail the cancel on Supabase write errors; Stripe is source of truth
        print(f"[subscription/cancel] Supabase update warning: {repr(exc)}")

    return {
        "status": "ok",
        "cancel_at_period_end": True,
        "current_period_end": period_end_iso,
    }


@app.post("/api/subscription/reactivate")
async def reactivate_subscription(user: dict[str, Any] = Depends(get_current_user)):
    stripe_subscription_id = user.get("stripe_subscription_id")
    if not stripe_subscription_id:
        raise HTTPException(status_code=400, detail="Geen abonnement gevonden.")

    require_env(STRIPE_SECRET_KEY, "STRIPE_SECRET_KEY")
    stripe.api_key = STRIPE_SECRET_KEY

    try:
        sub = stripe.Subscription.modify(
            stripe_subscription_id,
            cancel_at_period_end=False,
        )
        sub_dict = sub if isinstance(sub, dict) else sub.to_dict()
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Stripe reactivate failed: {exc}")

    period_end_iso = _stripe_sub_period_end_iso(sub_dict)

    try:
        await supabase_patch(
            f"/rest/v1/users?id=eq.{quote(user['id'], safe='')}",
            {
                "subscription_status": "active",
                "access_allowed": True,
                "cancels_at": None,
                "cancel_reason": None,
                "cancel_feedback": None,
            },
        )
    except Exception as exc:
        print(f"[subscription/reactivate] Supabase update warning: {repr(exc)}")

    return {
        "status": "ok",
        "cancel_at_period_end": False,
        "current_period_end": period_end_iso,
    }


@app.get("/api/stats/overview")
async def stats_overview(user: dict[str, Any] = Depends(get_current_user)):
    """
    Dashboard statistieken voor de ingelogde user.
    Defensief: als een sub-query faalt, geeft dat veld 0/None i.p.v. 500.
    """
    user_id = user.get("id")
    if not user_id:
        raise HTTPException(status_code=400, detail="User id ontbreekt")

    now = datetime.now(timezone.utc)
    month_start = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    prev_month_start = (month_start - timedelta(days=1)).replace(
        day=1, hour=0, minute=0, second=0, microsecond=0
    )
    seven_days_ago = now - timedelta(days=7)

    user_filter = f"user_id=eq.{quote(user_id, safe='')}"
    # URL-encode alles — anders wordt de `+` in de tz-offset als spatie gelezen
    # door PostgREST en faalt elke filter (count_rows geeft dan 0).
    month_iso = quote(month_start.isoformat(), safe='')
    prev_month_start_iso = quote(prev_month_start.isoformat(), safe='')
    prev_month_end_iso = quote(month_start.isoformat(), safe='')
    seven_days_iso = quote(seven_days_ago.isoformat(), safe='')

    async def count_rows(table: str, extra: str) -> int:
        try:
            data = await supabase_get(
                f"/rest/v1/{table}?{user_filter}&{extra}&select=id"
            )
            return len(data) if isinstance(data, list) else 0
        except Exception as exc:
            print(f"[stats] count {table} failed: {repr(exc)}")
            return 0

    async def fetch_timestamps(table: str, extra: str) -> list[str]:
        try:
            data = await supabase_get(
                f"/rest/v1/{table}?{user_filter}&{extra}&select=created_at"
            )
            if isinstance(data, list):
                return [r.get("created_at") for r in data if r.get("created_at")]
            return []
        except Exception as exc:
            print(f"[stats] fetch {table} failed: {repr(exc)}")
            return []

    # counts deze maand
    emails_classified = await count_rows(
        "emails", f"created_at=gte.{month_iso}"
    )
    drafts_ready = await count_rows(
        "drafts",
        f"created_at=gte.{month_iso}&status=neq.sent",
    )
    drafts_this_month = await count_rows(
        "drafts", f"created_at=gte.{month_iso}"
    )

    # counts vorige maand (voor delta)
    emails_prev = await count_rows(
        "emails",
        f"created_at=gte.{prev_month_start_iso}&created_at=lt.{prev_month_end_iso}",
    )
    drafts_prev = await count_rows(
        "drafts",
        f"created_at=gte.{prev_month_start_iso}&created_at=lt.{prev_month_end_iso}",
    )

    # time saved: 3 min per draft + 0.25 min per geclassificeerde mail
    time_saved_minutes = round(drafts_this_month * 3 + emails_classified * 0.25)
    time_saved_prev = round(drafts_prev * 3 + emails_prev * 0.25)

    time_saved_delta_pct: int | None = None
    if time_saved_prev > 0:
        time_saved_delta_pct = round(
            ((time_saved_minutes - time_saved_prev) / time_saved_prev) * 100
        )

    # weekly activity (laatste 7 dagen, Ma=0..Zo=6)
    weekly_activity = [0] * 7
    timestamps = await fetch_timestamps(
        "emails", f"created_at=gte.{seven_days_iso}"
    )
    for ts in timestamps:
        try:
            dt = datetime.fromisoformat(ts.replace("Z", "+00:00"))
            weekday = dt.weekday()
            if 0 <= weekday <= 6:
                weekly_activity[weekday] += 1
        except Exception:
            continue

    # statuslabels leven in Gmail — live ophalen via labels.get per label.
    # Gemiddeld kost dat ~5 Gmail API-calls per dashboard-load. Gemakkelijk
    # binnen quota. Alle calls zijn parallel (asyncio.gather).
    priority_caught_pct: int | None = None
    fast_reply_pct: int | None = None  # vereist reply-timing analyse, voor later
    followup_active = 0
    waiting_reply = 0
    spam_filtered = 0

    try:
        all_labels = await get_all_gmail_labels(user_id)
    except Exception as exc:
        print(f"[stats-overview] gmail labels list failed: {repr(exc)}")
        all_labels = {}

    async def _label_thread_count(label_name: str) -> int:
        meta = all_labels.get(label_name)
        if not meta or not meta.get("id"):
            return 0
        try:
            full = await gmail_get_json_for_user(
                user_id=user_id,
                url=f"{GMAIL_API_BASE}/labels/{meta['id']}",
            )
            return int(full.get("threadsTotal") or 0) if isinstance(full, dict) else 0
        except Exception as exc:
            print(f"[stats-overview] label '{label_name}' fetch failed: {repr(exc)}")
            return 0

    if all_labels:
        try:
            marketing_n, ignore_n, followup_n, waiting_n = await asyncio.gather(
                _label_thread_count("Marketing"),
                _label_thread_count("Ignore"),
                _label_thread_count("Follow Up"),
                _label_thread_count("Waiting On Reply"),
            )
            spam_filtered = marketing_n + ignore_n
            followup_active = followup_n
            waiting_reply = waiting_n
        except Exception as exc:
            print(f"[stats-overview] label-counts gather failed: {repr(exc)}")

    return {
        "time_saved_minutes": time_saved_minutes,
        "time_saved_delta_pct": time_saved_delta_pct,
        "drafts_ready": drafts_ready,
        "emails_classified": emails_classified,
        "priority_caught_pct": priority_caught_pct,
        "fast_reply_pct": fast_reply_pct,
        "spam_filtered": spam_filtered,
        "followup_active": followup_active,
        "waiting_reply": waiting_reply,
        "weekly_activity": weekly_activity,
    }


# ============================================================
# PUBLIC CHATBOT — voor website-bezoekers (homepage / mailbox-page)
# ============================================================

CHATBOT_SYSTEM_PROMPT = """Je bent **Floor**, de vriendelijke OfficeFlow-assistent. Je heet Floor en stelt jezelf zo voor als de gebruiker er expliciet om vraagt — niet bij elk bericht.

OfficeFlow is een Nederlandse SaaS voor MKB met 3 producten:

**1. Mailbox Manager (solo)** — €24/mnd of €240/jaar (2 mnd gratis)
- AI die Gmail-inbox begrijpt en 9 labels toepast (Priority, To Respond, Waiting On Reply, Follow Up, Done, FYI, Notification, Marketing, Ignore)
- Concept-antwoorden in jouw eigen schrijfstijl
- Follow-Up Radar (opvolg-concepten voor stille threads + open beloftes)
- Auto-archive optioneel
- Werkt IN Gmail — geen nieuwe app
- NOOIT auto-send. Drafts zijn altijd suggesties.
- Werkt nu alleen met Gmail; Outlook staat op de roadmap

**2. Mailbox Manager Teams** — voor bedrijven met meerdere inboxen
- Team Starter: €59/mnd, 3 mailboxen, tot 5 team-leden
- Team Pro: €149/mnd, 10 mailboxen, tot 20 leden, priority support
- Business: €299/mnd, 25 mailboxen, onbeperkt leden, API + SLA + dedicated manager
- Gedeeld dashboard, één factuur, per lid eigen stijlprofiel
- Gedeelde inboxen (info@, verkoop@) ondersteund

**3. Facturen & Offertes** — €19-89/mnd
- Compleet platform: offerte → factuur → werkregistratie → Moneybird-sync
- Voor Nederlandse Moneybird-gebruikers
- Abonnementen + terugkerende facturen + automatische betalingsopvolging

**Status (belangrijk):**
- Mailbox Manager wordt op dit moment door Google geverifieerd. Aanmelden kan, maar gebruikers krijgen een "niet-geverifieerd" waarschuwing tijdens OAuth — klik op "Geavanceerd" → "Doorgaan naar OfficeFlow" om verder te gaan.
- Betaalmethoden: iDEAL, Bancontact, Kaart, Apple Pay
- BTW 21% wordt automatisch toegevoegd via Stripe
- Alle abonnementen maandelijks opzegbaar
- Data: hosting in EU, AVG-conform

**Hoe je antwoordt:**
- Nederlands, vriendelijk, kort en concreet
- Bij prijs-vragen: noem de juiste tier
- Bij feature-vragen: wees specifiek
- Bij support / technische problemen: verwijs door naar support@officeflowcompany.com
- Bij vragen waar je het antwoord niet zeker weet: zeg dat eerlijk en verwijs door naar support@officeflowcompany.com

**Wat je NIET doet:**
- Beloftes over uitkomsten ("dit gaat 10 uur per week besparen")
- Korting geven of beloven
- Persoonlijke info opvragen
- Concurrenten benoemen of vergelijken
- Klagen of meegaan in negatieve reviews
- Drafts van klantmails schrijven (dat doet OfficeFlow zelf in Gmail)

Houd antwoorden onder de 4 zinnen tenzij de vraag echt om uitleg vraagt."""


class _ChatMessage(BaseModel):
    role: str  # 'user' | 'assistant'
    content: str


class _ChatBody(BaseModel):
    message: str
    history: list[_ChatMessage] = []


@app.post("/api/chat")
async def public_chat(body: _ChatBody, request: Request):
    """Public chatbot endpoint — geen auth nodig.
    Beperkt: max 500 chars per message, max 10 turns history."""
    api_key = require_env(OPENAI_API_KEY, "OPENAI_API_KEY")

    user_message = (body.message or "").strip()[:500]
    if not user_message:
        raise HTTPException(status_code=400, detail="Bericht is leeg.")

    # Bouw chat-history op (max 10 turns)
    msgs: list[dict[str, str]] = [{"role": "system", "content": CHATBOT_SYSTEM_PROMPT}]
    for h in (body.history or [])[-10:]:
        if h.role in ("user", "assistant") and h.content:
            msgs.append({"role": h.role, "content": h.content[:500]})
    msgs.append({"role": "user", "content": user_message})

    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            res = await client.post(
                "https://api.openai.com/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json",
                },
                json={
                    "model": "gpt-4o-mini",
                    "messages": msgs,
                    "max_tokens": 300,
                    "temperature": 0.4,
                },
            )
        data = res.json()
        if res.status_code >= 400:
            print(f"[chatbot] OpenAI error: {data}")
            raise HTTPException(status_code=502, detail="AI-service niet bereikbaar.")

        reply = data.get("choices", [{}])[0].get("message", {}).get("content", "").strip()
        if not reply:
            reply = "Sorry, ik kon dat niet verwerken. Mail support@officeflowcompany.com voor hulp."

        return {"reply": reply}
    except HTTPException:
        raise
    except Exception as exc:
        print(f"[chatbot] failed: {repr(exc)}")
        raise HTTPException(status_code=500, detail="Iets ging mis. Probeer opnieuw.")


@app.post("/webhooks/stripe")
async def stripe_webhook(request: Request):
    webhook_secret = require_env(STRIPE_WEBHOOK_SECRET, "STRIPE_WEBHOOK_SECRET")

    def stripe_obj_get(obj: Any, key: str, default: Any = None) -> Any:
        if obj is None:
            return default
        if isinstance(obj, dict):
            return obj.get(key, default)
        return getattr(obj, key, default)

    try:
        payload = await request.body()
        sig_header = request.headers.get("stripe-signature")

        if not sig_header:
            return JSONResponse(
                status_code=400,
                content={"error": "Missing Stripe signature header"},
            )

        event = stripe.Webhook.construct_event(
            payload=payload,
            sig_header=sig_header,
            secret=webhook_secret,
        )

        # Idempotency: Stripe delivers at-least-once. Two-layer dedup:
        #   1. In-memory cache (fastest, survives within one process)
        #   2. webhook_events DB table (survives restarts + multi-instance)
        # The DB layer is authoritative — if it says "already processed",
        # we skip even if the in-memory cache hasn't seen it yet.
        event_id = event.get("id") if isinstance(event, dict) else getattr(event, "id", None)
        event_type = event["type"]
        event_livemode = event.get("livemode") if isinstance(event, dict) else None

        # Fast path: in-memory check
        if stripe_event_already_processed(event_id):
            print(f"[stripe-webhook] Skipping duplicate event {event_id} (in-memory)")
            return JSONResponse(
                status_code=200,
                content={"received": True, "deduped": True, "source": "memory"},
            )

        # Authoritative path: DB-backed claim
        is_new = await supabase_record_webhook_event(
            stripe_event_id=event_id or "",
            event_type=event_type,
            livemode=event_livemode,
        )
        if not is_new:
            print(f"[stripe-webhook] Skipping duplicate event {event_id} (DB)")
            return JSONResponse(
                status_code=200,
                content={"received": True, "deduped": True, "source": "db"},
            )

        raw_data = event["data"]["object"]

        # Convert Stripe StripeObject → pure Python dict (JSON roundtrip)
        # Voorkomt KeyError: 0 bij dict() of iteratie van StripeObject
        import json as _json
        try:
            data = _json.loads(_json.dumps(raw_data, default=lambda o: getattr(o, "to_dict", lambda: dict(o))()))
        except Exception as exc:
            print(f"[webhook] failed to normalize event data: {repr(exc)}, falling back to raw")
            data = raw_data

        if event_type == "checkout.session.completed":
            from app import teams as teams_module

            # Harvest metadata keys defensief — probeer alle mogelijke paden
            session_metadata: dict[str, Any] = {}
            try:
                for k in ("product_family", "tier", "seats", "billing_period", "email", "team_name"):
                    raw_md = stripe_obj_get(data, "metadata")
                    if raw_md is None:
                        break
                    val = stripe_obj_get(raw_md, k)
                    if val is not None:
                        session_metadata[k] = val
                print(f"[teams-webhook] extracted metadata: {session_metadata}")
            except Exception as exc:
                print(f"[teams-webhook] metadata extract failed: {repr(exc)}")

            # Team checkout? Route naar teams.py
            if teams_module.is_team_checkout(session_metadata):
                try:
                    await teams_module.handle_team_checkout_completed(data)
                    return {"received": True}
                except Exception as exc:
                    import traceback
                    print(f"[teams-webhook] FAILED in team-route: {repr(exc)}")
                    print(traceback.format_exc())
                    # Niet re-raisen — laat solo-logica hieronder niet lopen bij team-checkouts
                    return {"received": True, "team_error": str(exc)}

            # Solo checkout (bestaand gedrag)
            customer_details = stripe_obj_get(data, "customer_details") or {}
            email = (
                stripe_obj_get(data, "customer_email")
                or stripe_obj_get(data, "client_reference_id")
                or stripe_obj_get(customer_details, "email")
            )
            customer_id = stripe_obj_get(data, "customer")
            subscription_id = stripe_obj_get(data, "subscription")

            if not email:
                return {"received": True}
            email = email.strip().lower()

            user = await supabase_get_user_by_email(email)
            if not user:
                user = await supabase_insert_user(email=email, full_name=None)

            await supabase_update_user_subscription(
                user_id=user["id"],
                subscription_status="active",
                access_allowed=True,
                stripe_customer_id=customer_id,
                stripe_subscription_id=subscription_id,
            )

        elif event_type == "customer.subscription.updated":
            from app import teams as teams_module

            cancel_at_period_end = bool(stripe_obj_get(data, "cancel_at_period_end", False))

            # Team-subscription? → route eerst daarheen; als handled, stop
            team_handled = await teams_module.handle_team_subscription_updated(
                data,
                cancel_at_period_end=cancel_at_period_end,
            )
            if team_handled:
                return {"received": True}

            # Solo-subscription (bestaand gedrag)
            status = stripe_obj_get(data, "status")
            customer_id = stripe_obj_get(data, "customer")
            subscription_id = stripe_obj_get(data, "id")

            if customer_id:
                user = await supabase_get_user_by_stripe_customer_id(customer_id)
                if user:
                    if cancel_at_period_end and status in ALLOWED_SUBSCRIPTION_STATUSES:
                        normalized_status = "canceling"
                        access_allowed = True
                    else:
                        normalized_status = status
                        access_allowed = status in ALLOWED_SUBSCRIPTION_STATUSES

                    await supabase_update_user_subscription(
                        user_id=user["id"],
                        subscription_status=normalized_status,
                        access_allowed=access_allowed,
                        stripe_customer_id=customer_id,
                        stripe_subscription_id=subscription_id,
                    )

                    mailbox = await supabase_get_mailbox_by_user_id(
                        user_id=user["id"],
                        provider="gmail",
                    )
                    if mailbox:
                        mailbox_status = "connected" if access_allowed else "canceled"
                        await supabase_update_mailbox_status(
                            mailbox_id=mailbox["id"],
                            status=mailbox_status,
                        )

        elif event_type == "customer.subscription.deleted":
            from app import teams as teams_module

            # Team-subscription? → route daarheen
            team_handled = await teams_module.handle_team_subscription_deleted(data)
            if team_handled:
                return {"received": True}

            # Solo-subscription (bestaand gedrag)
            customer_id = stripe_obj_get(data, "customer")
            subscription_id = stripe_obj_get(data, "id")

            if customer_id:
                user = await supabase_get_user_by_stripe_customer_id(customer_id)
                if user:
                    await supabase_update_user_subscription(
                        user_id=user["id"],
                        subscription_status="canceled",
                        access_allowed=False,
                        stripe_customer_id=customer_id,
                        stripe_subscription_id=subscription_id,
                    )

                    # Schedule 30-day data deletion (matches privacy policy)
                    now_utc = datetime.now(timezone.utc)
                    deletion_at = now_utc + timedelta(days=30)
                    try:
                        await supabase_patch(
                            f"/rest/v1/users?id=eq.{quote(user['id'], safe='')}",
                            {
                                "cancelled_at": now_utc.isoformat(),
                                "deletion_scheduled_at": deletion_at.isoformat(),
                            },
                        )
                    except Exception as exc:
                        print(f"[stripe-webhook] deletion schedule write failed: {repr(exc)}")

                    mailbox = await supabase_get_mailbox_by_user_id(
                        user_id=user["id"],
                        provider="gmail",
                    )
                    if mailbox:
                        await supabase_update_mailbox_status(
                            mailbox_id=mailbox["id"],
                            status="canceled",
                        )

        # Mark webhook event as successfully processed
        if event_id:
            await supabase_mark_webhook_event_processed(
                stripe_event_id=event_id,
                success=True,
            )

        return {"received": True}

    except stripe.error.SignatureVerificationError:
        return JSONResponse(status_code=400, content={"error": "Invalid Stripe signature"})
    except ValueError:
        return JSONResponse(status_code=400, content={"error": "Invalid Stripe payload"})
    except Exception as exc:
        # Mark event as failed so we can investigate later. Best-effort.
        try:
            if 'event_id' in locals() and event_id:
                await supabase_mark_webhook_event_processed(
                    stripe_event_id=event_id,
                    success=False,
                    error_message=f"{type(exc).__name__}: {str(exc)[:300]}",
                )
        except Exception:
            pass
        return JSONResponse(status_code=500, content={"error": f"Webhook handler failed: {str(exc)}"})

# =========================================================================
# OFFICEFLOW FEATURES: Morning Briefing + Follow-Up Radar + Relationship Intel
# Append this entire file to the END of main.py
# Then add ONE hook line in the sent-mail processing loop (see SECTION F)
# =========================================================================

import json as _json
from datetime import datetime as _dt, timedelta as _td, date as _date, timezone as _tz
from zoneinfo import ZoneInfo

# ---------------------------------------------------------------------------
# SECTION A — CONFIG
# ---------------------------------------------------------------------------

BRIEFING_DEFAULT_TIMEZONE = "Europe/Amsterdam"
BRIEFING_ENABLED_GLOBAL = os.getenv("BRIEFING_ENABLED", "true").lower() == "true"
COMMITMENT_DETECTION_ENABLED = os.getenv("COMMITMENT_DETECTION_ENABLED", "true").lower() == "true"


# ---------------------------------------------------------------------------
# SECTION B — FOLLOW-UP RADAR: commitment detection
# ---------------------------------------------------------------------------

async def detect_commitments_in_sent_mail(
    subject: str | None,
    body_text: str | None,
    sent_at_iso: str | None = None,
) -> dict[str, Any]:
    """
    Detects if a sent mail contains a commitment/promise (e.g. 'ik stuur morgen de offerte').
    Returns: {"has_commitment": bool, "action_text": str|null, "due_date": "YYYY-MM-DD"|null, "reason": str}
    """
    if not COMMITMENT_DETECTION_ENABLED:
        return {"has_commitment": False, "action_text": None, "due_date": None, "reason": "Detection disabled"}

    if not body_text or len(body_text.strip()) < 10:
        return {"has_commitment": False, "action_text": None, "due_date": None, "reason": "Body too short"}

    api_key = require_env(OPENAI_API_KEY, "OPENAI_API_KEY")

    try:
        reference_date = _dt.fromisoformat(sent_at_iso.replace("Z", "+00:00")).date() if sent_at_iso else _date.today()
    except Exception:
        reference_date = _date.today()

    prompt = f"""
Je analyseert een e-mail die DE GEBRUIKER zojuist heeft verstuurd.
Detecteer of de gebruiker een CONCRETE BELOFTE maakt die later actie van hem vereist.

Voorbeelden van beloftes:
- "Ik stuur je morgen de offerte"
- "Volgende week kom ik erop terug"
- "Vrijdag krijg je van mij de cijfers"
- "Ik check het en laat het je uiterlijk donderdag weten"

Geen beloftes (vragen / info / besluiten zonder toekomstige actie):
- "Dank voor je bericht"
- "Akkoord, we gaan door"
- "Kun jij vrijdag laten weten?"  (vraag aan ANDER, niet eigen belofte)

Referentiedatum (verstuurdatum): {reference_date.isoformat()}
Gebruik deze datum om relatieve tijdsaanduidingen (morgen, volgende week, vrijdag, over 3 dagen) om te zetten naar een concrete due_date.
Als er GEEN concrete deadline genoemd is maar wel een belofte, kies due_date = 7 dagen na referentiedatum.

Geef alleen geldige JSON terug:
{{
  "has_commitment": true,
  "action_text": "Offerte sturen",
  "due_date": "2026-04-22",
  "reason": "User zei 'ik stuur je morgen de offerte'"
}}

Of als er geen belofte is:
{{
  "has_commitment": false,
  "action_text": null,
  "due_date": null,
  "reason": "Enkel een dankwoord, geen toekomstige actie"
}}

Onderwerp: {subject}

E-mail:
{body_text[:3000]}
""".strip()

    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            response = await client.post(
                "https://api.openai.com/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json",
                },
                json={
                    "model": "gpt-4o-mini",
                    "messages": [{"role": "user", "content": prompt}],
                    "response_format": {"type": "json_object"},
                    "temperature": 0.1,
                },
            )
            response.raise_for_status()
            content = response.json()["choices"][0]["message"]["content"]
            result = _json.loads(content)

            # Normalize
            return {
                "has_commitment": bool(result.get("has_commitment", False)),
                "action_text": result.get("action_text"),
                "due_date": result.get("due_date"),
                "reason": result.get("reason", ""),
            }
    except Exception as exc:
        print(f"[commitment-detection] Failed: {repr(exc)}")
        return {"has_commitment": False, "action_text": None, "due_date": None, "reason": f"Error: {exc}"}


async def supabase_upsert_commitment(
    user_id: str,
    mailbox_id: str,
    gmail_thread_id: str,
    gmail_message_id: str | None,
    action_text: str,
    due_date: str | None,
    recipient_email: str | None,
    recipient_name: str | None,
    subject: str | None,
) -> dict[str, Any] | None:
    """Inserts or updates an active commitment for this thread."""
    try:
        # Check if an active commitment already exists for this thread
        existing = await supabase_get(
            f"/rest/v1/commitments?gmail_thread_id=eq.{gmail_thread_id}&status=eq.active&select=*",
        )
        if existing and len(existing) > 0:
            # Update existing
            commitment_id = existing[0]["id"]
            updated = await supabase_patch(
                f"/rest/v1/commitments?id=eq.{commitment_id}",
                {
                    "action_text": action_text,
                    "due_date": due_date,
                    "recipient_email": recipient_email,
                    "recipient_name": recipient_name,
                    "subject": subject,
                    "gmail_message_id": gmail_message_id,
                    "updated_at": _dt.now(_tz.utc).isoformat(),
                },
                prefer="return=representation",
            )
            return updated[0] if isinstance(updated, list) and updated else existing[0]

        # Insert new
        inserted = await supabase_post(
            "/rest/v1/commitments",
            {
                "user_id": user_id,
                "mailbox_id": mailbox_id,
                "gmail_thread_id": gmail_thread_id,
                "gmail_message_id": gmail_message_id,
                "action_text": action_text,
                "due_date": due_date,
                "recipient_email": recipient_email,
                "recipient_name": recipient_name,
                "subject": subject,
                "status": "active",
            },
            prefer="return=representation",
        )
        return inserted[0] if isinstance(inserted, list) and inserted else None
    except Exception as exc:
        print(f"[commitment-upsert] Failed: {repr(exc)}")
        return None


async def supabase_complete_commitments_for_thread(gmail_thread_id: str) -> int:
    """Marks active commitments for this thread as completed (used when a reply arrives)."""
    try:
        now_iso = _dt.now(_tz.utc).isoformat()
        await supabase_patch(
            f"/rest/v1/commitments?gmail_thread_id=eq.{gmail_thread_id}&status=eq.active",
            {"status": "completed", "completed_at": now_iso, "updated_at": now_iso},
        )
        return 1
    except Exception as exc:
        print(f"[commitment-complete] Failed: {repr(exc)}")
        return 0


async def process_sent_mail_for_commitment(
    user_id: str,
    mailbox_id: str,
    thread_id: str | None,
    message_id: str | None,
    subject: str | None,
    body_text: str | None,
    recipient_email: str | None,
    recipient_name: str | None,
    sent_at_iso: str | None,
    label_name_to_id: dict[str, str],
    current_label_ids: set[str],
) -> dict[str, Any]:
    """
    Wrapper: runs commitment detection on a sent mail. If detected, upserts DB record
    AND forces Follow Up label on the thread.
    Returns: {"detected": bool, "commitment": dict|null, "forced_follow_up": bool}
    """
    if not thread_id:
        return {"detected": False, "commitment": None, "forced_follow_up": False}

    detection = await detect_commitments_in_sent_mail(
        subject=subject, body_text=body_text, sent_at_iso=sent_at_iso
    )
    if not detection.get("has_commitment"):
        return {"detected": False, "commitment": None, "forced_follow_up": False}

    saved = await supabase_upsert_commitment(
        user_id=user_id,
        mailbox_id=mailbox_id,
        gmail_thread_id=thread_id,
        gmail_message_id=message_id,
        action_text=detection["action_text"] or "Opvolgen",
        due_date=detection.get("due_date"),
        recipient_email=recipient_email,
        recipient_name=recipient_name,
        subject=subject,
    )

    # Force Follow Up label (overrides any other status on thread)
    forced = False
    try:
        if "Follow Up" in label_name_to_id:
            await sync_thread_status(
                user_id=user_id,
                thread_id=thread_id,
                current_message_id=message_id or "",
                current_label_ids=current_label_ids,
                label_name_to_id=label_name_to_id,
                target_label_name="Follow Up",
            )
            forced = True
    except Exception as exc:
        print(f"[commitment-label] Force Follow Up failed: {repr(exc)}")

    return {"detected": True, "commitment": saved, "forced_follow_up": forced}


# ---------------------------------------------------------------------------
# SECTION C — MORNING BRIEFING
# ---------------------------------------------------------------------------

async def get_mailbox_briefing_config(mailbox_id: str) -> dict[str, Any]:
    try:
        rows = await supabase_get(
            f"/rest/v1/mailboxes?id=eq.{mailbox_id}&select=briefing_enabled,briefing_hour,briefing_minute,briefing_timezone,briefing_last_sent_at,email_address",
        )
        return rows[0] if rows else {}
    except Exception:
        return {}


async def count_labels_for_briefing(
    user_id: str, label_name_to_id: dict[str, str]
) -> dict[str, int]:
    """Counts threads in inbox per custom label. Returns dict {label_name: count}."""
    counts: dict[str, int] = {}
    for label_name in ("Priority", "To Respond", "Follow Up"):
        label_id = label_name_to_id.get(label_name)
        if not label_id:
            counts[label_name] = 0
            continue
        try:
            data = await gmail_get_json_for_user(
                user_id=user_id,
                url=f"{GMAIL_API_BASE}/threads",
                params={"labelIds": ["INBOX", label_id], "maxResults": 50},
            )
            counts[label_name] = data.get("resultSizeEstimate", len(data.get("threads", []) or []))
        except Exception:
            counts[label_name] = 0
    return counts


async def get_upcoming_commitments(user_id: str, days_ahead: int = 7) -> list[dict[str, Any]]:
    try:
        today = _date.today().isoformat()
        cutoff = (_date.today() + _td(days=days_ahead)).isoformat()
        rows = await supabase_get(
            f"/rest/v1/commitments?user_id=eq.{user_id}&status=eq.active&due_date=lte.{cutoff}&order=due_date.asc&select=*",
        )
        return rows or []
    except Exception as exc:
        print(f"[briefing] commitments fetch failed: {repr(exc)}")
        return []


async def get_top_priority_threads(
    user_id: str, label_name_to_id: dict[str, str], limit: int = 3
) -> list[dict[str, Any]]:
    """Returns top 3 threads from Priority > To Respond labels for the briefing."""
    top: list[dict[str, Any]] = []
    for label_name in ("Priority", "To Respond"):
        if len(top) >= limit:
            break
        label_id = label_name_to_id.get(label_name)
        if not label_id:
            continue
        try:
            data = await gmail_get_json_for_user(
                user_id=user_id,
                url=f"{GMAIL_API_BASE}/threads",
                params={"labelIds": ["INBOX", label_id], "maxResults": limit},
            )
            for thread in data.get("threads", []) or []:
                if len(top) >= limit:
                    break
                thread_id = thread.get("id")
                snippet = thread.get("snippet", "")
                # Fetch first message for subject/sender
                try:
                    thr = await gmail_get_json_for_user(
                        user_id=user_id,
                        url=f"{GMAIL_API_BASE}/threads/{thread_id}",
                    )
                    msgs = thr.get("messages", []) or []
                    if not msgs:
                        continue
                    last_msg = msgs[-1]
                    headers = last_msg.get("payload", {}).get("headers", [])
                    subject = get_header_value(headers, "Subject") or "(geen onderwerp)"
                    from_header = get_header_value(headers, "From") or ""
                    top.append({
                        "label": label_name,
                        "subject": subject[:120],
                        "from": from_header[:120],
                        "snippet": snippet[:150],
                        "thread_id": thread_id,
                    })
                except Exception:
                    continue
        except Exception:
            continue
    return top[:limit]


async def get_briefing_impact_stats(
    user_id: str,
    label_name_to_id: dict[str, str],
) -> dict[str, int]:
    """Computes 'what OfficeFlow did for you' stats for yesterday + this week.

    Returns dict with:
      - drafts_yesterday, mails_organized_yesterday, commitments_yesterday, minutes_saved_yesterday
      - drafts_week, mails_organized_week, commitments_week, minutes_saved_week
    """
    stats = {
        "drafts_yesterday": 0,
        "mails_organized_yesterday": 0,
        "commitments_yesterday": 0,
        "minutes_saved_yesterday": 0,
        "drafts_week": 0,
        "mails_organized_week": 0,
        "commitments_week": 0,
        "minutes_saved_week": 0,
    }

    # Gmail: drafts created in last 24h
    try:
        data = await gmail_get_json_for_user(
            user_id=user_id,
            url=f"{GMAIL_API_BASE}/threads",
            params={"q": "in:drafts newer_than:1d", "maxResults": 100},
        )
        stats["drafts_yesterday"] = int(data.get("resultSizeEstimate", len(data.get("threads", []) or [])))
    except Exception:
        pass

    # Gmail: drafts created in last 7d
    try:
        data = await gmail_get_json_for_user(
            user_id=user_id,
            url=f"{GMAIL_API_BASE}/threads",
            params={"q": "in:drafts newer_than:7d", "maxResults": 200},
        )
        stats["drafts_week"] = int(data.get("resultSizeEstimate", len(data.get("threads", []) or [])))
    except Exception:
        pass

    # Gmail: mails organized (classified with one of our labels) in last 24h
    classified_yesterday = 0
    classified_week = 0
    for ln in ("Priority", "To Respond", "Follow Up", "Done", "FYI", "Notification", "Marketing"):
        lid = label_name_to_id.get(ln)
        if not lid:
            continue
        try:
            d1 = await gmail_get_json_for_user(
                user_id=user_id,
                url=f"{GMAIL_API_BASE}/threads",
                params={"labelIds": [lid], "q": "newer_than:1d", "maxResults": 100},
            )
            classified_yesterday += int(d1.get("resultSizeEstimate", 0))
        except Exception:
            pass
        try:
            d7 = await gmail_get_json_for_user(
                user_id=user_id,
                url=f"{GMAIL_API_BASE}/threads",
                params={"labelIds": [lid], "q": "newer_than:7d", "maxResults": 200},
            )
            classified_week += int(d7.get("resultSizeEstimate", 0))
        except Exception:
            pass
    stats["mails_organized_yesterday"] = classified_yesterday
    stats["mails_organized_week"] = classified_week

    # Supabase: commitments caught in last 24h / last 7d
    try:
        yesterday_iso = quote((_dt.now(_tz.utc) - _td(days=1)).isoformat(), safe='')
        rows = await supabase_get(
            f"/rest/v1/commitments?user_id=eq.{user_id}&created_at=gte.{yesterday_iso}&select=id",
        )
        stats["commitments_yesterday"] = len(rows or [])
    except Exception:
        pass

    try:
        week_iso = quote((_dt.now(_tz.utc) - _td(days=7)).isoformat(), safe='')
        rows = await supabase_get(
            f"/rest/v1/commitments?user_id=eq.{user_id}&created_at=gte.{week_iso}&select=id",
        )
        stats["commitments_week"] = len(rows or [])
    except Exception:
        pass

    # Time saved estimates:
    #   - AI draft ready to review: ~3 min saved per mail (don't have to compose from scratch)
    #   - Classified / organized mail: ~0.5 min saved (less mental triage)
    #   - Commitment caught: ~2 min saved (avoids forgetting + search time later)
    stats["minutes_saved_yesterday"] = int(
        stats["drafts_yesterday"] * 3
        + stats["mails_organized_yesterday"] * 0.5
        + stats["commitments_yesterday"] * 2
    )
    stats["minutes_saved_week"] = int(
        stats["drafts_week"] * 3
        + stats["mails_organized_week"] * 0.5
        + stats["commitments_week"] * 2
    )

    return stats


def build_briefing_html(
    user_first_name: str,
    counts: dict[str, int],
    commitments: list[dict[str, Any]],
    top_threads: list[dict[str, Any]],
    today_str: str,
    user_email: str = "",
    impact: dict[str, int] | None = None,
) -> tuple[str, str]:
    """Returns (subject, html_body) for the briefing email."""
    total_action = counts.get("Priority", 0) + counts.get("To Respond", 0) + counts.get("Follow Up", 0)
    impact = impact or {}
    minutes_y = int(impact.get("minutes_saved_yesterday", 0) or 0)

    # Subject: highlight time saved if meaningful, else fallback to action count
    if minutes_y >= 15:
        subject = f"OfficeFlow — {minutes_y} min bespaard · {total_action} mails vragen aandacht"
    else:
        subject = f"OfficeFlow Briefing — {today_str} · {total_action} mails vragen aandacht"

    # Gmail deep links — use /u/0/ path so URL matches existing Gmail tab
    # (query strings like ?authuser=... force a fresh page load in a new tab;
    # hash-only routes under /u/0/ navigate inside the already-open Gmail app).
    from urllib.parse import quote as _urlquote
    gmail_base = "https://mail.google.com/mail/u/0/"

    def label_url(label_name: str) -> str:
        return f"{gmail_base}#label/{_urlquote(label_name.replace(' ', '+'), safe='+')}"

    def thread_url(thread_id: str) -> str:
        return f"{gmail_base}#inbox/{thread_id}"

    def commitment_url(thread_id: str) -> str:
        return f"{gmail_base}#label/Follow+Up/{thread_id}"

    priority_url = label_url("Priority")
    to_respond_url = label_url("To Respond")
    follow_up_url = label_url("Follow Up")

    commitment_rows = ""
    if commitments:
        for c in commitments[:10]:
            due = c.get("due_date") or "geen deadline"
            action = (c.get("action_text") or "Opvolgen")[:120]
            recipient = c.get("recipient_name") or c.get("recipient_email") or "onbekend"
            subj = (c.get("subject") or "")[:100]
            cthread_id = c.get("gmail_thread_id") or ""
            clink = commitment_url(cthread_id) if cthread_id else follow_up_url
            commitment_rows += f"""
            <tr>
              <td style="padding:10px 14px;border-bottom:1px solid #f1f5f9;vertical-align:top;font-size:13px;color:#475569;white-space:nowrap;">{due}</td>
              <td style="padding:10px 14px;border-bottom:1px solid #f1f5f9;vertical-align:top;font-size:13px;color:#0f172a;">
                <a href="{clink}" target="_top" style="text-decoration:none;color:inherit;display:block;">
                  <div style="font-weight:600;color:#0f172a;">{action}</div>
                  <div style="color:#64748b;font-size:12px;margin-top:2px;">naar {recipient}{' — ' + subj if subj else ''}</div>
                </a>
              </td>
            </tr>"""
    else:
        commitment_rows = '<tr><td colspan="2" style="padding:14px;color:#94a3b8;font-size:13px;text-align:center;">Geen openstaande beloftes — alles onder controle.</td></tr>'

    thread_rows = ""
    if top_threads:
        for t in top_threads:
            tid = t.get("thread_id") or ""
            turl = thread_url(tid) if tid else label_url(t.get("label") or "Priority")
            thread_rows += f"""
            <tr>
              <td style="padding:0;border-bottom:1px solid #f1f5f9;">
                <a href="{turl}" target="_top" style="display:block;padding:10px 14px;text-decoration:none;color:inherit;">
                  <div style="display:inline-block;padding:2px 8px;border-radius:6px;background:#fef3c7;color:#92400e;font-size:11px;font-weight:700;letter-spacing:.3px;">{t.get('label','').upper()}</div>
                  <div style="font-weight:600;font-size:14px;color:#0f172a;margin-top:6px;">{t.get('subject','')}</div>
                  <div style="color:#64748b;font-size:12px;margin-top:2px;">{t.get('from','')}</div>
                  <div style="color:#475569;font-size:12px;margin-top:4px;line-height:1.4;">{t.get('snippet','')}</div>
                </a>
              </td>
            </tr>"""
    else:
        thread_rows = '<tr><td style="padding:14px;color:#94a3b8;font-size:13px;text-align:center;">Geen urgente mails — mooi begin van je dag.</td></tr>'

    # --- Derived impact values for the hero section ---
    drafts_y = int(impact.get("drafts_yesterday", 0) or 0)
    mails_y = int(impact.get("mails_organized_yesterday", 0) or 0)
    commits_y = int(impact.get("commitments_yesterday", 0) or 0)
    drafts_w = int(impact.get("drafts_week", 0) or 0)
    commits_w = int(impact.get("commitments_week", 0) or 0)
    minutes_w = int(impact.get("minutes_saved_week", 0) or 0)

    # Human-friendly "time saved" string
    def _fmt_minutes(m: int) -> str:
        if m <= 0:
            return "0 min"
        if m < 60:
            return f"{m} min"
        hrs = m // 60
        mins = m % 60
        return f"{hrs}u {mins:02d}m" if mins else f"{hrs} uur"

    time_saved_str = _fmt_minutes(minutes_y)

    # Headline copy adapts to volume
    if minutes_y >= 30:
        hero_headline = f"OfficeFlow heeft {time_saved_str} voor je bespaard."
        hero_sub = f"{drafts_y} drafts klaargezet · {mails_y} mails geordend · {commits_y} beloftes opgevangen."
    elif drafts_y + commits_y + mails_y > 0:
        hero_headline = "Rustige 24 uur — alles netjes verwerkt."
        hero_sub = f"{drafts_y} drafts · {mails_y} mails geordend · {commits_y} beloftes opgevangen."
    else:
        hero_headline = "Je inbox is klaar voor vandaag."
        hero_sub = "Geen nieuwe verwerkte mails in de afgelopen 24 uur."

    # Week momentum copy
    week_line = (
        f"<strong style=\"color:#0f172a;\">{drafts_w}</strong> drafts · "
        f"<strong style=\"color:#0f172a;\">{commits_w}</strong> beloftes · "
        f"<strong style=\"color:#16a34a;\">{_fmt_minutes(minutes_w)} bespaard</strong>"
    )

    html = f"""<!DOCTYPE html>
<html>
<body style="margin:0;padding:0;background:#f8fafc;font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Roboto,sans-serif;">
  <table width="100%" cellpadding="0" cellspacing="0" style="background:#f8fafc;padding:32px 12px;">
    <tr><td align="center">
      <table width="600" cellpadding="0" cellspacing="0" style="max-width:600px;background:#ffffff;border-radius:14px;overflow:hidden;box-shadow:0 2px 16px rgba(15,23,42,.04);">

        <tr><td style="padding:28px 32px 20px 32px;border-bottom:1px solid #f1f5f9;">
          <div style="font-size:12px;color:#f97316;font-weight:700;letter-spacing:1.5px;">OFFICEFLOW · BRIEFING</div>
          <div style="font-size:22px;color:#0f172a;font-weight:700;margin-top:6px;">Goedemorgen{', ' + user_first_name if user_first_name else ''}.</div>
          <div style="font-size:14px;color:#64748b;margin-top:4px;">{today_str}</div>
        </td></tr>

        <tr><td style="padding:20px 32px 4px 32px;">
          <table width="100%" cellpadding="0" cellspacing="0" style="border-radius:12px;overflow:hidden;">
            <tr><td style="padding:22px 24px;background:#0f172a;background-image:linear-gradient(135deg,#0f172a 0%,#1e293b 55%,#312e81 100%);color:#ffffff;">
              <div style="font-size:11px;color:#fb923c;font-weight:700;letter-spacing:1.5px;">TERWIJL JIJ WEG WAS</div>
              <div style="font-size:22px;font-weight:800;color:#ffffff;margin-top:6px;line-height:1.25;">{hero_headline}</div>
              <div style="font-size:13px;color:#cbd5e1;margin-top:6px;line-height:1.5;">{hero_sub}</div>
              <table width="100%" cellpadding="0" cellspacing="0" style="margin-top:16px;">
                <tr>
                  <td width="33%" style="padding:0 4px 0 0;">
                    <div style="background:rgba(255,255,255,.07);border:1px solid rgba(255,255,255,.06);border-radius:8px;padding:12px;text-align:center;">
                      <div style="font-size:22px;font-weight:800;color:#ffffff;line-height:1;">{drafts_y}</div>
                      <div style="font-size:10px;color:#cbd5e1;font-weight:600;letter-spacing:.4px;margin-top:4px;">AI DRAFTS</div>
                    </div>
                  </td>
                  <td width="33%" style="padding:0 2px;">
                    <div style="background:rgba(255,255,255,.07);border:1px solid rgba(255,255,255,.06);border-radius:8px;padding:12px;text-align:center;">
                      <div style="font-size:22px;font-weight:800;color:#ffffff;line-height:1;">{mails_y}</div>
                      <div style="font-size:10px;color:#cbd5e1;font-weight:600;letter-spacing:.4px;margin-top:4px;">MAILS GEORDEND</div>
                    </div>
                  </td>
                  <td width="33%" style="padding:0 0 0 4px;">
                    <div style="background:rgba(255,255,255,.07);border:1px solid rgba(255,255,255,.06);border-radius:8px;padding:12px;text-align:center;">
                      <div style="font-size:22px;font-weight:800;color:#ffffff;line-height:1;">{commits_y}</div>
                      <div style="font-size:10px;color:#cbd5e1;font-weight:600;letter-spacing:.4px;margin-top:4px;">BELOFTES</div>
                    </div>
                  </td>
                </tr>
              </table>
            </td></tr>
          </table>
        </td></tr>

        <tr><td style="padding:24px 32px 8px 32px;">
          <div style="font-size:13px;color:#0f172a;font-weight:700;letter-spacing:.3px;text-transform:uppercase;margin-bottom:12px;">Wat ligt er op je bord?</div>
          <table width="100%" cellpadding="0" cellspacing="0">
            <tr>
              <td width="33%" style="padding:0 6px 0 0;">
                <a href="{priority_url}" target="_top" style="text-decoration:none;display:block;">
                  <div style="background:#fef2f2;border:1px solid #fecaca;border-radius:10px;padding:14px;text-align:center;">
                    <div style="font-size:28px;font-weight:800;color:#b91c1c;">{counts.get('Priority',0)}</div>
                    <div style="font-size:11px;color:#991b1b;font-weight:600;letter-spacing:.3px;margin-top:2px;">PRIORITY</div>
                  </div>
                </a>
              </td>
              <td width="33%" style="padding:0 3px;">
                <a href="{to_respond_url}" target="_top" style="text-decoration:none;display:block;">
                  <div style="background:#eff6ff;border:1px solid #bfdbfe;border-radius:10px;padding:14px;text-align:center;">
                    <div style="font-size:28px;font-weight:800;color:#1d4ed8;">{counts.get('To Respond',0)}</div>
                    <div style="font-size:11px;color:#1e40af;font-weight:600;letter-spacing:.3px;margin-top:2px;">TO RESPOND</div>
                  </div>
                </a>
              </td>
              <td width="33%" style="padding:0 0 0 6px;">
                <a href="{follow_up_url}" target="_top" style="text-decoration:none;display:block;">
                  <div style="background:#fff7ed;border:1px solid #fed7aa;border-radius:10px;padding:14px;text-align:center;">
                    <div style="font-size:28px;font-weight:800;color:#c2410c;">{counts.get('Follow Up',0)}</div>
                    <div style="font-size:11px;color:#9a3412;font-weight:600;letter-spacing:.3px;margin-top:2px;">FOLLOW UP</div>
                  </div>
                </a>
              </td>
            </tr>
          </table>
        </td></tr>

        <tr><td style="padding:24px 32px 8px 32px;">
          <div style="font-size:13px;color:#0f172a;font-weight:700;letter-spacing:.3px;text-transform:uppercase;margin-bottom:12px;">Top 3 vandaag</div>
          <table width="100%" cellpadding="0" cellspacing="0" style="border:1px solid #e2e8f0;border-radius:10px;overflow:hidden;">
            {thread_rows}
          </table>
        </td></tr>

        <tr><td style="padding:20px 32px 8px 32px;">
          <div style="font-size:13px;color:#0f172a;font-weight:700;letter-spacing:.3px;text-transform:uppercase;margin-bottom:12px;">Jouw openstaande beloftes</div>
          <table width="100%" cellpadding="0" cellspacing="0" style="border:1px solid #e2e8f0;border-radius:10px;overflow:hidden;">
            <thead><tr>
              <th align="left" style="background:#f8fafc;padding:8px 14px;font-size:11px;color:#64748b;font-weight:700;letter-spacing:.3px;border-bottom:1px solid #e2e8f0;">DEADLINE</th>
              <th align="left" style="background:#f8fafc;padding:8px 14px;font-size:11px;color:#64748b;font-weight:700;letter-spacing:.3px;border-bottom:1px solid #e2e8f0;">ACTIE</th>
            </tr></thead>
            <tbody>
              {commitment_rows}
            </tbody>
          </table>
        </td></tr>

        <tr><td style="padding:12px 32px 24px 32px;">
          <table width="100%" cellpadding="0" cellspacing="0" style="background:#f8fafc;border-radius:10px;border:1px solid #e2e8f0;">
            <tr>
              <td style="padding:14px 18px;">
                <div style="font-size:11px;color:#64748b;font-weight:700;letter-spacing:.4px;text-transform:uppercase;">Deze week</div>
                <div style="font-size:14px;color:#0f172a;margin-top:4px;line-height:1.5;">{week_line}</div>
              </td>
            </tr>
          </table>
        </td></tr>

        <tr><td style="padding:16px 32px 24px 32px;background:#fafbfc;border-top:1px solid #f1f5f9;">
          <div style="font-size:11px;color:#94a3b8;text-align:center;line-height:1.6;">
            Deze briefing wordt elke ochtend door OfficeFlow voor je samengesteld.<br>
            Geen auto-send. Geen auto-archive. Jij blijft altijd in controle.
          </div>
        </td></tr>

      </table>
    </td></tr>
  </table>
</body>
</html>"""
    return subject, html


def build_briefing_raw_email(to_email: str, subject: str, html_body: str) -> str:
    """Builds a base64url-encoded raw MIME email for Gmail API send."""
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    msg = MIMEMultipart("alternative")
    msg["To"] = to_email
    msg["From"] = to_email
    msg["Subject"] = subject
    plain = "Bekijk deze briefing in een HTML-capable client."
    msg.attach(MIMEText(plain, "plain", "utf-8"))
    msg.attach(MIMEText(html_body, "html", "utf-8"))
    raw = base64.urlsafe_b64encode(msg.as_bytes()).decode("utf-8").rstrip("=")
    return raw


async def send_briefing_for_mailbox(email: str) -> dict[str, Any]:
    """Builds and sends the daily briefing to the given mailbox."""
    context = await get_gmail_context_by_email(email)
    user = context["user"]
    mailbox = context["mailbox"]

    label_map = await get_all_gmail_labels(user["id"])
    label_name_to_id = {name: label["id"] for name, label in label_map.items()}

    counts = await count_labels_for_briefing(user["id"], label_name_to_id)
    commitments = await get_upcoming_commitments(user["id"], days_ahead=7)
    top_threads = await get_top_priority_threads(user["id"], label_name_to_id, limit=3)
    impact = await get_briefing_impact_stats(user["id"], label_name_to_id)

    user_first_name = (user.get("full_name") or user.get("email") or "").split(" ")[0].split("@")[0]
    today_str = _dt.now(ZoneInfo(BRIEFING_DEFAULT_TIMEZONE)).strftime("%A %d %B %Y").capitalize()

    subject, html_body = build_briefing_html(
        user_first_name=user_first_name,
        counts=counts,
        commitments=commitments,
        top_threads=top_threads,
        today_str=today_str,
        user_email=email,
        impact=impact,
    )

    raw = build_briefing_raw_email(to_email=email, subject=subject, html_body=html_body)

    send_response = await gmail_post_json_for_user(
        user_id=user["id"],
        url=f"{GMAIL_API_BASE}/messages/send",
        payload={"raw": raw},
    )

    # Update briefing_last_sent_at
    try:
        await supabase_patch(
            f"/rest/v1/mailboxes?id=eq.{mailbox['id']}",
            {"briefing_last_sent_at": _dt.now(_tz.utc).isoformat()},
        )
    except Exception:
        pass

    return {
        "status": "sent",
        "email": email,
        "subject": subject,
        "gmail_message_id": send_response.get("id"),
        "counts": counts,
        "commitments_count": len(commitments),
        "top_threads_count": len(top_threads),
    }


async def briefing_loop():
    """Runs every 2 min. Sends briefing to each mailbox at its configured time (once per day)."""
    await asyncio.sleep(30)
    while True:
        try:
            if not BRIEFING_ENABLED_GLOBAL:
                await asyncio.sleep(120)
                continue

            mailboxes = await get_all_active_mailboxes()
            for mailbox in mailboxes:
                try:
                    email = mailbox.get("email_address")
                    if not email:
                        continue
                    if not mailbox.get("briefing_enabled", True):
                        continue

                    tz_name = mailbox.get("briefing_timezone") or BRIEFING_DEFAULT_TIMEZONE
                    hour = int(mailbox.get("briefing_hour") or 7)
                    minute = int(mailbox.get("briefing_minute") or 30)

                    try:
                        now_local = _dt.now(ZoneInfo(tz_name))
                    except Exception:
                        now_local = _dt.now(ZoneInfo(BRIEFING_DEFAULT_TIMEZONE))

                    last_sent_raw = mailbox.get("briefing_last_sent_at")
                    already_today = False
                    if last_sent_raw:
                        try:
                            last_sent = _dt.fromisoformat(last_sent_raw.replace("Z", "+00:00"))
                            if last_sent.astimezone(now_local.tzinfo).date() == now_local.date():
                                already_today = True
                        except Exception:
                            pass

                    if already_today:
                        continue

                    # Only send within a 5-min window after target time
                    target = now_local.replace(hour=hour, minute=minute, second=0, microsecond=0)
                    if now_local < target:
                        continue
                    if (now_local - target) > _td(minutes=15):
                        # Window missed for today; will send tomorrow
                        continue

                    print(f"[briefing] Sending to {email}...")
                    result = await send_briefing_for_mailbox(email)
                    print(f"[briefing] Sent to {email}: {result.get('gmail_message_id')}")
                except Exception as exc:
                    print(f"[briefing] Failed for {mailbox.get('email_address')}: {repr(exc)}")
        except Exception as exc:
            print(f"[briefing] Loop error: {repr(exc)}")

        await asyncio.sleep(120)  # check every 2 min


# ---------------------------------------------------------------------------
# SECTION D — RELATIONSHIP INTELLIGENCE
# ---------------------------------------------------------------------------

async def get_relationship_for_contact(user_id: str, contact_email: str) -> dict[str, Any]:
    """
    Returns aggregated intel for a given contact email:
    - total threads with this contact
    - breakdown by label (Priority, To Respond, Waiting On Reply, Follow Up, Done)
    - last contact date
    - active commitments to this contact
    """
    contact_email = (contact_email or "").strip().lower()
    if not contact_email:
        raise HTTPException(status_code=400, detail="contact_email required")

    # Query Gmail for threads involving this contact
    query = f"(from:{contact_email} OR to:{contact_email})"
    data = await gmail_get_json_for_user(
        user_id=user_id,
        url=f"{GMAIL_API_BASE}/threads",
        params={"q": query, "maxResults": 50},
    )
    thread_stubs = data.get("threads", []) or []

    label_map = await get_all_gmail_labels(user_id)
    id_to_name = {label["id"]: name for name, label in label_map.items()}
    tracked_labels = {"Priority", "To Respond", "Waiting On Reply", "Follow Up", "Done", "FYI", "Notification", "Marketing", "Ignore"}

    label_breakdown: dict[str, int] = {name: 0 for name in tracked_labels}
    last_contact_ts: int | None = None
    last_contact_subject: str | None = None
    last_contact_direction: str | None = None  # "incoming" or "outgoing"
    thread_summaries: list[dict[str, Any]] = []

    for stub in thread_stubs[:30]:  # cap for performance
        thread_id = stub.get("id")
        if not thread_id:
            continue
        try:
            thr = await gmail_get_json_for_user(
                user_id=user_id,
                url=f"{GMAIL_API_BASE}/threads/{thread_id}",
            )
        except Exception:
            continue

        messages = thr.get("messages", []) or []
        if not messages:
            continue

        last_msg = messages[-1]
        internal_date = int(last_msg.get("internalDate", "0"))
        last_msg_headers = last_msg.get("payload", {}).get("headers", [])
        subject = get_header_value(last_msg_headers, "Subject") or "(geen onderwerp)"
        from_header = get_header_value(last_msg_headers, "From") or ""
        is_outgoing = contact_email not in from_header.lower()

        # Aggregate labels from all messages in thread
        thread_label_names: set[str] = set()
        for m in messages:
            for lid in m.get("labelIds", []) or []:
                name = id_to_name.get(lid)
                if name in tracked_labels:
                    thread_label_names.add(name)

        for name in thread_label_names:
            label_breakdown[name] = label_breakdown.get(name, 0) + 1

        if last_contact_ts is None or internal_date > last_contact_ts:
            last_contact_ts = internal_date
            last_contact_subject = subject
            last_contact_direction = "outgoing" if is_outgoing else "incoming"

        thread_summaries.append({
            "thread_id": thread_id,
            "subject": subject[:120],
            "last_ts": internal_date,
            "labels": sorted(thread_label_names),
            "direction": "outgoing" if is_outgoing else "incoming",
        })

    thread_summaries.sort(key=lambda t: t["last_ts"], reverse=True)

    # Commitments for this contact
    try:
        commitments = await supabase_get(
            f"/rest/v1/commitments?user_id=eq.{user_id}&recipient_email=eq.{contact_email}&status=eq.active&order=due_date.asc&select=*",
        )
    except Exception:
        commitments = []

    last_contact_iso = None
    if last_contact_ts:
        try:
            last_contact_iso = _dt.fromtimestamp(last_contact_ts / 1000, tz=_tz.utc).isoformat()
        except Exception:
            pass

    return {
        "contact_email": contact_email,
        "total_threads": len(thread_stubs),
        "shown_threads": len(thread_summaries),
        "label_breakdown": label_breakdown,
        "last_contact_at": last_contact_iso,
        "last_contact_subject": last_contact_subject,
        "last_contact_direction": last_contact_direction,
        "active_commitments": commitments or [],
        "recent_threads": thread_summaries[:15],
    }


# ---------------------------------------------------------------------------
# SECTION E — ENDPOINTS
# ---------------------------------------------------------------------------

@app.post("/internal/send-briefing")
async def http_send_briefing(
    email: str = Body(..., embed=True),
    _: None = Depends(require_internal_secret),
):
    """Manual trigger: sends morning briefing immediately to the given mailbox."""
    return await send_briefing_for_mailbox(email)


@app.post("/internal/detect-commitment")
async def http_detect_commitment(
    email: str = Body(...),
    subject: str | None = Body(default=None),
    body_text: str = Body(...),
    _: None = Depends(require_internal_secret),
):
    """Manual trigger: runs commitment detection on arbitrary text (for testing)."""
    return await detect_commitments_in_sent_mail(subject=subject, body_text=body_text)


@app.get("/commitments")
async def http_list_commitments(email: str, status: str = "active"):
    """Lists commitments for a mailbox. Filter by status (active / completed / cancelled)."""
    context = await get_gmail_context_by_email(email)
    user = context["user"]
    if status not in ("active", "completed", "cancelled"):
        raise HTTPException(status_code=400, detail="status must be active / completed / cancelled")
    rows = await supabase_get(
        f"/rest/v1/commitments?user_id=eq.{user['id']}&status=eq.{status}&order=due_date.asc&select=*",
    )
    return {"commitments": rows or []}


@app.post("/commitments/{commitment_id}/complete")
async def http_complete_commitment(commitment_id: str, email: str = Body(..., embed=True)):
    context = await get_gmail_context_by_email(email)
    user = context["user"]
    now_iso = _dt.now(_tz.utc).isoformat()
    updated = await supabase_patch(
        f"/rest/v1/commitments?id=eq.{commitment_id}&user_id=eq.{user['id']}",
        {"status": "completed", "completed_at": now_iso, "updated_at": now_iso},
        prefer="return=representation",
    )
    return {"status": "ok", "commitment": updated[0] if isinstance(updated, list) and updated else None}


@app.get("/relationships/{contact_email}")
async def http_relationship(contact_email: str, email: str):
    """Relationship Intelligence: aggregated stats for a contact. `email` = user's mailbox."""
    context = await get_gmail_context_by_email(email)
    user = context["user"]
    return await get_relationship_for_contact(user_id=user["id"], contact_email=contact_email)


@app.post("/briefing/settings")
async def http_briefing_settings(
    email: str = Body(...),
    enabled: bool | None = Body(default=None),
    hour: int | None = Body(default=None),
    minute: int | None = Body(default=None),
    timezone_name: str | None = Body(default=None),
):
    """Update briefing preferences for a mailbox."""
    context = await get_gmail_context_by_email(email)
    mailbox = context["mailbox"]
    patch: dict[str, Any] = {}
    if enabled is not None:
        patch["briefing_enabled"] = enabled
    if hour is not None:
        patch["briefing_hour"] = max(0, min(23, hour))
    if minute is not None:
        patch["briefing_minute"] = max(0, min(59, minute))
    if timezone_name:
        patch["briefing_timezone"] = timezone_name
    if not patch:
        return {"status": "noop"}
    updated = await supabase_patch(
        f"/rest/v1/mailboxes?id=eq.{mailbox['id']}",
        patch,
        prefer="return=representation",
    )
    return {"status": "ok", "mailbox": updated[0] if isinstance(updated, list) and updated else None}


# ---------------------------------------------------------------------------
# SECTION G — FOLLOW-UP RADAR: auto-nudge on overdue commitments
# ---------------------------------------------------------------------------

FOLLOW_UP_RADAR_ENABLED_GLOBAL = os.getenv("FOLLOW_UP_RADAR_ENABLED", "true").lower() == "true"
FOLLOW_UP_RADAR_MAX_NUDGES_PER_RUN = int(os.getenv("FOLLOW_UP_RADAR_MAX_NUDGES_PER_RUN", "10"))
FOLLOW_UP_RADAR_RUN_HOUR = int(os.getenv("FOLLOW_UP_RADAR_RUN_HOUR", "8"))  # local time


async def get_mailbox_radar_config(mailbox_id: str) -> dict[str, Any]:
    try:
        rows = await supabase_get(
            f"/rest/v1/mailboxes?id=eq.{mailbox_id}"
            f"&select=radar_enabled,radar_grace_days,radar_last_run_at,"
            f"briefing_timezone,email_address",
        )
        return rows[0] if rows else {}
    except Exception as exc:
        print(f"[radar] get config failed for mailbox {mailbox_id}: {repr(exc)}")
        return {}


async def supabase_get_overdue_commitments_for_mailbox(
    mailbox_id: str,
    grace_days: int = 0,
) -> list[dict[str, Any]]:
    """
    Returns active commitments whose due_date + grace_days is in the past
    AND that have not yet been nudged AND are not suppressed.
    """
    try:
        cutoff = (_date.today() - _td(days=grace_days)).isoformat()
        rows = await supabase_get(
            f"/rest/v1/commitments"
            f"?mailbox_id=eq.{mailbox_id}"
            f"&status=eq.active"
            f"&nudge_suppressed=eq.false"
            f"&nudge_sent_at=is.null"
            f"&due_date=lte.{cutoff}"
            f"&order=due_date.asc"
            f"&select=*"
        )
        return rows if isinstance(rows, list) else []
    except Exception as exc:
        print(f"[radar] fetch overdue commitments failed for mailbox {mailbox_id}: {repr(exc)}")
        return []


async def supabase_mark_commitment_nudged(
    commitment_id: str,
    gmail_draft_id: str | None,
) -> None:
    try:
        now_iso = _dt.now(_tz.utc).isoformat()
        await supabase_patch(
            f"/rest/v1/commitments?id=eq.{commitment_id}",
            {
                "nudge_sent_at": now_iso,
                "nudge_count": 1,
                "nudge_draft_id": gmail_draft_id,
                "updated_at": now_iso,
            },
        )
    except Exception as exc:
        print(f"[radar] mark-nudged failed for commitment {commitment_id}: {repr(exc)}")


async def supabase_mark_radar_run(mailbox_id: str) -> None:
    try:
        await supabase_patch(
            f"/rest/v1/mailboxes?id=eq.{mailbox_id}",
            {"radar_last_run_at": _dt.now(_tz.utc).isoformat()},
        )
    except Exception as exc:
        print(f"[radar] mark-run failed for mailbox {mailbox_id}: {repr(exc)}")


async def mark_commitments_completed_on_reply(gmail_thread_id: str) -> None:
    """Convenience wrapper used by inbound-processing when a reply arrives."""
    if not gmail_thread_id:
        return
    try:
        await supabase_complete_commitments_for_thread(gmail_thread_id)
    except Exception as exc:
        print(f"[radar] complete-on-reply failed for thread {gmail_thread_id}: {repr(exc)}")


async def generate_nudge_draft(
    user_id: str,
    commitment: dict[str, Any],
) -> str:
    """
    Produces a short, warm nudge body. Reuses user's style profile + settings
    via generate_ai_reply so the tone matches their normal voice.
    """
    recipient_name = (commitment.get("recipient_name") or "").strip()
    recipient_email = (commitment.get("recipient_email") or "").strip()
    action_text = (commitment.get("action_text") or "opvolgen").strip()
    due_date = commitment.get("due_date")
    subject = commitment.get("subject") or "Opvolging"

    salutation_target = recipient_name or recipient_email or "daar"

    pseudo_body = (
        f"[INTERN NUDGE-CONTEXT — dit is GEEN echte inkomende e-mail]\n"
        f"De gebruiker heeft eerder beloofd: '{action_text}'.\n"
        f"Deadline was: {due_date or 'niet exact gespecificeerd'}.\n"
        f"Er is nog geen opvolging verstuurd.\n\n"
        f"Schrijf een korte, natuurlijke follow-up naar {salutation_target}.\n"
        f"Richtlijnen voor deze specifieke draft:\n"
        f"- Verontschuldig je niet overdreven; wees to-the-point.\n"
        f"- Bevestig de belofte kort en geef een realistische nieuwe status of vraag.\n"
        f"- Geen nieuwe harde deadlines verzinnen; vraag eventueel om een kort uitstel als dat past.\n"
        f"- Max 4–6 zinnen.\n"
        f"- Geen onderwerpregel."
    )

    try:
        reply = await generate_ai_reply(
            user_id=user_id,
            subject=subject,
            sender=recipient_email or recipient_name or None,
            body_text=pseudo_body,
        )
        return reply or ""
    except Exception as exc:
        print(f"[radar] generate_nudge_draft failed: {repr(exc)}")
        fallback = (
            f"Hoi {recipient_name or ''},\n\n"
            f"Ik kom even terug op {action_text.lower()}. Ik wil je hierover nog graag bijpraten — "
            f"zou jij kunnen laten weten wat een handig moment is?\n\n"
            f"Groet"
        )
        return fallback.strip()


async def get_thread_message_headers(user_id: str, thread_id: str) -> dict[str, Any]:
    """
    Fetches the latest message in a thread and returns headers we need to
    build a properly threaded reply (Message-ID + References).
    """
    try:
        thread_data = await gmail_get_json_for_user(
            user_id=user_id,
            url=f"{GMAIL_API_BASE}/threads/{thread_id}",
            params={"format": "metadata", "metadataHeaders": ["Message-ID", "References", "Subject", "From", "To"]},
        )
        messages = thread_data.get("messages", []) or []
        if not messages:
            return {}
        last = messages[-1]
        headers = last.get("payload", {}).get("headers", []) or []
        return {
            "message_id_header": get_header_value(headers, "Message-ID"),
            "references_header": get_header_value(headers, "References"),
            "subject": get_header_value(headers, "Subject"),
            "gmail_message_id": last.get("id"),
            "current_label_ids": set(last.get("labelIds", []) or []),
        }
    except Exception as exc:
        print(f"[radar] get_thread_message_headers failed for thread {thread_id}: {repr(exc)}")
        return {}


async def process_follow_up_radar_for_mailbox(email: str) -> dict[str, Any]:
    """
    Scans overdue commitments for this mailbox. For each:
      1. Re-verifies the thread still needs a nudge.
      2. Escalates thread label -> Priority.
      3. Generates a threaded Gmail draft (no auto-send).
      4. Persists nudge_sent_at on the commitment.
    Idempotent — will not re-nudge a commitment that already has nudge_sent_at.
    """
    if not FOLLOW_UP_RADAR_ENABLED_GLOBAL:
        return {"status": "disabled_global", "nudges_created": 0}

    user = await supabase_get_user_by_email(email)
    if not user:
        return {"status": "no_user", "nudges_created": 0}

    mailbox = await supabase_get_mailbox_by_user_and_email(user["id"], email)
    if not mailbox:
        return {"status": "no_mailbox", "nudges_created": 0}

    config = await get_mailbox_radar_config(mailbox["id"])
    if not config.get("radar_enabled", True):
        return {"status": "disabled_for_mailbox", "nudges_created": 0}

    grace_days = int(config.get("radar_grace_days") or 0)
    overdue = await supabase_get_overdue_commitments_for_mailbox(
        mailbox_id=mailbox["id"], grace_days=grace_days,
    )
    if not overdue:
        await supabase_mark_radar_run(mailbox["id"])
        return {"status": "no_overdue", "nudges_created": 0}

    overdue = overdue[:FOLLOW_UP_RADAR_MAX_NUDGES_PER_RUN]

    try:
        labels = await get_all_gmail_labels(user["id"])
        label_name_to_id = {
            name: lbl["id"] for name, lbl in labels.items() if isinstance(lbl, dict) and "id" in lbl
        }
    except Exception as exc:
        print(f"[radar] label fetch failed for {email}: {repr(exc)}")
        label_name_to_id = {}

    nudges_created = 0
    errors: list[str] = []

    for commitment in overdue:
        try:
            commitment_id = commitment.get("id")
            thread_id = commitment.get("gmail_thread_id")
            recipient_email = commitment.get("recipient_email")

            if not commitment_id or not thread_id or not recipient_email:
                errors.append(f"commitment {commitment_id}: missing thread_id or recipient_email")
                continue

            thread_info = await get_thread_message_headers(user["id"], thread_id)
            if not thread_info:
                errors.append(f"commitment {commitment_id}: thread fetch failed")
                continue

            if "Priority" in label_name_to_id:
                try:
                    await sync_thread_status(
                        user_id=user["id"],
                        thread_id=thread_id,
                        current_message_id=thread_info.get("gmail_message_id") or "",
                        current_label_ids=thread_info.get("current_label_ids") or set(),
                        label_name_to_id=label_name_to_id,
                        target_label_name="Priority",
                    )
                except Exception as exc:
                    print(f"[radar] escalate label failed for commitment {commitment_id}: {repr(exc)}")

            nudge_body = await generate_nudge_draft(user_id=user["id"], commitment=commitment)
            if not nudge_body or len(nudge_body.strip()) < 10:
                errors.append(f"commitment {commitment_id}: empty nudge body")
                continue

            draft_result = await create_gmail_threaded_draft(
                user_id=user["id"],
                to_email=recipient_email,
                subject=thread_info.get("subject") or commitment.get("subject") or "Opvolging",
                body=nudge_body,
                thread_id=thread_id,
                original_message_id_header=thread_info.get("message_id_header"),
                references_header=thread_info.get("references_header"),
            )
            gmail_draft_id = (draft_result or {}).get("id")

            await supabase_mark_commitment_nudged(
                commitment_id=commitment_id,
                gmail_draft_id=gmail_draft_id,
            )
            nudges_created += 1
        except Exception as exc:
            errors.append(f"commitment {commitment.get('id')}: {repr(exc)}")
            print(f"[radar] per-commitment failure: {repr(exc)}")

    await supabase_mark_radar_run(mailbox["id"])

    return {
        "status": "ok",
        "email": email,
        "candidates": len(overdue),
        "nudges_created": nudges_created,
        "errors": errors,
    }


async def radar_loop():
    """
    Background loop: checks every 5 min. For each mailbox whose local time has
    crossed FOLLOW_UP_RADAR_RUN_HOUR and that has not run today, triggers the radar.
    """
    await asyncio.sleep(45)
    while True:
        try:
            if not FOLLOW_UP_RADAR_ENABLED_GLOBAL:
                await asyncio.sleep(300)
                continue

            mailboxes = await get_all_active_mailboxes()
            for mailbox in mailboxes:
                try:
                    email = mailbox.get("email_address")
                    if not email:
                        continue
                    if not mailbox.get("radar_enabled", True):
                        continue

                    tz_name = mailbox.get("briefing_timezone") or BRIEFING_DEFAULT_TIMEZONE
                    try:
                        now_local = _dt.now(ZoneInfo(tz_name))
                    except Exception:
                        now_local = _dt.now(ZoneInfo(BRIEFING_DEFAULT_TIMEZONE))

                    if now_local.hour < FOLLOW_UP_RADAR_RUN_HOUR:
                        continue

                    last_run_raw = mailbox.get("radar_last_run_at")
                    already_today = False
                    if last_run_raw:
                        try:
                            last_run = _dt.fromisoformat(last_run_raw.replace("Z", "+00:00"))
                            if last_run.astimezone(now_local.tzinfo).date() == now_local.date():
                                already_today = True
                        except Exception:
                            pass
                    if already_today:
                        continue

                    print(f"[radar] Running for {email}...")
                    result = await process_follow_up_radar_for_mailbox(email)
                    print(f"[radar] {email}: {result}")
                    try:
                        result_s = await process_silence_radar_for_mailbox(email)
                        print(f"[radar-silence] {email}: {result_s}")
                    except Exception as exc_s:
                        print(f"[radar-silence] loop-call failed for {email}: {repr(exc_s)}")
                except Exception as exc:
                    print(f"[radar] mailbox loop failure for {mailbox.get('email_address')}: {repr(exc)}")
        except Exception as exc:
            print(f"[radar] loop error: {repr(exc)}")

        await asyncio.sleep(300)  # 5 min cadence


# ---------------------------------------------------------------------------
# Internal / testing endpoints for the radar
# ---------------------------------------------------------------------------

@app.post("/internal/follow-up-radar/run")
async def http_run_follow_up_radar(
    email: str = Body(..., embed=True),
    _: None = Depends(require_internal_secret),
):
    """Manually trigger the radar for one mailbox. Use for QA/testing."""
    return await process_follow_up_radar_for_mailbox(email)


@app.post("/internal/follow-up-radar/suppress")
async def http_suppress_commitment_nudge(
    commitment_id: str = Body(..., embed=True),
    _: None = Depends(require_internal_secret),
):
    """User can suppress nudging for a specific commitment (future UI hook)."""
    try:
        updated = await supabase_patch(
            f"/rest/v1/commitments?id=eq.{commitment_id}",
            {"nudge_suppressed": True, "updated_at": _dt.now(_tz.utc).isoformat()},
            prefer="return=representation",
        )
        return {"status": "ok", "commitment": updated[0] if isinstance(updated, list) and updated else None}
    except Exception as exc:
        return {"status": "error", "error": repr(exc)}


@app.post("/internal/follow-up-radar/reset")
async def http_reset_commitment_nudge(
    commitment_id: str = Body(..., embed=True),
    _: None = Depends(require_internal_secret),
):
    """Clear nudge state so the commitment can be nudged again on next run."""
    try:
        updated = await supabase_patch(
            f"/rest/v1/commitments?id=eq.{commitment_id}",
            {
                "nudge_sent_at": None,
                "nudge_count": 0,
                "nudge_draft_id": None,
                "nudge_suppressed": False,
                "updated_at": _dt.now(_tz.utc).isoformat(),
            },
            prefer="return=representation",
        )
        return {"status": "ok", "commitment": updated[0] if isinstance(updated, list) and updated else None}
    except Exception as exc:
        return {"status": "error", "error": repr(exc)}


# ---------------------------------------------------------------------------
# SECTION H — SILENCE RADAR: nudge on "Waiting On Reply" threads gone quiet
# ---------------------------------------------------------------------------

SILENCE_RADAR_ENABLED_GLOBAL = os.getenv("SILENCE_RADAR_ENABLED", "true").lower() == "true"
SILENCE_RADAR_MAX_NUDGES_PER_RUN = int(os.getenv("SILENCE_RADAR_MAX_NUDGES_PER_RUN", "10"))
SILENCE_RADAR_DEFAULT_THRESHOLD_DAYS = int(os.getenv("SILENCE_RADAR_DEFAULT_THRESHOLD_DAYS", "5"))


async def supabase_upsert_awaiting_reply(
    user_id: str,
    mailbox_id: str,
    gmail_thread_id: str,
    last_user_message_id: str | None,
    last_user_sent_at: str,
    recipient_email: str | None,
    recipient_name: str | None,
    subject: str | None,
) -> None:
    """
    Upsert awaiting_reply on (mailbox_id, gmail_thread_id).
    A new sent message on the same thread RESETS the silence clock:
      - last_user_sent_at moves forward
      - nudge_sent_at/nudge_count reset so we nudge again if silence continues
      - nudge_suppressed is NOT reset (it's omitted from payload → preserved)
    """
    if not (user_id and mailbox_id and gmail_thread_id and last_user_sent_at):
        return
    try:
        now_iso = _dt.now(_tz.utc).isoformat()
        payload = {
            "user_id": user_id,
            "mailbox_id": mailbox_id,
            "gmail_thread_id": gmail_thread_id,
            "last_user_message_id": last_user_message_id,
            "last_user_sent_at": last_user_sent_at,
            "recipient_email": recipient_email,
            "recipient_name": recipient_name,
            "subject": subject,
            "status": "active",
            "replied_at": None,
            "nudge_sent_at": None,
            "nudge_count": 0,
            "nudge_draft_id": None,
            "updated_at": now_iso,
        }
        await supabase_post(
            "/rest/v1/awaiting_replies?on_conflict=mailbox_id,gmail_thread_id",
            payload,
            prefer="resolution=merge-duplicates,return=representation",
        )
    except Exception as exc:
        print(f"[silence] upsert awaiting_reply failed for thread {gmail_thread_id}: {repr(exc)}")


async def supabase_cancel_awaiting_reply_for_thread(gmail_thread_id: str) -> None:
    """Mark any active awaiting_reply for this thread as cancelled (label moved away from WOR)."""
    if not gmail_thread_id:
        return
    try:
        now_iso = _dt.now(_tz.utc).isoformat()
        await supabase_patch(
            f"/rest/v1/awaiting_replies?gmail_thread_id=eq.{gmail_thread_id}&status=eq.active",
            {"status": "cancelled", "updated_at": now_iso},
        )
    except Exception as exc:
        print(f"[silence] cancel awaiting_reply failed for thread {gmail_thread_id}: {repr(exc)}")


async def supabase_mark_awaiting_reply_replied(gmail_thread_id: str) -> None:
    """Mark awaiting_reply as replied once an external reply arrives on the thread."""
    if not gmail_thread_id:
        return
    try:
        now_iso = _dt.now(_tz.utc).isoformat()
        await supabase_patch(
            f"/rest/v1/awaiting_replies?gmail_thread_id=eq.{gmail_thread_id}&status=eq.active",
            {"status": "replied", "replied_at": now_iso, "updated_at": now_iso},
        )
    except Exception as exc:
        print(f"[silence] mark-replied failed for thread {gmail_thread_id}: {repr(exc)}")


async def mark_awaiting_replies_replied_on_reply(gmail_thread_id: str) -> None:
    """Convenience wrapper for the inbound-reply hook."""
    await supabase_mark_awaiting_reply_replied(gmail_thread_id)


async def supabase_get_stale_awaiting_replies_for_mailbox(
    mailbox_id: str,
    threshold_days: int,
) -> list[dict[str, Any]]:
    """
    Active, non-suppressed, not-yet-nudged rows whose last_user_sent_at is older
    than (now - threshold_days).
    """
    try:
        cutoff = quote((_dt.now(_tz.utc) - _td(days=threshold_days)).isoformat(), safe='')
        rows = await supabase_get(
            f"/rest/v1/awaiting_replies"
            f"?mailbox_id=eq.{mailbox_id}"
            f"&status=eq.active"
            f"&nudge_suppressed=eq.false"
            f"&nudge_sent_at=is.null"
            f"&last_user_sent_at=lte.{cutoff}"
            f"&order=last_user_sent_at.asc"
            f"&select=*"
        )
        return rows if isinstance(rows, list) else []
    except Exception as exc:
        print(f"[silence] fetch stale awaiting_replies failed for mailbox {mailbox_id}: {repr(exc)}")
        return []


async def supabase_mark_awaiting_reply_nudged(
    awaiting_reply_id: str,
    gmail_draft_id: str | None,
) -> None:
    try:
        now_iso = _dt.now(_tz.utc).isoformat()
        await supabase_patch(
            f"/rest/v1/awaiting_replies?id=eq.{awaiting_reply_id}",
            {
                "nudge_sent_at": now_iso,
                "nudge_count": 1,
                "nudge_draft_id": gmail_draft_id,
                "updated_at": now_iso,
            },
        )
    except Exception as exc:
        print(f"[silence] mark-nudged failed for awaiting_reply {awaiting_reply_id}: {repr(exc)}")


async def get_mailbox_silence_config(mailbox_id: str) -> dict[str, Any]:
    try:
        rows = await supabase_get(
            f"/rest/v1/mailboxes?id=eq.{mailbox_id}"
            f"&select=silence_radar_enabled,silence_threshold_days,email_address"
        )
        return rows[0] if rows else {}
    except Exception as exc:
        print(f"[silence] get config failed for mailbox {mailbox_id}: {repr(exc)}")
        return {}


async def generate_silence_nudge_draft(
    user_id: str,
    row: dict[str, Any],
    days_silent: int,
) -> str:
    """Short, friendly bump reply in user's voice."""
    recipient_name = (row.get("recipient_name") or "").strip()
    recipient_email = (row.get("recipient_email") or "").strip()
    subject = row.get("subject") or "Korte check"
    salutation_target = recipient_name or recipient_email or "daar"

    pseudo_body = (
        f"[INTERN NUDGE-CONTEXT — dit is GEEN echte inkomende e-mail]\n"
        f"De gebruiker heeft {days_silent} dag(en) geleden gemaild naar {salutation_target} "
        f"over '{subject}'.\n"
        f"Er is nog geen reactie binnengekomen.\n\n"
        f"Schrijf een korte, vriendelijke herinnering (bump).\n"
        f"Richtlijnen voor deze draft:\n"
        f"- Niet verwijtend; ga uit van een drukke inbox aan de andere kant.\n"
        f"- Verwijs kort naar de eerdere mail.\n"
        f"- Vraag of er iets onduidelijk is of wanneer een reactie te verwachten is.\n"
        f"- Max 3-5 zinnen.\n"
        f"- Geen onderwerpregel."
    )
    try:
        reply = await generate_ai_reply(
            user_id=user_id,
            subject=subject,
            sender=recipient_email or recipient_name or None,
            body_text=pseudo_body,
        )
        return reply or ""
    except Exception as exc:
        print(f"[silence] generate_silence_nudge_draft failed: {repr(exc)}")
        fallback = (
            f"Hoi {recipient_name or ''},\n\n"
            f"Ik wilde m'n eerdere mail even bij je onder de aandacht brengen. "
            f"Weet je al wanneer je hier een reactie op kunt geven? "
            f"Laat het gerust weten als er iets onduidelijk is.\n\n"
            f"Groet"
        )
        return fallback.strip()


async def process_silence_radar_for_mailbox(email: str) -> dict[str, Any]:
    """
    Scans awaiting_replies that have gone silent past threshold. For each:
      1. Verifies thread still has "Waiting On Reply" label in Gmail.
      2. Escalates thread label -> Priority.
      3. Generates a threaded Gmail draft (no auto-send).
      4. Persists nudge_sent_at on the awaiting_reply row.
    Idempotent — nudge_sent_at IS NULL filter prevents re-nudging.
    """
    if not SILENCE_RADAR_ENABLED_GLOBAL:
        return {"status": "disabled_global", "nudges_created": 0}

    user = await supabase_get_user_by_email(email)
    if not user:
        return {"status": "no_user", "nudges_created": 0}

    mailbox = await supabase_get_mailbox_by_user_and_email(user["id"], email)
    if not mailbox:
        return {"status": "no_mailbox", "nudges_created": 0}

    config = await get_mailbox_silence_config(mailbox["id"])
    if not config.get("silence_radar_enabled", True):
        return {"status": "disabled_for_mailbox", "nudges_created": 0}

    threshold_days = int(
        config.get("silence_threshold_days") or SILENCE_RADAR_DEFAULT_THRESHOLD_DAYS
    )
    stale = await supabase_get_stale_awaiting_replies_for_mailbox(
        mailbox_id=mailbox["id"], threshold_days=threshold_days,
    )
    if not stale:
        return {"status": "no_stale", "nudges_created": 0}

    stale = stale[:SILENCE_RADAR_MAX_NUDGES_PER_RUN]

    try:
        labels = await get_all_gmail_labels(user["id"])
        label_name_to_id = {
            name: lbl["id"] for name, lbl in labels.items() if isinstance(lbl, dict) and "id" in lbl
        }
    except Exception as exc:
        print(f"[silence] label fetch failed for {email}: {repr(exc)}")
        label_name_to_id = {}

    waiting_label_id = label_name_to_id.get("Waiting On Reply") or label_name_to_id.get(
        "OfficeFlow/Waiting On Reply"
    )

    nudges_created = 0
    errors: list[str] = []

    for row in stale:
        try:
            row_id = row.get("id")
            thread_id = row.get("gmail_thread_id")
            recipient_email = row.get("recipient_email")

            if not row_id or not thread_id or not recipient_email:
                errors.append(f"awaiting_reply {row_id}: missing thread_id or recipient_email")
                continue

            thread_info = await get_thread_message_headers(user["id"], thread_id)
            if not thread_info:
                errors.append(f"awaiting_reply {row_id}: thread fetch failed")
                continue

            # Verify the thread is still genuinely Waiting On Reply.
            current_label_ids = thread_info.get("current_label_ids") or set()
            if waiting_label_id and waiting_label_id not in current_label_ids:
                await supabase_cancel_awaiting_reply_for_thread(thread_id)
                continue

            # Compute days silent for prompt.
            try:
                last_sent_raw = row.get("last_user_sent_at") or ""
                last_sent_dt = _dt.fromisoformat(last_sent_raw.replace("Z", "+00:00"))
                days_silent = max(1, (_dt.now(_tz.utc) - last_sent_dt).days)
            except Exception:
                days_silent = threshold_days

            if "Priority" in label_name_to_id:
                try:
                    await sync_thread_status(
                        user_id=user["id"],
                        thread_id=thread_id,
                        current_message_id=thread_info.get("gmail_message_id") or "",
                        current_label_ids=current_label_ids,
                        label_name_to_id=label_name_to_id,
                        target_label_name="Priority",
                    )
                except Exception as exc:
                    print(f"[silence] escalate label failed for awaiting_reply {row_id}: {repr(exc)}")

            nudge_body = await generate_silence_nudge_draft(
                user_id=user["id"], row=row, days_silent=days_silent,
            )
            if not nudge_body or len(nudge_body.strip()) < 10:
                errors.append(f"awaiting_reply {row_id}: empty nudge body")
                continue

            draft_result = await create_gmail_threaded_draft(
                user_id=user["id"],
                to_email=recipient_email,
                subject=thread_info.get("subject") or row.get("subject") or "Korte check",
                body=nudge_body,
                thread_id=thread_id,
                original_message_id_header=thread_info.get("message_id_header"),
                references_header=thread_info.get("references_header"),
            )
            gmail_draft_id = (draft_result or {}).get("id")

            await supabase_mark_awaiting_reply_nudged(
                awaiting_reply_id=row_id,
                gmail_draft_id=gmail_draft_id,
            )
            nudges_created += 1
        except Exception as exc:
            errors.append(f"awaiting_reply {row.get('id')}: {repr(exc)}")
            print(f"[silence] per-row failure: {repr(exc)}")

    return {
        "status": "ok",
        "email": email,
        "candidates": len(stale),
        "nudges_created": nudges_created,
        "threshold_days": threshold_days,
        "errors": errors,
    }


# ---------------------------------------------------------------------------
# Internal / testing endpoints for the silence radar
# ---------------------------------------------------------------------------

@app.post("/internal/silence-radar/run")
async def http_run_silence_radar(
    email: str = Body(..., embed=True),
    _: None = Depends(require_internal_secret),
):
    """Manually trigger the silence radar for one mailbox. Use for QA/testing."""
    return await process_silence_radar_for_mailbox(email)


@app.post("/internal/silence-radar/suppress")
async def http_suppress_awaiting_reply_nudge(
    awaiting_reply_id: str = Body(..., embed=True),
    _: None = Depends(require_internal_secret),
):
    """User can suppress nudging for a specific awaiting_reply (future UI hook)."""
    try:
        updated = await supabase_patch(
            f"/rest/v1/awaiting_replies?id=eq.{awaiting_reply_id}",
            {"nudge_suppressed": True, "updated_at": _dt.now(_tz.utc).isoformat()},
            prefer="return=representation",
        )
        return {"status": "ok", "awaiting_reply": updated[0] if isinstance(updated, list) and updated else None}
    except Exception as exc:
        return {"status": "error", "error": repr(exc)}


@app.post("/internal/silence-radar/reset")
async def http_reset_awaiting_reply_nudge(
    awaiting_reply_id: str = Body(..., embed=True),
    _: None = Depends(require_internal_secret),
):
    """Clear nudge state so the awaiting_reply can be nudged again on next run."""
    try:
        updated = await supabase_patch(
            f"/rest/v1/awaiting_replies?id=eq.{awaiting_reply_id}",
            {
                "nudge_sent_at": None,
                "nudge_count": 0,
                "nudge_draft_id": None,
                "nudge_suppressed": False,
                "status": "active",
                "replied_at": None,
                "updated_at": _dt.now(_tz.utc).isoformat(),
            },
            prefer="return=representation",
        )
        return {"status": "ok", "awaiting_reply": updated[0] if isinstance(updated, list) and updated else None}
    except Exception as exc:
        return {"status": "error", "error": repr(exc)}


# ---------------------------------------------------------------------------
# SECTION I — RADAR DASHBOARD API (auth'd endpoints used by the user dashboard)
# ---------------------------------------------------------------------------

async def _radar_default_settings() -> dict[str, Any]:
    return {
        "radar_enabled": True,
        "radar_grace_days": 0,
        "silence_radar_enabled": True,
        "silence_threshold_days": SILENCE_RADAR_DEFAULT_THRESHOLD_DAYS,
        "last_run_at": None,
    }


@app.get("/api/teams/me")
async def api_get_my_teams(user: dict[str, Any] = Depends(get_current_user)):
    """Geeft alle teams terug waar de ingelogde user lid van is."""
    from app import teams as teams_module
    return await teams_module.list_teams_for_user(user["id"])


class _TeamInvitePayload(BaseModel):
    team_id: str
    email: str
    role: str = "member"


@app.post("/api/teams/invite")
async def api_invite_to_team(
    body: _TeamInvitePayload,
    user: dict[str, Any] = Depends(get_current_user),
):
    """Nodig een e-mailadres uit voor het opgegeven team. Alleen admin."""
    from app import teams as teams_module
    return await teams_module.invite_member_to_team(
        admin_user_id=user["id"],
        team_id=body.team_id,
        invite_email=body.email,
        role=body.role,
    )


@app.delete("/api/teams/{team_id}/members/{target_user_id}")
async def api_remove_team_member(
    team_id: str,
    target_user_id: str,
    user: dict[str, Any] = Depends(get_current_user),
):
    """Verwijder een lid uit een team. Alleen admin. Owner kan niet verwijderd worden."""
    from app import teams as teams_module
    return await teams_module.remove_member_from_team(
        admin_user_id=user["id"],
        team_id=team_id,
        target_user_id=target_user_id,
    )


class _RoleChangePayload(BaseModel):
    role: str


@app.patch("/api/teams/{team_id}/members/{target_user_id}/role")
async def api_change_team_member_role(
    team_id: str,
    target_user_id: str,
    body: _RoleChangePayload,
    user: dict[str, Any] = Depends(get_current_user),
):
    """Wijzig role van een team-lid (admin / member). Alleen admin."""
    from app import teams as teams_module
    return await teams_module.change_member_role(
        admin_user_id=user["id"],
        team_id=team_id,
        target_user_id=target_user_id,
        new_role=body.role,
    )


@app.post("/api/teams/{team_id}/billing-portal")
async def api_team_billing_portal(
    team_id: str,
    user: dict[str, Any] = Depends(get_current_user),
):
    """Genereer een Stripe Customer Portal URL voor team-billing beheer.
    Alleen admin. Klant kan via de portal plan wijzigen, payment method updaten, cancelen."""
    from app import teams as teams_module
    import stripe as stripe_sdk

    # Check admin
    membership_rows = await supabase_get(
        f"/rest/v1/team_members?team_id=eq.{quote(team_id, safe='')}"
        f"&user_id=eq.{quote(user['id'], safe='')}&select=role"
    )
    if not isinstance(membership_rows, list) or not membership_rows:
        raise HTTPException(status_code=403, detail="Geen lid van dit team.")
    if membership_rows[0].get("role") != "admin":
        raise HTTPException(status_code=403, detail="Alleen admin kan billing beheren.")

    # Haal team op voor stripe_customer_id
    team_rows = await supabase_get(
        f"/rest/v1/teams?id=eq.{quote(team_id, safe='')}&select=stripe_customer_id"
    )
    if not isinstance(team_rows, list) or not team_rows:
        raise HTTPException(status_code=404, detail="Team niet gevonden.")
    customer_id = team_rows[0].get("stripe_customer_id")
    if not customer_id:
        raise HTTPException(status_code=400, detail="Team heeft geen Stripe-customer gekoppeld.")

    stripe_secret = os.getenv("STRIPE_SECRET_KEY")
    if not stripe_secret:
        raise HTTPException(status_code=500, detail="STRIPE_SECRET_KEY ontbreekt.")
    stripe_sdk.api_key = stripe_secret

    try:
        session = stripe_sdk.billing_portal.Session.create(
            customer=customer_id,
            return_url="https://officeflowcompany.com/team.html",
        )
        return {"url": session.url}
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Stripe portal error: {str(exc)}")


@app.get("/api/teams/{team_id}/stats")
async def api_team_stats(
    team_id: str,
    user: dict[str, Any] = Depends(get_current_user),
):
    """Aggregated stats voor alle mailboxen in dit team (laatste 7 dagen)."""
    # Check lid
    membership_rows = await supabase_get(
        f"/rest/v1/team_members?team_id=eq.{quote(team_id, safe='')}"
        f"&user_id=eq.{quote(user['id'], safe='')}&select=role"
    )
    if not isinstance(membership_rows, list) or not membership_rows:
        raise HTTPException(status_code=403, detail="Geen lid van dit team.")

    # Haal alle mailboxen in dit team
    mailbox_rows = await supabase_get(
        f"/rest/v1/mailboxes?team_id=eq.{quote(team_id, safe='')}&select=id,email_address,user_id,status"
    )
    mailboxes = mailbox_rows if isinstance(mailbox_rows, list) else []

    # 7-dagen cutoff
    now_utc = datetime.now(timezone.utc)
    cutoff_iso = (now_utc - timedelta(days=7)).isoformat()

    per_mailbox = []
    total_mails = 0
    total_drafts = 0

    for mb in mailboxes:
        mb_id = mb.get("id")
        if not mb_id:
            continue

        # Mails geclassificeerd laatste 7 dagen
        mails_count = 0
        drafts_count = 0
        try:
            mails_count = await supabase_get_count(
                f"/rest/v1/messages?mailbox_id=eq.{quote(mb_id, safe='')}"
                f"&created_at=gte.{quote(cutoff_iso, safe='')}"
            )
        except Exception:
            pass
        try:
            drafts_count = await supabase_get_count(
                f"/rest/v1/drafts?mailbox_id=eq.{quote(mb_id, safe='')}"
                f"&created_at=gte.{quote(cutoff_iso, safe='')}"
            )
        except Exception:
            pass

        total_mails += mails_count
        total_drafts += drafts_count

        per_mailbox.append({
            "mailbox_id": mb_id,
            "email_address": mb.get("email_address"),
            "status": mb.get("status"),
            "mails_7d": mails_count,
            "drafts_7d": drafts_count,
        })

    return {
        "team_id": team_id,
        "period_days": 7,
        "total_mails": total_mails,
        "total_drafts": total_drafts,
        "mailbox_count": len(mailboxes),
        "per_mailbox": per_mailbox,
    }


@app.get("/api/radar/overview")
async def radar_overview(user: dict[str, Any] = Depends(get_current_user)):
    """
    Consolidated payload for the dashboard Follow-Up Radar section:
      - settings (enabled flags, threshold, last run)
      - stats (watching counts, nudges this week, total nudges)
      - recent_activity (last 5 nudges across both scopes, newest first)
    Defensive: degrades to zeros + defaults if subqueries fail.
    """
    user_id = user.get("id")
    if not user_id:
        raise HTTPException(status_code=400, detail="User id ontbreekt")

    mailbox = await supabase_get_mailbox_by_user_id(user_id=user_id, provider="gmail")
    if not mailbox:
        return {
            "settings": await _radar_default_settings(),
            "stats": {
                "watching_commitments": 0,
                "watching_awaiting_replies": 0,
                "nudges_this_week": 0,
                "nudges_total": 0,
            },
            "recent_activity": [],
            "_mailbox": False,
        }

    mailbox_id = mailbox["id"]
    settings = {
        "radar_enabled": bool(mailbox.get("radar_enabled", True)),
        "radar_grace_days": int(mailbox.get("radar_grace_days") or 0),
        "silence_radar_enabled": bool(mailbox.get("silence_radar_enabled", True)),
        "silence_threshold_days": int(
            mailbox.get("silence_threshold_days") or SILENCE_RADAR_DEFAULT_THRESHOLD_DAYS
        ),
        "last_run_at": mailbox.get("radar_last_run_at"),
    }

    now = _dt.now(_tz.utc)
    seven_days_iso = quote((now - _td(days=7)).isoformat(), safe='')

    async def _safe_count(url: str) -> int:
        try:
            rows = await supabase_get(url)
            return len(rows) if isinstance(rows, list) else 0
        except Exception as exc:
            print(f"[radar-overview] count failed: {repr(exc)}")
            return 0

    watching_commitments = await _safe_count(
        f"/rest/v1/commitments?mailbox_id=eq.{mailbox_id}"
        f"&status=eq.active&nudge_sent_at=is.null&nudge_suppressed=eq.false&select=id"
    )
    watching_awaiting = await _safe_count(
        f"/rest/v1/awaiting_replies?mailbox_id=eq.{mailbox_id}"
        f"&status=eq.active&nudge_sent_at=is.null&nudge_suppressed=eq.false&select=id"
    )
    nudges_week_commit = await _safe_count(
        f"/rest/v1/commitments?mailbox_id=eq.{mailbox_id}"
        f"&nudge_sent_at=gte.{seven_days_iso}&select=id"
    )
    nudges_week_await = await _safe_count(
        f"/rest/v1/awaiting_replies?mailbox_id=eq.{mailbox_id}"
        f"&nudge_sent_at=gte.{seven_days_iso}&select=id"
    )
    nudges_total_commit = await _safe_count(
        f"/rest/v1/commitments?mailbox_id=eq.{mailbox_id}"
        f"&nudge_sent_at=not.is.null&select=id"
    )
    nudges_total_await = await _safe_count(
        f"/rest/v1/awaiting_replies?mailbox_id=eq.{mailbox_id}"
        f"&nudge_sent_at=not.is.null&select=id"
    )

    stats = {
        "watching_commitments": watching_commitments,
        "watching_awaiting_replies": watching_awaiting,
        "nudges_this_week": nudges_week_commit + nudges_week_await,
        "nudges_total": nudges_total_commit + nudges_total_await,
    }

    # Recent activity — last 5 nudges, merged from both tables
    recent: list[dict[str, Any]] = []
    try:
        commit_rows = await supabase_get(
            f"/rest/v1/commitments?mailbox_id=eq.{mailbox_id}"
            f"&nudge_sent_at=not.is.null"
            f"&order=nudge_sent_at.desc&limit=5"
            f"&select=nudge_sent_at,recipient_name,recipient_email,subject,action_text,due_date,gmail_thread_id"
        )
        if isinstance(commit_rows, list):
            for r in commit_rows:
                days_overdue = None
                try:
                    if r.get("due_date") and r.get("nudge_sent_at"):
                        due = _dt.fromisoformat(str(r["due_date"])).date()
                        nudged = _dt.fromisoformat(r["nudge_sent_at"].replace("Z", "+00:00")).date()
                        days_overdue = max(0, (nudged - due).days)
                except Exception:
                    pass
                recent.append({
                    "type": "commitment",
                    "at": r.get("nudge_sent_at"),
                    "recipient_name": r.get("recipient_name"),
                    "recipient_email": r.get("recipient_email"),
                    "subject": r.get("subject") or r.get("action_text"),
                    "action_text": r.get("action_text"),
                    "days_overdue": days_overdue,
                    "gmail_thread_id": r.get("gmail_thread_id"),
                })
    except Exception as exc:
        print(f"[radar-overview] commit recent failed: {repr(exc)}")

    try:
        await_rows = await supabase_get(
            f"/rest/v1/awaiting_replies?mailbox_id=eq.{mailbox_id}"
            f"&nudge_sent_at=not.is.null"
            f"&order=nudge_sent_at.desc&limit=5"
            f"&select=nudge_sent_at,recipient_name,recipient_email,subject,last_user_sent_at,gmail_thread_id"
        )
        if isinstance(await_rows, list):
            for r in await_rows:
                days_silent = None
                try:
                    if r.get("last_user_sent_at") and r.get("nudge_sent_at"):
                        sent = _dt.fromisoformat(r["last_user_sent_at"].replace("Z", "+00:00"))
                        nudged = _dt.fromisoformat(r["nudge_sent_at"].replace("Z", "+00:00"))
                        days_silent = max(1, (nudged - sent).days)
                except Exception:
                    pass
                recent.append({
                    "type": "silence",
                    "at": r.get("nudge_sent_at"),
                    "recipient_name": r.get("recipient_name"),
                    "recipient_email": r.get("recipient_email"),
                    "subject": r.get("subject"),
                    "action_text": None,
                    "days_silent": days_silent,
                    "gmail_thread_id": r.get("gmail_thread_id"),
                })
    except Exception as exc:
        print(f"[radar-overview] await recent failed: {repr(exc)}")

    recent.sort(key=lambda r: r.get("at") or "", reverse=True)
    recent = recent[:5]

    return {
        "settings": settings,
        "stats": stats,
        "recent_activity": recent,
        "_mailbox": True,
    }


@app.patch("/api/radar/settings")
async def radar_update_settings(
    body: dict[str, Any] = Body(default_factory=dict),
    user: dict[str, Any] = Depends(get_current_user),
):
    """
    Update radar settings for the user's mailbox.
    Accepts any subset of:
      radar_enabled (bool), radar_grace_days (int 0-14),
      silence_radar_enabled (bool), silence_threshold_days (int 1-30)
    """
    user_id = user.get("id")
    if not user_id:
        raise HTTPException(status_code=400, detail="User id ontbreekt")

    mailbox = await supabase_get_mailbox_by_user_id(user_id=user_id, provider="gmail")
    if not mailbox:
        raise HTTPException(status_code=404, detail="Geen gekoppelde mailbox")

    patch: dict[str, Any] = {}
    if "radar_enabled" in body:
        patch["radar_enabled"] = bool(body["radar_enabled"])
    if "radar_grace_days" in body:
        try:
            v = int(body["radar_grace_days"])
            patch["radar_grace_days"] = max(0, min(14, v))
        except Exception:
            pass
    if "silence_radar_enabled" in body:
        patch["silence_radar_enabled"] = bool(body["silence_radar_enabled"])
    if "silence_threshold_days" in body:
        try:
            v = int(body["silence_threshold_days"])
            patch["silence_threshold_days"] = max(1, min(30, v))
        except Exception:
            pass

    if not patch:
        return {"status": "noop", "applied": {}}

    try:
        await supabase_patch(
            f"/rest/v1/mailboxes?id=eq.{mailbox['id']}",
            patch,
        )
    except Exception as exc:
        print(f"[radar-settings] patch failed: {repr(exc)}")
        raise HTTPException(status_code=500, detail=f"Update mislukt: {exc}")

    return {"status": "ok", "applied": patch}


@app.post("/api/radar/run-now")
async def radar_run_now(user: dict[str, Any] = Depends(get_current_user)):
    """
    Manually trigger both radar scopes for the user's mailbox.
    Returns combined nudge counts — drafts land in Gmail, never auto-sent.
    """
    user_id = user.get("id")
    if not user_id:
        raise HTTPException(status_code=400, detail="User id ontbreekt")

    mailbox = await supabase_get_mailbox_by_user_id(user_id=user_id, provider="gmail")
    if not mailbox:
        raise HTTPException(status_code=404, detail="Geen gekoppelde mailbox")
    email = mailbox.get("email_address")
    if not email:
        raise HTTPException(status_code=404, detail="Mailbox mist email_address")

    follow_up = await process_follow_up_radar_for_mailbox(email)
    silence = await process_silence_radar_for_mailbox(email)

    return {
        "status": "ok",
        "follow_up": follow_up,
        "silence": silence,
        "total_nudges_created": (
            (follow_up.get("nudges_created") or 0) + (silence.get("nudges_created") or 0)
        ),
    }


# ---------------------------------------------------------------------------
# SECTION I.b — AUTO-ARCHIVE LOW-VALUE MAIL (settings + one-shot cleanup)
# ---------------------------------------------------------------------------

@app.get("/api/archive/settings")
async def archive_get_settings(user: dict[str, Any] = Depends(get_current_user)):
    """Current auto-archive settings + counts of low-value mail still in Inbox."""
    user_id = user.get("id")
    if not user_id:
        raise HTTPException(status_code=400, detail="User id ontbreekt")

    mailbox = await supabase_get_mailbox_by_user_id(user_id=user_id, provider="gmail")
    if not mailbox:
        raise HTTPException(status_code=404, detail="Geen gekoppelde mailbox")

    notif_on = bool(mailbox.get("notification_auto_archive", False))
    counted_labels = sorted(LOW_VALUE_LABELS) + (
        [NOTIFICATION_AUTO_ARCHIVE_LABEL] if notif_on else []
    )

    # Count leftover low-value mail currently still in Inbox (best-effort).
    in_inbox_counts: dict[str, int] = {}
    total_in_inbox = 0
    try:
        for label_name in counted_labels:
            q = f'in:inbox label:"{label_name}"'
            data = await gmail_get_json_for_user(
                user_id=user_id,
                url=f"{GMAIL_API_BASE}/messages",
                params={"q": q, "maxResults": 500},
            )
            n = len(data.get("messages", []) or [])
            in_inbox_counts[label_name] = n
            total_in_inbox += n
    except Exception as exc:
        print(f"[archive-settings] inbox count failed: {repr(exc)}")

    return {
        "status": "ok",
        "auto_archive_low_value": bool(mailbox.get("auto_archive_low_value", True)),
        "notification_auto_archive": notif_on,
        "low_value_labels": sorted(LOW_VALUE_LABELS),
        "notification_label": NOTIFICATION_AUTO_ARCHIVE_LABEL,
        "in_inbox_counts": in_inbox_counts,
        "total_in_inbox": total_in_inbox,
        "trusted_sender_protection": True,
    }


@app.patch("/api/archive/settings")
async def archive_update_settings(
    body: dict[str, Any] = Body(default_factory=dict),
    user: dict[str, Any] = Depends(get_current_user),
):
    """Toggle auto-archive for low-value labels and Notification opt-in."""
    user_id = user.get("id")
    if not user_id:
        raise HTTPException(status_code=400, detail="User id ontbreekt")

    mailbox = await supabase_get_mailbox_by_user_id(user_id=user_id, provider="gmail")
    if not mailbox:
        raise HTTPException(status_code=404, detail="Geen gekoppelde mailbox")

    patch: dict[str, Any] = {}
    if "auto_archive_low_value" in body:
        patch["auto_archive_low_value"] = bool(body["auto_archive_low_value"])
    if "notification_auto_archive" in body:
        patch["notification_auto_archive"] = bool(body["notification_auto_archive"])

    if not patch:
        return {"status": "noop", "applied": {}}

    try:
        await supabase_patch(
            f"/rest/v1/mailboxes?id=eq.{mailbox['id']}",
            patch,
        )
    except Exception as exc:
        print(f"[archive-settings] patch failed: {repr(exc)}")
        raise HTTPException(status_code=500, detail=f"Update mislukt: {exc}")

    return {"status": "ok", "applied": patch}


@app.post("/api/cleanup/low-value")
async def archive_cleanup_low_value(user: dict[str, Any] = Depends(get_current_user)):
    """One-shot sweep: remove INBOX from every existing low-value mail.

    - Does NOT delete anything. Mails stay in All Mail + keep their status label.
    - Trust-list: skip any mail whose sender is a known contact (reply history
      or prior Priority / To Respond / Follow Up label).
    - Notification is only swept when notification_auto_archive is on for
      the mailbox.
    """
    user_id = user.get("id")
    if not user_id:
        raise HTTPException(status_code=400, detail="User id ontbreekt")

    mailbox = await supabase_get_mailbox_by_user_id(user_id=user_id, provider="gmail")
    if not mailbox:
        raise HTTPException(status_code=404, detail="Geen gekoppelde mailbox")

    notif_on = bool(mailbox.get("notification_auto_archive", False))
    sweep_labels = sorted(LOW_VALUE_LABELS) + (
        [NOTIFICATION_AUTO_ARCHIVE_LABEL] if notif_on else []
    )

    per_label: dict[str, int] = {}
    per_label_skipped: dict[str, int] = {}
    total_archived = 0
    total_skipped = 0
    errors = 0
    trust_cache: dict[str, bool] = {}

    for label_name in sweep_labels:
        try:
            q = f'in:inbox label:"{label_name}"'
            data = await gmail_get_json_for_user(
                user_id=user_id,
                url=f"{GMAIL_API_BASE}/messages",
                params={"q": q, "maxResults": 500},
            )
        except Exception as exc:
            print(f"[cleanup-low-value] list {label_name}: {repr(exc)}")
            errors += 1
            per_label[label_name] = 0
            per_label_skipped[label_name] = 0
            continue

        archived = 0
        skipped = 0
        for msg in data.get("messages", []) or []:
            mid = msg.get("id")
            if not mid:
                continue

            # Fetch sender for trust-list check
            try:
                detail = await gmail_get_json_for_user(
                    user_id=user_id,
                    url=f"{GMAIL_API_BASE}/messages/{mid}",
                    params={"format": "metadata", "metadataHeaders": "From"},
                )
                from_header = get_header_value(
                    detail.get("payload", {}).get("headers", []), "From"
                )
                sender_email = extract_email_address(from_header) or ""
            except Exception as exc:
                print(f"[cleanup-low-value] header fetch msg {mid}: {repr(exc)}")
                sender_email = ""

            try:
                trusted = await is_trusted_sender(
                    user_id=user_id,
                    sender_email=sender_email,
                    cache=trust_cache,
                )
            except Exception as exc:
                print(f"[cleanup-low-value] trust check msg {mid}: {repr(exc)}")
                trusted = False

            if trusted:
                skipped += 1
                continue

            try:
                await modify_gmail_message_labels(
                    user_id=user_id,
                    gmail_message_id=mid,
                    remove_label_ids=["INBOX"],
                )
                archived += 1
            except Exception as exc:
                print(f"[cleanup-low-value] strip INBOX msg {mid}: {repr(exc)}")
                errors += 1

        per_label[label_name] = archived
        per_label_skipped[label_name] = skipped
        total_archived += archived
        total_skipped += skipped

    return {
        "status": "ok",
        "archived_per_label": per_label,
        "skipped_trusted_per_label": per_label_skipped,
        "total_archived": total_archived,
        "total_skipped_trusted": total_skipped,
        "errors": errors,
    }


# ---------------------------------------------------------------------------
# SECTION I.c — RECENT DRAFTS (dashboard feed with confidence dot)
# ---------------------------------------------------------------------------

@app.get("/api/recent-drafts")
async def api_recent_drafts(
    limit: int = 20,
    user: dict[str, Any] = Depends(get_current_user),
):
    """Last N drafts OfficeFlow generated for this user, newest first.

    Returns a flat list ready for dashboard rendering. `confidence` may be
    null for drafts created before the feature shipped — the UI renders a
    neutral grey dot in that case.
    """
    user_id = user.get("id")
    if not user_id:
        raise HTTPException(status_code=400, detail="User id ontbreekt")

    # Clamp limit so we never pull more than 50 rows in one call.
    if limit < 1:
        limit = 1
    if limit > 50:
        limit = 50

    # Pull a wider window than `limit` so live-Gmail filtering of stale drafts
    # still leaves enough rows to fill the dashboard list AND we can return
    # an accurate `active_count` for the WACHTEND-badge on the dashboard.
    fetch_limit = 500
    try:
        drafts_rows = await supabase_get(
            "/rest/v1/drafts"
            f"?user_id=eq.{quote(user_id, safe='')}"
            f"&select=id,email_id,subject,confidence,status,gmail_draft_id,created_at"
            f"&order=created_at.desc"
            f"&limit={fetch_limit}"
        )
    except Exception as exc:
        print(f"[recent-drafts] drafts fetch failed: {repr(exc)}")
        raise HTTPException(status_code=500, detail="Kon concepten niet ophalen")

    drafts_rows = drafts_rows or []

    # Drop drafts already marked deleted/sent from a previous round.
    drafts_rows = [
        d for d in drafts_rows
        if (d.get("status") or "").lower() not in ("deleted", "sent")
    ]

    # Live-check: ask Gmail which drafts still exist for this user. One call
    # to /drafts (maxResults=500) is enough to cover any realistic inbox.
    # Drafts the user deleted manually OR sent disappear from this list, so
    # we filter our DB rows accordingly. Best-effort — if Gmail fails, we
    # show what we have.
    valid_gmail_draft_ids: set[str] | None = None
    try:
        gmail_drafts_resp = await gmail_get_json_for_user(
            user_id=user_id,
            url=f"{GMAIL_API_BASE}/drafts",
            params={"maxResults": 500},
        )
        if isinstance(gmail_drafts_resp, dict):
            valid_gmail_draft_ids = {
                d.get("id")
                for d in (gmail_drafts_resp.get("drafts") or [])
                if d.get("id")
            }
    except Exception as exc:
        print(f"[recent-drafts] gmail drafts list failed: {repr(exc)}")

    if valid_gmail_draft_ids is not None:
        fresh_rows: list[dict[str, Any]] = []
        stale_db_ids: list[str] = []
        for draft in drafts_rows:
            gid = draft.get("gmail_draft_id")
            # Drafts without a gmail_draft_id are old/legacy rows — keep them
            # so they don't silently disappear from the UI. Only drafts that
            # had a Gmail-id but are no longer there get filtered out.
            if gid and gid not in valid_gmail_draft_ids:
                if draft.get("id"):
                    stale_db_ids.append(draft["id"])
            else:
                fresh_rows.append(draft)
        drafts_rows = fresh_rows

        # Persist the 'deleted' status so future calls skip these without
        # needing to ask Gmail again.
        if stale_db_ids:
            try:
                in_clause = ",".join(quote(sid, safe="") for sid in stale_db_ids)
                await supabase_patch(
                    f"/rest/v1/drafts?id=in.({in_clause})",
                    {"status": "deleted"},
                )
                print(f"[recent-drafts] marked {len(stale_db_ids)} drafts as deleted")
            except Exception as exc:
                print(f"[recent-drafts] mark-deleted failed: {repr(exc)}")

    # True count of active drafts after Gmail-filter — drives the WACHTEND
    # badge + sidebar counters, regardless of display limit.
    active_count = len(drafts_rows)

    # Cap to requested limit AFTER filtering so the user always gets the
    # freshest N drafts that still exist.
    drafts_rows = drafts_rows[:limit]

    # Pull the email rows we need so we can attach thread/message ids to
    # each draft. One batched query keeps this cheap even for limit=50.
    email_ids = [d.get("email_id") for d in drafts_rows if d.get("email_id")]
    emails_by_id: dict[str, dict[str, Any]] = {}
    if email_ids:
        in_clause = ",".join(quote(eid, safe="") for eid in email_ids)
        try:
            emails_rows = await supabase_get(
                "/rest/v1/emails"
                f"?id=in.({in_clause})"
                f"&select=id,gmail_thread_id,gmail_message_id,subject"
            )
            for row in emails_rows or []:
                emails_by_id[row["id"]] = row
        except Exception as exc:
            print(f"[recent-drafts] emails fetch failed: {repr(exc)}")

    # Use the user's mailbox email in the Gmail URL so the link opens the
    # right account even when the browser has multiple Google sessions.
    # The ?authuser=<email> query parameter is Google's documented way to
    # force a specific account; the /u/0/ segment is a placeholder.
    mailbox = await supabase_get_mailbox_by_user_id(user_id=user_id, provider="gmail")
    mailbox_email = (mailbox or {}).get("email_address") or ""
    auth_qs = f"?authuser={quote(mailbox_email, safe='')}" if mailbox_email else ""

    items: list[dict[str, Any]] = []
    for draft in drafts_rows:
        email_row = emails_by_id.get(draft.get("email_id") or "") or {}
        thread_id = email_row.get("gmail_thread_id")
        # Gmail web URL to open the thread, scoped to the connected mailbox
        # via authuser so the browser routes to the right Google account.
        gmail_thread_url = (
            f"https://mail.google.com/mail/u/0/{auth_qs}#all/{thread_id}"
            if thread_id else None
        )
        items.append({
            "id": draft.get("id"),
            "email_id": draft.get("email_id"),
            "subject": draft.get("subject") or email_row.get("subject") or "(geen onderwerp)",
            "confidence": draft.get("confidence"),
            "status": draft.get("status"),
            "gmail_draft_id": draft.get("gmail_draft_id"),
            "gmail_thread_id": thread_id,
            "gmail_thread_url": gmail_thread_url,
            "created_at": draft.get("created_at"),
        })

    # Quick counts so the dashboard can show a header like "2 check · 1 review".
    confidence_counts = {"high": 0, "medium": 0, "low": 0, "unknown": 0}
    for item in items:
        bucket = item["confidence"] if item["confidence"] in ("high", "medium", "low") else "unknown"
        confidence_counts[bucket] += 1

    return {
        "status": "ok",
        "items": items,
        "total": len(items),
        "active_count": active_count,
        "confidence_counts": confidence_counts,
    }


# ---------------------------------------------------------------------------
# SECTION I.d — FIRST-RUN WOW MOMENT (one-time welcome modal stats)
# ---------------------------------------------------------------------------

# Conservative estimate: average time saved per AI-drafted reply = 3 minutes
# (reading + thinking + composing) and per classified mail = 15s (triage).
TIME_SAVED_PER_DRAFT_MIN = 3
TIME_SAVED_PER_CLASSIFIED_SEC = 15


def _format_time_saved(total_minutes: float) -> str:
    """Pretty string like '3u 20m' / '47m' / '12s'."""
    if total_minutes < 1:
        secs = max(1, int(round(total_minutes * 60)))
        return f"{secs}s"
    total = int(round(total_minutes))
    hours = total // 60
    mins = total % 60
    if hours and mins:
        return f"{hours}u {mins}m"
    if hours:
        return f"{hours}u"
    return f"{mins}m"


@app.get("/api/first-run-stats")
async def api_first_run_stats(user: dict[str, Any] = Depends(get_current_user)):
    """Return the hero numbers shown in the welcome modal.

    - emails_classified: inserts into public.emails for this user.
    - drafts_created: drafts this user has in public.drafts.
    - time_saved_minutes: estimated minutes saved (drafts × 3m + classified × 15s).
    - seen: true if the user has already dismissed the modal.
    """
    user_id = user.get("id")
    if not user_id:
        raise HTTPException(status_code=400, detail="User id ontbreekt")

    # Has the user already seen the modal?
    onboarding = await supabase_get_onboarding_state(user_id)
    seen = bool(onboarding and onboarding.get("first_run_seen_at"))

    # Count classified emails (head+count trick to avoid downloading rows).
    try:
        emails_count_resp = await supabase_get_count(
            f"/rest/v1/emails?user_id=eq.{quote(user_id, safe='')}&select=id"
        )
    except Exception as exc:
        print(f"[first-run] emails count failed: {repr(exc)}")
        emails_count_resp = 0

    try:
        drafts_count_resp = await supabase_get_count(
            f"/rest/v1/drafts?user_id=eq.{quote(user_id, safe='')}&select=id"
        )
    except Exception as exc:
        print(f"[first-run] drafts count failed: {repr(exc)}")
        drafts_count_resp = 0

    emails_classified = int(emails_count_resp or 0)
    drafts_created = int(drafts_count_resp or 0)

    minutes_saved = (
        drafts_created * TIME_SAVED_PER_DRAFT_MIN
        + emails_classified * (TIME_SAVED_PER_CLASSIFIED_SEC / 60.0)
    )

    return {
        "seen": seen,
        "emails_classified": emails_classified,
        "drafts_created": drafts_created,
        "time_saved_minutes": round(minutes_saved, 1),
        "time_saved_human": _format_time_saved(minutes_saved),
    }


@app.post("/api/first-run-seen")
async def api_first_run_seen(user: dict[str, Any] = Depends(get_current_user)):
    """Mark the welcome modal as dismissed. Idempotent — safe to call twice."""
    user_id = user.get("id")
    if not user_id:
        raise HTTPException(status_code=400, detail="User id ontbreekt")
    await supabase_mark_first_run_seen(user_id)
    return {"status": "ok"}


# ---------------------------------------------------------------------------
# SECTION I.e — WEEKLY RECAP EMAIL (Saturday 09:00 local)
# ---------------------------------------------------------------------------

WEEKLY_RECAP_SEND_HOUR = 9   # Saturday local time
WEEKLY_RECAP_SEND_WEEKDAY = 5  # Monday=0 … Saturday=5
WEEKLY_RECAP_ENABLED_GLOBAL = os.getenv("WEEKLY_RECAP_ENABLED", "true").lower() == "true"


async def get_weekly_recap_stats(user_id: str) -> dict[str, Any]:
    """Aggregate last 7 days: emails processed, drafts created, busiest day.

    All counts are scoped to `user_id`. Returns zero-filled keys when the DB
    has nothing so the email template never has to handle nulls.
    """
    since_iso = (datetime.now(timezone.utc) - timedelta(days=7)).isoformat()

    emails_rows: list[dict[str, Any]] = []
    drafts_count = 0

    try:
        emails_rows = await supabase_get(
            "/rest/v1/emails"
            f"?user_id=eq.{quote(user_id, safe='')}"
            f"&created_at=gte.{quote(since_iso, safe='')}"
            f"&select=created_at"
            f"&limit=5000"
        ) or []
    except Exception as exc:
        print(f"[weekly-recap] emails fetch failed: {repr(exc)}")

    try:
        drafts_count = await supabase_get_count(
            "/rest/v1/drafts"
            f"?user_id=eq.{quote(user_id, safe='')}"
            f"&created_at=gte.{quote(since_iso, safe='')}"
            f"&select=id"
        ) or 0
    except Exception as exc:
        print(f"[weekly-recap] drafts count failed: {repr(exc)}")

    # Bucket emails by weekday to find busiest day.
    weekday_counts: dict[int, int] = {i: 0 for i in range(7)}
    for row in emails_rows:
        raw = row.get("created_at")
        if not raw:
            continue
        try:
            iso = raw.replace("Z", "+00:00") if isinstance(raw, str) else raw
            dt = datetime.fromisoformat(iso)
            if dt.tzinfo is None:
                dt = dt.replace(tzinfo=timezone.utc)
            weekday_counts[dt.astimezone(ZoneInfo(BRIEFING_DEFAULT_TIMEZONE)).weekday()] += 1
        except Exception:
            continue

    busiest_day_idx = max(weekday_counts, key=weekday_counts.get) if weekday_counts else 0
    busiest_day_count = weekday_counts[busiest_day_idx]
    weekday_names_nl = ["maandag", "dinsdag", "woensdag", "donderdag", "vrijdag", "zaterdag", "zondag"]
    busiest_day_label = weekday_names_nl[busiest_day_idx]

    emails_count = len(emails_rows)
    minutes_saved = (
        drafts_count * TIME_SAVED_PER_DRAFT_MIN
        + emails_count * (TIME_SAVED_PER_CLASSIFIED_SEC / 60.0)
    )

    return {
        "emails_processed": emails_count,
        "drafts_created": drafts_count,
        "time_saved_minutes": round(minutes_saved, 1),
        "time_saved_human": _format_time_saved(minutes_saved),
        "busiest_day_label": busiest_day_label,
        "busiest_day_count": busiest_day_count,
    }


def build_weekly_recap_html(first_name: str, stats: dict[str, Any], user_email: str) -> tuple[str, str]:
    """Return (subject, html_body) for the weekly recap email."""
    hello = f"Hoi {first_name}," if first_name else "Hoi,"
    subject = f"OfficeFlow weekrecap — je bespaarde {stats['time_saved_human']}"

    busiest_line = (
        f"Je drukste dag was <strong>{stats['busiest_day_label']}</strong> "
        f"met <strong>{stats['busiest_day_count']}</strong> binnenkomende mails."
        if stats["busiest_day_count"] > 0
        else "Rustige week — weinig mail binnengekomen."
    )

    html = f"""
    <html><body style="margin:0;padding:0;background:#f6f7f9;font-family:-apple-system,Segoe UI,Roboto,sans-serif;color:#111827;">
      <div style="max-width:560px;margin:0 auto;padding:32px 24px;">
        <div style="background:#fff;border-radius:16px;padding:32px 28px;box-shadow:0 1px 3px rgba(0,0,0,.04);border:1px solid #eef0f4;">
          <div style="font-size:11px;letter-spacing:.12em;text-transform:uppercase;color:#9ca3af;font-weight:700;">Weekrecap</div>
          <h1 style="font-size:22px;margin:6px 0 4px 0;color:#111827;">{hello}</h1>
          <p style="margin:0 0 22px 0;color:#4b5563;font-size:15px;line-height:1.5;">
            Hier is wat OfficeFlow deze week voor je deed.
          </p>

          <div style="display:flex;gap:10px;margin-bottom:20px;flex-wrap:wrap;">
            <div style="flex:1;min-width:140px;background:#fff7ed;border:1px solid #fed7aa;border-radius:12px;padding:14px 16px;">
              <div style="font-size:12px;color:#9a3412;font-weight:600;margin-bottom:4px;">Mails verwerkt</div>
              <div style="font-size:24px;font-weight:800;color:#7c2d12;">{stats['emails_processed']}</div>
            </div>
            <div style="flex:1;min-width:140px;background:#ecfdf5;border:1px solid #bbf7d0;border-radius:12px;padding:14px 16px;">
              <div style="font-size:12px;color:#065f46;font-weight:600;margin-bottom:4px;">Concepten klaargezet</div>
              <div style="font-size:24px;font-weight:800;color:#064e3b;">{stats['drafts_created']}</div>
            </div>
            <div style="flex:1;min-width:140px;background:#eff6ff;border:1px solid #bfdbfe;border-radius:12px;padding:14px 16px;">
              <div style="font-size:12px;color:#1e40af;font-weight:600;margin-bottom:4px;">Tijd bespaard</div>
              <div style="font-size:24px;font-weight:800;color:#1e3a8a;">{stats['time_saved_human']}</div>
            </div>
          </div>

          <p style="margin:0 0 18px 0;color:#374151;font-size:14px;line-height:1.55;">
            {busiest_line}
          </p>

          <a href="https://officeflowcompany.com/dashboard.html"
             style="display:inline-block;background:#ea580c;color:#fff;text-decoration:none;padding:12px 20px;border-radius:10px;font-weight:600;font-size:14px;">
            Bekijk je dashboard
          </a>

          <p style="margin:26px 0 0 0;color:#9ca3af;font-size:12px;line-height:1.5;">
            Je ontvangt deze mail elke zaterdag. Uitschakelen kan in je
            <a href="https://officeflowcompany.com/dashboard.html#account" style="color:#ea580c;text-decoration:none;">account-instellingen</a>.
          </p>
        </div>
        <div style="text-align:center;font-size:11px;color:#9ca3af;margin-top:14px;">
          OfficeFlow · {user_email}
        </div>
      </div>
    </body></html>
    """
    return subject, html


async def send_weekly_recap_for_mailbox(email: str) -> dict[str, Any]:
    """Build and send the Saturday weekly recap to the given mailbox."""
    context = await get_gmail_context_by_email(email)
    user = context["user"]
    mailbox = context["mailbox"]

    stats = await get_weekly_recap_stats(user["id"])

    first_name = (user.get("full_name") or user.get("email") or "").split(" ")[0].split("@")[0]
    subject, html_body = build_weekly_recap_html(first_name, stats, email)

    raw = build_briefing_raw_email(to_email=email, subject=subject, html_body=html_body)

    send_response = await gmail_post_json_for_user(
        user_id=user["id"],
        url=f"{GMAIL_API_BASE}/messages/send",
        payload={"raw": raw},
    )

    try:
        await supabase_patch(
            f"/rest/v1/mailboxes?id=eq.{mailbox['id']}",
            {"weekly_recap_last_sent_at": utc_now_iso()},
        )
    except Exception as exc:
        print(f"[weekly-recap] timestamp patch failed: {repr(exc)}")

    return {
        "status": "sent",
        "email": email,
        "subject": subject,
        "gmail_message_id": send_response.get("id"),
        "stats": stats,
    }


async def weekly_recap_loop():
    """Every 2 min check each mailbox. Send Saturday 09:00 local, once per ISO week."""
    await asyncio.sleep(45)
    while True:
        try:
            if not WEEKLY_RECAP_ENABLED_GLOBAL:
                await asyncio.sleep(120)
                continue

            mailboxes = await get_all_active_mailboxes()
            for mailbox in mailboxes:
                try:
                    email = mailbox.get("email_address")
                    if not email:
                        continue
                    if not mailbox.get("weekly_recap_enabled", True):
                        continue

                    tz_name = mailbox.get("briefing_timezone") or BRIEFING_DEFAULT_TIMEZONE
                    try:
                        now_local = _dt.now(ZoneInfo(tz_name))
                    except Exception:
                        now_local = _dt.now(ZoneInfo(BRIEFING_DEFAULT_TIMEZONE))

                    # Only Saturday
                    if now_local.weekday() != WEEKLY_RECAP_SEND_WEEKDAY:
                        continue
                    # Only after 09:00 local
                    if now_local.hour < WEEKLY_RECAP_SEND_HOUR:
                        continue

                    # Dedupe within the same ISO week
                    last_sent_raw = mailbox.get("weekly_recap_last_sent_at")
                    if last_sent_raw:
                        try:
                            last_sent = _dt.fromisoformat(last_sent_raw.replace("Z", "+00:00"))
                            last_local = last_sent.astimezone(now_local.tzinfo)
                            if last_local.isocalendar()[:2] == now_local.isocalendar()[:2]:
                                continue  # already sent this week
                        except Exception:
                            pass

                    # Window: from 09:00 until end of Saturday
                    target = now_local.replace(hour=WEEKLY_RECAP_SEND_HOUR, minute=0, second=0, microsecond=0)
                    if now_local < target:
                        continue

                    print(f"[weekly-recap] Sending to {email}...")
                    result = await send_weekly_recap_for_mailbox(email)
                    print(f"[weekly-recap] Sent to {email}: {result.get('gmail_message_id')}")
                except Exception as exc:
                    print(f"[weekly-recap] Failed for {mailbox.get('email_address')}: {repr(exc)}")
        except Exception as exc:
            print(f"[weekly-recap] Loop error: {repr(exc)}")

        await asyncio.sleep(120)


@app.get("/api/preferences/weekly-recap")
async def weekly_recap_get_settings(user: dict[str, Any] = Depends(get_current_user)):
    """Return whether the weekly recap email is enabled for the user's mailbox."""
    user_id = user.get("id")
    if not user_id:
        raise HTTPException(status_code=400, detail="User id ontbreekt")

    mailbox = await supabase_get_mailbox_by_user_id(user_id=user_id, provider="gmail")
    if not mailbox:
        raise HTTPException(status_code=404, detail="Geen gekoppelde mailbox")

    return {
        "status": "ok",
        "enabled": bool(mailbox.get("weekly_recap_enabled", True)),
        "last_sent_at": mailbox.get("weekly_recap_last_sent_at"),
    }


@app.patch("/api/preferences/weekly-recap")
async def weekly_recap_update_settings(
    body: dict[str, Any] = Body(default_factory=dict),
    user: dict[str, Any] = Depends(get_current_user),
):
    """Toggle the weekly recap email on/off."""
    user_id = user.get("id")
    if not user_id:
        raise HTTPException(status_code=400, detail="User id ontbreekt")

    if "enabled" not in body:
        return {"status": "noop"}

    mailbox = await supabase_get_mailbox_by_user_id(user_id=user_id, provider="gmail")
    if not mailbox:
        raise HTTPException(status_code=404, detail="Geen gekoppelde mailbox")

    try:
        await supabase_patch(
            f"/rest/v1/mailboxes?id=eq.{mailbox['id']}",
            {"weekly_recap_enabled": bool(body["enabled"])},
        )
    except Exception as exc:
        print(f"[weekly-recap-settings] patch failed: {repr(exc)}")
        raise HTTPException(status_code=500, detail=f"Update mislukt: {exc}")

    return {"status": "ok", "enabled": bool(body["enabled"])}


@app.post("/api/weekly-recap/run-now")
async def http_send_weekly_recap(email: str = Body(..., embed=True)):
    """Manual trigger for testing — sends the recap immediately regardless of day/time.

    Wraps the send in try/except so a failure surfaces as a readable JSON
    payload instead of a generic 500. HTTPExceptions (404/400/403 from
    get_gmail_context_by_email / ensure_user_has_access) are re-raised so
    FastAPI returns their real status code.
    """
    try:
        return await send_weekly_recap_for_mailbox(email)
    except HTTPException:
        raise
    except Exception as exc:
        # Surface the underlying error so we can debug from the client.
        import traceback
        tb = traceback.format_exc()
        print(f"[weekly-recap] manual trigger failed for {email}: {repr(exc)}\n{tb}")
        return {
            "status": "error",
            "email": email,
            "error_type": type(exc).__name__,
            "error_message": str(exc),
        }


# ---------------------------------------------------------------------------
# SECTION F — STARTUP HOOK (add briefing + radar loops to background tasks)
# ---------------------------------------------------------------------------

@app.on_event("startup")
async def start_features_background_tasks():
    print("Features background tasks started (briefing + radar + silence + weekly recap)")
    asyncio.create_task(briefing_loop())
    asyncio.create_task(radar_loop())
    asyncio.create_task(weekly_recap_loop())
